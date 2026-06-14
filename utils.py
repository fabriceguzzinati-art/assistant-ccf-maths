# ============================================================
# utils.py — PARTIE 1 : Imports + Constantes + Filières + Gamification
# ============================================================

import os
import re
import json
import requests
from datetime import datetime
from io import BytesIO
from zoneinfo import ZoneInfo
from collections import defaultdict

import PIL.Image
import google.generativeai as genai
from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Chemin absolu du dossier racine du projet
APP_DIR  = os.path.dirname(os.path.abspath(__file__))
BANQUE_DIR = os.path.join(APP_DIR, "banque")


# ============================================================
# 1. DONNÉES OFFICIELLES — BO MATHS BAC PRO
# ============================================================

NIVEAUX_CATEGORIES = {
    "Collège":        ["6ème", "5ème", "4ème", "3ème"],
    "Lycée Général":  ["2nde", "1ère", "Terminale"],
    "Bac Pro":        ["2nde Pro", "1ère Pro", "Term Pro"],
    "CAP":            ["1ère année CAP", "2ème année CAP"],
    "BTS":            ["BTS 1", "BTS 2"],
}

MATIERES = [
    "Mathématiques", "Français", "Histoire-Géographie", "SVT",
    "Physique-Chimie", "Anglais", "Espagnol", "SES",
    "Philosophie", "Sciences de l'ingénieur", "EMC",
]

CHAPITRES_MATHS_BAC_PRO = {
    "2nde Pro": [
        "Calcul numérique et algébrique",
        "Puissances et notations scientifiques",
        "Proportionnalité et pourcentages",
        "Équations du 1er degré à une inconnue",
        "Inéquations du 1er degré",
        "Géométrie plane — figures usuelles",
        "Trigonométrie dans le triangle rectangle",
        "Vecteurs — notions de base",
        "Notion de fonction — représentation graphique",
        "Fonctions linéaires et affines",
        "Statistiques descriptives — série à une variable",
        "Notion de probabilité — expériences aléatoires",
    ],
    "1ère Pro": [
        "Statistique à deux variables — ajustement affine et coefficient de détermination",
        "Probabilités — événements, tableaux croisés, probabilités conditionnelles",
        "Suites numériques — suites arithmétiques",
        "Résolution graphique d'équations et d'inéquations f(x)=g(x)",
        "Fonctions polynômes de degré 2 — racines, signe, forme factorisée",
        "Fonction dérivée — variations, extremums, fonction inverse",
        "Calculs commerciaux et financiers — intérêts simples, coûts (filières sans physique-chimie)",
        "Géométrie dans l'espace — solides usuels et sections par un plan",
        "Vecteurs du plan — coordonnées, opérations, norme (groupements A et B)",
        "Trigonométrie — cercle trigonométrique, fonctions sinus et cosinus (groupements A et B)",
        "Algorithmique et programmation Python — listes, fonctions, boucles",
        "Automatismes — calcul, grandeurs, lecture graphique",
    ],
    "Term Pro": [
        "Statistiques à deux variables — ajustements non affines, changements de variable",
        "Probabilités — arbres pondérés, formule des probabilités totales, indépendance",
        "Suites géométriques — terme général, sens de variation, somme",
        "Fonctions polynômes de degré 3 — dérivée, variations, extremums",
        "Fonctions exponentielles de base q et logarithme décimal",
        "Calculs commerciaux et financiers — intérêts composés, amortissements (filières sans physique-chimie)",
        "Vecteurs dans l'espace — coordonnées, norme, colinéarité (groupement B)",
        "Trigonométrie — équations, vecteurs de Fresnel (groupement A)",
        "Algorithmique et programmation Python — approfondissement listes et fonctions",
        "Automatismes — probabilités, suites, dérivation, vecteurs",
        "Calcul intégral — primitives, intégrale, aire (programme complémentaire)",
        "Fonctions logarithme népérien et exponentielle de base e (programme complémentaire)",
        "Nombres complexes — forme algébrique et trigonométrique (programme complémentaire)",
        "Produit scalaire de deux vecteurs du plan (programme complémentaire)",
    ],
}

CHAPITRES_PAR_MATIERE_GENERAL = {
    "Mathématiques":        ["Nombres et opérations", "Fractions", "Équations", "Fonctions", "Géométrie", "Statistiques", "Probabilités", "Algorithmique"],
    "Français":             ["Grammaire", "Orthographe", "Conjugaison", "Analyse littéraire", "Argumentation", "Expression écrite", "Oral"],
    "Histoire-Géographie":  ["Antiquité", "Moyen Âge", "Époque Moderne", "Époque Contemporaine", "Géographie de la France", "Géographie mondiale", "EMC"],
    "SVT":                  ["Cellule et génétique", "Évolution", "Corps humain et santé", "Écosystèmes", "Géologie"],
    "Physique-Chimie":      ["Mécanique", "Électricité", "Optique", "Thermodynamique", "Chimie organique", "Chimie des solutions"],
    "Anglais":              ["Compréhension écrite", "Expression écrite", "Compréhension orale", "Expression orale", "Grammaire", "Vocabulaire professionnel"],
    "Espagnol":             ["Compréhension écrite", "Expression écrite", "Grammaire", "Civilisation", "Vocabulaire"],
    "SES":                  ["Économie", "Sociologie", "Science politique", "Mondialisation"],
    "Philosophie":          ["La conscience", "Le langage", "La liberté", "La morale", "La politique", "La vérité", "L'art"],
    "Sciences de l'ingénieur": ["Mécanique", "Électronique", "Informatique industrielle", "Matériaux"],
    "EMC":                  ["Démocratie et citoyenneté", "Droits et libertés", "Laïcité", "Engagement"],
}

# Compatibilité avec l'ancien app_eleve.py
THEMES = NIVEAUX_CATEGORIES
NIVEAUXCATEGORIES = NIVEAUX_CATEGORIES
CHAPITRESMATHSBACPRO = CHAPITRES_MATHS_BAC_PRO
CHAPITRESPARMATIEREGENERAL = CHAPITRES_PAR_MATIERE_GENERAL

# ============================================================
# 2. FILIÈRES PRO
# ============================================================

LISTE_FILIERES = ["ASSP", "MCVB", "MCVA", "AGORA", "Autre (Préciser ci-dessous)"]

CONTEXTES_FILIERES = {
    "ASSP": {
        "nom_complet": "Accompagnement, Soins et Services à la Personne",
        "contextes_maths": [
            "Calcul de doses et de dilutions médicamenteuses",
            "Gestion de plannings d'intervenants à domicile",
            "Statistiques sur données de santé (IMC, fréquence cardiaque)",
            "Calcul de coûts de prise en charge",
            "Lecture et interprétation de graphiques médicaux",
        ],
        "mots_cles": "soins, patients, résidents, personnes âgées, handicap, domicile, EHPAD, pharmacie, hygiène",
    },
    "MCVB": {
        "nom_complet": "Métiers du Commerce et de la Vente — option B (Prospection Clientèle)",
        "contextes_maths": [
            "Calcul de marges, taux de remise et prix de vente",
            "Statistiques de vente et de prospection",
            "Gestion d'un budget commercial",
            "Pourcentages d'évolution du chiffre d'affaires",
            "Représentations graphiques des ventes",
        ],
        "mots_cles": "vente, client, prospection, chiffre d'affaires, commission, remise, catalogue",
    },
    "MCVA": {
        "nom_complet": "Métiers du Commerce et de la Vente — option A (Animation et Gestion de l'Espace Commercial)",
        "contextes_maths": [
            "Calcul de surfaces de vente et d'implantation linéaire",
            "Statistiques de fréquentation et taux de transformation",
            "Gestion des stocks — taux de rotation, coefficient multiplicateur",
            "Calcul de TVA, de prix TTC et HT",
            "Optimisation de l'espace commercial",
        ],
        "mots_cles": "magasin, rayon, linéaire, stock, inventaire, promotion, merchandising, caisse",
    },
    "AGORA": {
        "nom_complet": "Assistance à la Gestion des Organisations et de leurs Activités",
        "contextes_maths": [
            "Calcul de salaires, cotisations et charges sociales",
            "Gestion de budgets d'entreprise",
            "Statistiques sur données RH",
            "Facturation et suivi comptable de base",
            "Calcul d'intérêts pour emprunts professionnels",
        ],
        "mots_cles": "entreprise, comptabilité, salaire, facture, devis, ressources humaines, secrétariat",
    },
}

COMPETENCES_CCF = [
    {
        "nom": "S'approprier",
        "indicateurs": [
            "Rechercher, extraire et organiser l'information.",
            "Traduire des informations, des codages.",
        ],
    },
    {
        "nom": "Analyser / Raisonner",
        "indicateurs": [
            "Émettre des conjectures, formuler des hypothèses.",
            "Proposer, choisir une méthode de résolution ou un protocole expérimental.",
            "Élaborer un algorithme.",
        ],
    },
    {
        "nom": "Réaliser",
        "indicateurs": [
            "Mettre en œuvre une méthode de résolution, des algorithmes ou un protocole expérimental en respectant les règles de sécurité.",
            "Utiliser un modèle, représenter, calculer.",
            "Expérimenter, faire une simulation.",
        ],
    },
    {
        "nom": "Valider",
        "indicateurs": [
            "Exploiter et interpréter des résultats ou des observations de façon critique et argumentée.",
            "Contrôler la vraisemblance d'une conjecture, de la valeur d'une mesure.",
            "Valider un modèle ou une hypothèse.",
            "Mener un raisonnement logique et établir une conclusion.",
        ],
    },
    {
        "nom": "Communiquer",
        "indicateurs": [
            "Rendre compte d'un résultat, à l'oral ou à l'écrit en utilisant des outils et un langage approprié.",
            "Expliquer une démarche.",
        ],
    },
]


# ============================================================
# 3. GAMIFICATION — Constantes
# ============================================================

SEUIL_VALIDATION   = 2
ORDRE_NIVEAUX_DIFF = ["🟢 Débutant", "🟡 Moyen", "🟠 Confirmé", "🔴 Expert"]

MASCOTTES = {
    "🟢 Débutant": {"animal": "🐣", "nom": "Poussin", "couleur": "#22c55e"},
    "🟡 Moyen":    {"animal": "🦊", "nom": "Renard",  "couleur": "#f59e0b"},
    "🟠 Confirmé": {"animal": "🦁", "nom": "Lion",    "couleur": "#f97316"},
    "🔴 Expert":   {"animal": "🐉", "nom": "Dragon",  "couleur": "#ef4444"},
}

MESSAGES_VICTOIRE = {
    "🟢 Débutant": "🐣 Poussin vaincu ! Tu maîtrises les bases — le Renard t'attend !",
    "🟡 Moyen":    "🦊 Renard vaincu ! Tu commences à être redoutable — au Lion !",
    "🟠 Confirmé": "🦁 Lion vaincu ! Tu es vraiment solide — ose affronter le Dragon !",
    "🔴 Expert":   "🐉 Dragon vaincu ! Tu es un expert — bravo, c'est le sommet !",
}

XP_PAR_ACTION = {
    "Exercice-🟢 Débutant": 10,
    "Exercice-🟡 Moyen":    20,
    "Exercice-🟠 Confirmé": 35,
    "Exercice-🔴 Expert":   50,
    "CCF":                  40,
    "Boss":                 80,
}

RANGS = [
    (0,    "🥉", "Apprenti"),
    (100,  "🥈", "Initié"),
    (300,  "🥇", "Confirmé"),
    (600,  "💎", "Expert"),
    (1000, "👑", "Maître"),
]

BADGES = [
    # Débuts
    {"id": "premier_pas",       "emoji": "🌱",  "nom": "Premier Pas",         "description": "Terminer son premier exercice",                          "categorie": "Débuts"},
    {"id": "premiere_victoire", "emoji": "🎯",  "nom": "Première Victoire",   "description": "Obtenir sa première bonne évaluation (😊 ou 🌟)",        "categorie": "Débuts"},
    {"id": "gamer",             "emoji": "🎮",  "nom": "Gamer",               "description": "Terminer un exercice interactif",                        "categorie": "Débuts"},
    # Régularité
    {"id": "en_feu",            "emoji": "🔥",  "nom": "En Feu",              "description": "Atteindre un streak de 3 jours consécutifs",             "categorie": "Régularité"},
    {"id": "flamme",            "emoji": "🔥🔥","nom": "Flamme",              "description": "Atteindre un streak de 7 jours consécutifs",             "categorie": "Régularité"},
    # Performance
    {"id": "sans_faute",        "emoji": "💯",  "nom": "Sans Faute",          "description": "Obtenir 100% sur un exercice interactif",                "categorie": "Performance"},
    {"id": "perfectionniste",   "emoji": "🌟",  "nom": "Perfectionniste",     "description": "Obtenir 5 évaluations « Très bien »",                   "categorie": "Performance"},
    {"id": "assidu",            "emoji": "💪",  "nom": "Assidu",              "description": "Terminer 10 exercices avec bonne évaluation",            "categorie": "Performance"},
    # Boss
    {"id": "premier_boss",      "emoji": "⚔️",  "nom": "Premier Boss",        "description": "Vaincre son premier boss",                               "categorie": "Boss"},
    {"id": "chasseur_dragons",  "emoji": "🐉",  "nom": "Chasseur de Dragons", "description": "Vaincre le boss Expert (Dragon)",                        "categorie": "Boss"},
    # Exploration
    {"id": "explorateur",       "emoji": "🗺️",  "nom": "Explorateur",         "description": "Travailler 3 chapitres différents",                     "categorie": "Exploration"},
    {"id": "encyclopediste",    "emoji": "📚",  "nom": "Encyclopédiste",      "description": "Travailler 5 chapitres différents",                     "categorie": "Exploration"},
    {"id": "candidat_ccf",      "emoji": "📋",  "nom": "Candidat CCF",        "description": "Terminer un sujet CCF",                                  "categorie": "Exploration"},
    {"id": "champion",          "emoji": "🏆",  "nom": "Champion",            "description": "Maîtriser complètement un chapitre (4 niveaux + boss)",  "categorie": "Exploration"},
]


# ============================================================
# 4. HELPERS — Chapitres
# ============================================================

def get_chapitres(matiere: str, niveau: str, categorie: str) -> list:
    if matiere == "Mathématiques" and categorie == "Bac Pro" and niveau in CHAPITRES_MATHS_BAC_PRO:
        return CHAPITRES_MATHS_BAC_PRO[niveau]
    return CHAPITRES_PAR_MATIERE_GENERAL.get(matiere, ["Chapitre général"])


def build_contexte_filiere(filiere: str) -> str:
    if filiere in CONTEXTES_FILIERES:
        ctx    = CONTEXTES_FILIERES[filiere]
        lignes = "\n".join(f"- {c}" for c in ctx["contextes_maths"])
        return f"\n**Filière : {ctx['nom_complet']}**\nUnivers : {ctx['mots_cles']}\nContextes maths :\n{lignes}\n"
    return f"\nFilière : {filiere}\n" if filiere else ""
    
# ============================================================
# utils.py — PARTIE 2 : Prompts + Gemini + Grist + Fonctions métier
# ============================================================

# ── À coller à la suite de la Partie 1 ──────────────────────


# ============================================================
# 5. PROMPTS
# ============================================================

NIVEAUX_DIFFICULTE = {
    "🟢 Débutant": "Questions très guidées, micro-étapes, résultats intermédiaires donnés.",
    "🟡 Moyen":    "Semi-guidé, quelques repères fournis, mise en situation simple.",
    "🟠 Confirmé": "Autonome, mise en situation réaliste, raisonnement attendu.",
    "🔴 Expert":   "Transfert de compétences, données brutes, problème ouvert ambitieux.",
}

SYSTEM_EXERCICES = """\
Tu es un professeur expert en pédagogie différenciée pour lycée professionnel (Bac Pro).
Tes élèves ont un niveau en mathématiques fragile : tes énoncés sont toujours clairs, bienveillants et ancrés dans des contextes professionnels concrets.

En fonction du niveau de difficulté demandé, adapte PRÉCISÉMENT la structure suivante :

━━ DÉBUTANT ━━
- Questions très guidées, découpées en micro-étapes (une opération par question).
- Résultats intermédiaires fournis pour permettre de continuer même en cas d'erreur.
- Vocabulaire ultra-simplifié, aucun terme technique sans définition immédiate.
- Rappel de cours détaillé avec exemple résolu pas à pas.
- Exercice d'application : calcul direct, données déjà extraites.
- Mise en situation : contexte simple, une seule inconnue.
- Pas de problème ouvert — remplacer par une question bilan guidée.

━━ MOYEN ━━
- Questions semi-guidées avec quelques repères (formule rappelée, première étape donnée).
- Rappel de cours synthétique avec un exemple.
- Exercice d'application : 3 questions progressives.
- Mise en situation professionnelle simple avec tableau de données.
- Problème ouvert court (1 question de synthèse guidée).

━━ CONFIRMÉ ━━
- Questions autonomes, aucune aide dans l'énoncé.
- Rappel de cours en points clés uniquement (pas d'exemple résolu).
- Exercice d'application : questions progressives avec barème.
- Mise en situation professionnelle réaliste et complète.
- Problème ouvert avec raisonnement attendu.

━━ EXPERT ━━
- Questions ouvertes sans guidage, transfert de compétences vers une situation nouvelle.
- Rappel de cours : absent ou très succinct (2 lignes max).
- Exercice d'application : données brutes à extraire soi-même.
- Mise en situation complexe avec plusieurs informations à croiser.
- Problème ouvert ambitieux — mais toujours réaliste pour un élève de Bac Pro.

Dans tous les cas : JAMAIS de calcul hors programme Bac Pro, JAMAIS de piège inutile.
La correction détaillée doit être adaptée au même niveau.

Structure de sortie (Markdown) :
1. **Rappel de cours** (adapté au niveau)
2. **Exercice d'application** (adapté au niveau)
3. **Exercice de mise en situation** (contexte professionnel de la filière)
4. **Problème ouvert** (adapté au niveau)
5. **Corrections détaillées** (avec le niveau de détail approprié)\
"""


def build_prompt_exercices(niveau, categorie, matiere, chapitre, consignes, filiere="", difficulte="🟡 Moyen"):
    ctx        = build_contexte_filiere(filiere)
    diff_label = difficulte.split(" ", 1)[-1].upper()
    user = (
        f"Génère un contenu pédagogique de niveau **{diff_label}** pour :\n"
        f"- Niveau scolaire : {niveau} ({categorie})\n"
        f"- Matière : {matiere}\n"
        f"- Chapitre : {chapitre}\n"
        f"{ctx}"
        f"- Instructions : {consignes or 'Aucune'}\n\n"
        f"Applique scrupuleusement les consignes du niveau {diff_label} définies dans tes instructions."
    )
    return SYSTEM_EXERCICES, user


def build_prompt_ccf_entrainement(niveau, categorie, matiere, chapitre, consignes,
                                   filiere="", avec_corrige=True, chapitre_b=""):
    ctx          = build_contexte_filiere(filiere)
    bloc_corrige = "\n### CORRIGÉ DÉTAILLÉ *(document professeur)*\nCorrection complète de chaque question.\n" if avec_corrige else ""
    chap_a_label = f"Partie A : {chapitre}"
    chap_b_label = f"Partie B : {chapitre_b}" if chapitre_b else "Partie B : thème mathématique complémentaire au choix"
    return f"""Tu es un professeur de mathématiques expert en Bac Pro et en évaluation CCF conforme au BO.

Génère un SUJET D'ENTRAÎNEMENT AU CCF pour :
- Niveau : {niveau} ({categorie})
- Matière : {matiere}
- {chap_a_label}
- {chap_b_label}
{ctx}
- Instructions : {consignes or 'Aucune'}

## RÈGLES STRICTES DU FORMAT CCF OFFICIEL :

1. La MISE EN SITUATION présente le contexte professionnel avec données chiffrées réalistes.
2. La PROBLÉMATIQUE est UNE SEULE question (avec ?) encadrée en gras. C'est la seule question avec un "?".
3. Les questions sont numérotées avec le nom de la compétence en majuscules : S'APPROPRIER / RÉALISER / ANALYSER-RAISONNER / VALIDER / COMMUNIQUER.
4. Chaque question se termine par des lignes de réponse (______).
5. La dernière question COMMUNIQUER répond à la problématique initiale.
6. L'évaluation est sur 10 points, répartis en niveaux 0/1/2 par compétence.
7. Les deux parties s'appuient sur LA MÊME mise en situation.

### MISE EN SITUATION PROFESSIONNELLE
[Description du contexte avec données chiffrées et document support]

### PROBLÉMATIQUE
**[UNE SEULE question centrale se terminant par ?]**

### PARTIE A — [Titre lié à : {chapitre}]
1. **S'APPROPRIER** - [question] ______
2. **RÉALISER** - [question] ______
3. **ANALYSER / RAISONNER** - [question] ______
4. **RÉALISER** 🛎️ APPELER L'EXAMINATEUR - [question] ______
5. **VALIDER** - [question] ______
6. **COMMUNIQUER** - Répondre à la problématique. ______

### PARTIE B — [{chapitre_b if chapitre_b else "second thème mathématique"}]
[Mêmes règles]
{bloc_corrige}
Réponds entièrement en Markdown."""


def build_prompt_ccf_officiel(niveau, categorie, matiere, chapitre, consignes,
                               filiere="", duree="45 min", num_sit="1",
                               avec_corrige=True, chapitre_b=""):
    ctx         = build_contexte_filiere(filiere)
    nom_filiere = CONTEXTES_FILIERES[filiere]["nom_complet"] if filiere in CONTEXTES_FILIERES else filiere
    bloc_corrige = "\n### CORRIGÉ DÉTAILLÉ *(document professeur — NE PAS DISTRIBUER)*\n[Correction complète]\n" if avec_corrige else ""
    chap_b_label = f"Partie B : {chapitre_b}" if chapitre_b else "Partie B : thème complémentaire en cohérence avec la situation"
    return f"""Tu es un professeur de mathématiques expert en Bac Pro et en évaluation CCF conforme au BO.

Génère un SUJET DE CCF OFFICIEL pour :
- Niveau : {niveau} ({categorie}) — Filière : {nom_filiere}
- Matière : {matiere}
- Partie A : {chapitre} — {chap_b_label}
- Situation n° {num_sit} — Durée : {duree}
{ctx}
- Instructions : {consignes or 'Aucune'}

## RÈGLES STRICTES :
1. MISE EN SITUATION avec données chiffrées et document support nommé.
2. PROBLÉMATIQUE : UNE SEULE question avec ? — c'est la SEULE avec un point d'interrogation.
3. Questions numérotées avec compétences BO en gras : **S'APPROPRIER** / **RÉALISER** / **ANALYSER / RAISONNER** / **VALIDER** / **COMMUNIQUER**.
4. Toutes les 5 compétences évaluées au moins une fois.
5. Lignes de réponse (______) après chaque question.
6. Dernière question **COMMUNIQUER** répond à la problématique.
7. 🛎️ APPELER L'EXAMINATEUR quand calculatrice requise.
8. Noté sur /10 avec niveaux 0/1/2 par compétence.
9. Les deux parties sur LA MÊME mise en situation.
{bloc_corrige}
Réponds entièrement en Markdown avec mise en page soignée."""


def build_prompt_boss(niveau, filiere, chapitre, niveau_valide):
    ctx        = build_contexte_filiere(filiere)
    diff_label = niveau_valide.split(" ", 1)[-1].upper()
    mascotte   = MASCOTTES[niveau_valide]["animal"]
    return f"""Tu es un professeur expert en Bac Pro qui crée un DÉFI BOSS pour un élève.

L'élève vient de valider le niveau {diff_label} sur le chapitre "{chapitre}".

Crée un sujet d'exercices BOSS pour :
- Niveau scolaire : {niveau} (Bac Pro) — Matière : Mathématiques
- Chapitre : {chapitre} — Difficulté : {diff_label} AVANCÉ
{ctx}

RÈGLES DU BOSS :
- Mise en situation professionnelle réaliste et complète, sans données guidées.
- 4 à 5 questions progressives sans aide, sans formules rappelées.
- La dernière question demande un raisonnement complet et une conclusion rédigée.
- PAS DE CORRIGÉ — l'élève doit s'auto-corriger ou demander au professeur.
- Ton encourageant mais exigeant.

### ⚔️ DÉFI BOSS — {mascotte} {MASCOTTES[niveau_valide]["nom"]}
*Prouve que tu maîtrises vraiment ce chapitre !*

Réponds entièrement en Markdown."""


# ============================================================
# 6. APPEL API GEMINI
# ============================================================

def call_gemini(api_key: str, prompt, image=None) -> str:
    genai.configure(api_key=api_key.strip())
    if isinstance(prompt, tuple):
        system_instruction, user_prompt = prompt
        model   = genai.GenerativeModel(model_name="gemini-2.5-flash",
                                        system_instruction=system_instruction)
        content = user_prompt
    else:
        model   = genai.GenerativeModel(model_name="gemini-2.5-flash")
        content = prompt
    if image:
        img      = PIL.Image.open(image)
        response = model.generate_content([content, img])
    else:
        response = model.generate_content(content)
    if response and response.text:
        return response.text
    return "L'IA a répondu mais le texte est vide. Réessayez."


# ============================================================
# 7. INTÉGRATION GRIST
# ============================================================

def _grist_config() -> tuple:
    """Retourne (api_key, doc_id, base_url) depuis st.secrets."""
    import streamlit as st
    return (
        st.secrets.get("GRIST_API_KEY", ""),
        st.secrets.get("GRIST_DOC_ID", ""),
        st.secrets.get("GRIST_URL", "https://grist.numerique.gouv.fr"),
    )


def envoyer_grist(code_eleve, type_activite, meta, auto_evaluation="",
                  genre="", avatar=""):
    """Envoie une ligne dans Suivi_eleves. Affiche un warning sidebar si erreur."""
    import streamlit as st
    try:
        api_key, doc_id, base_url = _grist_config()
        if not api_key or not doc_id:
            return
        now = datetime.now(ZoneInfo("Europe/Paris"))
        url = f"{base_url}/api/docs/{doc_id}/tables/Suivi_eleves/records"
        headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
        payload = {"records": [{"fields": {
            "code_eleve":        str(code_eleve),
            "date":              now.strftime("%Y-%m-%d"),
            "heure":             now.strftime("%H:%M"),
            "type_activite":     type_activite,
            "classe":            str(meta.get("niveau", "")),
            "filiere":           str(meta.get("filiere", "")),
            "matiere":           str(meta.get("matiere", "")),
            "chapitre":          str(meta.get("chapitre", "")),
            "niveau_difficulte": str(meta.get("difficulte", "")),
            "auto_evaluation":   str(auto_evaluation),
            "score_auto":        str(meta.get("score_auto", "")),
            "source":            str(meta.get("source", "Gemini")),
            "genre":             str(genre),
            "avatar":            str(avatar),
        }}]}
        resp = requests.post(url, headers=headers, json=payload, timeout=5)
        if resp.status_code not in (200, 201):
            st.sidebar.warning(f"⚠️ Grist ({resp.status_code}) : {resp.text[:150]}")
    except Exception as e:
        st.sidebar.warning(f"⚠️ Erreur Grist : {e}")


def envoyer_proposition_grist(code_eleve, meta, contenu, auto_evaluation="", json_path=""):
    """Envoie un sujet généré dans Banque_propositions."""
    import streamlit as st
    try:
        api_key, doc_id, base_url = _grist_config()
        if not api_key or not doc_id:
            return False
        now = datetime.now(ZoneInfo("Europe/Paris"))
        url = f"{base_url}/api/docs/{doc_id}/tables/Banque_propositions/records"
        headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
        payload = {"records": [{"fields": {
            "code_eleve":        str(code_eleve),
            "date":              now.strftime("%Y-%m-%d"),
            "heure":             now.strftime("%H:%M"),
            "niveau":            str(meta.get("niveau", "")),
            "filiere":           str(meta.get("filiere", "")),
            "matiere":           str(meta.get("matiere", "")),
            "chapitre":          str(meta.get("chapitre", "")),
            "niveau_difficulte": str(meta.get("difficulte", "")),
            "contenu":           str(contenu)[:5000],
            "json_path":         str(json_path),
            "auto_evaluation":   str(auto_evaluation),
            "statut":            "en attente",
        }}]}
        resp = requests.post(url, headers=headers, json=payload, timeout=8)
        return resp.status_code in (200, 201)
    except Exception as e:
        st.sidebar.warning(f"⚠️ Erreur proposition Grist : {e}")
        return False


def lire_progression_grist(code_eleve: str) -> list:
    """Lit toutes les lignes Suivi_eleves pour un élève."""
    try:
        import urllib.parse
        api_key, doc_id, base_url = _grist_config()
        if not api_key or not doc_id or not code_eleve:
            return []
        filtre = urllib.parse.quote(json.dumps({"code_eleve": [code_eleve]}))
        url     = f"{base_url}/api/docs/{doc_id}/tables/Suivi_eleves/records?filter={filtre}"
        headers = {"Authorization": f"Bearer {api_key}"}
        resp    = requests.get(url, headers=headers, timeout=8)
        if resp.status_code == 200:
            return [r["fields"] for r in resp.json().get("records", [])]
    except Exception:
        pass
    return []


def lire_classement_grist() -> list:
    """Lit les 10 meilleurs élèves par XP."""
    try:
        api_key, doc_id, base_url = _grist_config()
        if not api_key or not doc_id:
            return []
        url     = f"{base_url}/api/docs/{doc_id}/tables/Suivi_eleves/records"
        headers = {"Authorization": f"Bearer {api_key}"}
        resp    = requests.get(url, headers=headers, timeout=8)
        if resp.status_code != 200:
            return []
        records   = [r["fields"] for r in resp.json().get("records", [])]
        par_eleve = defaultdict(list)
        for r in records:
            code = r.get("code_eleve", "")
            if code and code != "anonyme":
                par_eleve[code].append(r)
        classement = []
        for code, recs in par_eleve.items():
            classement.append({
                "code":   code,
                "xp":     calculer_xp(recs),
                "streak": calculer_streak(recs).get("streak", 0),
            })
        return sorted(classement, key=lambda x: x["xp"], reverse=True)[:10]
    except Exception:
        return []


# ============================================================
# 8. FONCTIONS MÉTIER — Gamification
# ============================================================

def calculer_xp(records: list) -> int:
    total = 0
    for r in records:
        if r.get("auto_evaluation", "") not in ("😊 Bien", "🌟 Très bien"):
            continue
        type_a = r.get("type_activite", "")
        diff   = r.get("niveau_difficulte", "")
        if "Boss" in type_a:
            total += XP_PAR_ACTION["Boss"]
        elif "CCF" in type_a:
            total += XP_PAR_ACTION["CCF"]
        else:
            total += XP_PAR_ACTION.get(f"Exercice-{diff}", 10)
    return total


def calculer_streak(records: list) -> dict:
    from datetime import date, timedelta
    dates_ok = set()
    for r in records:
        if r.get("auto_evaluation", "") in ("😊 Bien", "🌟 Très bien"):
            d = r.get("date", "")
            if d:
                try:
                    dates_ok.add(date.fromisoformat(str(d)[:10]))
                except ValueError:
                    pass
    if not dates_ok:
        return {"streak": 0, "record": 0, "nouveau_record": False, "derniere_date": None}

    today, yesterday    = date.today(), date.today() - timedelta(days=1)
    dates_ok_sorted     = sorted(dates_ok, reverse=True)
    derniere            = dates_ok_sorted[0]

    if derniere < yesterday:
        streak_actuel = 0
    else:
        streak_actuel, curseur = 1, derniere
        for d in dates_ok_sorted[1:]:
            if curseur - d == timedelta(days=1):
                streak_actuel += 1
                curseur = d
            else:
                break

    record, seq_cur = 0, 1
    for i in range(1, len(dates_ok_sorted)):
        if dates_ok_sorted[i-1] - dates_ok_sorted[i] == timedelta(days=1):
            seq_cur += 1
        else:
            record  = max(record, seq_cur)
            seq_cur = 1
    record = max(record, seq_cur)

    return {
        "streak":         streak_actuel,
        "record":         record,
        "nouveau_record": streak_actuel > 0 and streak_actuel >= record,
        "derniere_date":  derniere,
    }


def calculer_progression(records: list, chapitre: str) -> dict:
    bonnes, boss_ok = defaultdict(int), set()
    for r in records:
        if r.get("chapitre", "") != chapitre:
            continue
        diff, eval_v, type_a = (r.get("niveau_difficulte", ""),
                                r.get("auto_evaluation", ""),
                                r.get("type_activite", ""))
        if eval_v in ("😊 Bien", "🌟 Très bien"):
            if "Boss" in type_a:
                boss_ok.add(diff)
            else:
                bonnes[diff] += 1
    return {
        diff: {
            "nb_bonnes":   bonnes.get(diff, 0),
            "valide":      bonnes.get(diff, 0) >= SEUIL_VALIDATION,
            "boss_vaincu": diff in boss_ok,
        }
        for diff in ORDRE_NIVEAUX_DIFF
    }


def niveau_suivant(diff: str):
    idx = ORDRE_NIVEAUX_DIFF.index(diff)
    return ORDRE_NIVEAUX_DIFF[idx + 1] if idx < len(ORDRE_NIVEAUX_DIFF) - 1 else None


def calculer_objectif(records: list, chapitre_actif: str, difficulte_active: str) -> dict:
    from datetime import date
    s        = calculer_streak(records)
    streak   = s.get("streak", 0)
    derniere = s.get("derniere_date")

    if derniere and derniere < date.today() and streak > 0:
        return {"emoji": "🔥", "couleur": "#ef4444", "priorite": 3,
                "message": f"Ton streak de {streak} jour{'s' if streak > 1 else ''} est en danger !"}

    if chapitre_actif and records:
        prog = calculer_progression(records, chapitre_actif)
        for diff in ORDRE_NIVEAUX_DIFF:
            p = prog.get(diff, {})
            if p.get("valide") and not p.get("boss_vaincu"):
                return {"emoji": "⚔️", "couleur": "#7c3aed", "priorite": 2,
                        "message": f"Le boss {MASCOTTES[diff]['animal']} t'attend sur '{chapitre_actif[:35]}' !"}

    if chapitre_actif and records:
        prog = calculer_progression(records, chapitre_actif)
        for diff in ORDRE_NIVEAUX_DIFF:
            p, nb = prog.get(diff, {}), prog.get(diff, {}).get("nb_bonnes", 0)
            if 0 < nb < SEUIL_VALIDATION and not p.get("valide"):
                reste = SEUIL_VALIDATION - nb
                return {"emoji": "🎯", "couleur": "#f59e0b", "priorite": 1,
                        "message": f"Plus que {reste} bonne{'s' if reste > 1 else ''} éval en {diff.split()[-1]} !"}

    if streak == 0:
        return {"emoji": "🚀", "couleur": "#6366f1", "priorite": 0,
                "message": "Commence ta session du jour — même 1 exercice suffit !"}
    return {"emoji": "💪", "couleur": "#22c55e", "priorite": 0,
            "message": f"Super ! {streak} jour{'s' if streak > 1 else ''} d'affilée. Continue !"}


def calculer_badges(records: list, streak: int) -> set:
    if not records:
        return set()
    bonnes      = [r for r in records if r.get("auto_evaluation","") in ("😊 Bien","🌟 Très bien")]
    tres_bien   = [r for r in records if r.get("auto_evaluation","") == "🌟 Très bien"]
    boss_ok     = [r for r in bonnes   if "Boss"        in r.get("type_activite","")]
    ccf_ok      = [r for r in bonnes   if "CCF"         in r.get("type_activite","")]
    interactifs = [r for r in records  if "interactif"  in r.get("type_activite","").lower()]
    chapitres   = set(r.get("chapitre","") for r in records if r.get("chapitre"))

    debloque = set()
    if records:             debloque.add("premier_pas")
    if bonnes:              debloque.add("premiere_victoire")
    if interactifs:         debloque.add("gamer")
    if streak >= 3:         debloque.add("en_feu")
    if streak >= 7:         debloque.add("flamme")
    if any("(100%)" in r.get("type_activite","") for r in records):
                            debloque.add("sans_faute")
    if len(tres_bien) >= 5: debloque.add("perfectionniste")
    if len(bonnes) >= 10:   debloque.add("assidu")
    if boss_ok:             debloque.add("premier_boss")
    if any("🔴 Expert" in r.get("niveau_difficulte","") or
           "Expert"    in r.get("type_activite","")
           for r in boss_ok):
                            debloque.add("chasseur_dragons")
    if len(chapitres) >= 3: debloque.add("explorateur")
    if len(chapitres) >= 5: debloque.add("encyclopediste")
    if ccf_ok:              debloque.add("candidat_ccf")

    chapitres_boss_ok = set(r.get("chapitre","") for r in boss_ok if r.get("chapitre"))
    for chap in chapitres_boss_ok:
        prog = calculer_progression(records, chap)
        if all(prog[d]["valide"] for d in ORDRE_NIVEAUX_DIFF):
            debloque.add("champion")
            break

    return debloque


# ============================================================
# 9. BANQUE D'EXERCICES
# ============================================================

def _normalise(text: str) -> str:
    import unicodedata
    text = unicodedata.normalize("NFD", text)
    text = "".join(c for c in text if unicodedata.category(c) != "Mn")
    return re.sub(r"[^a-z0-9]", "", text.lower())


def _slug(text: str) -> str:
    text = text.replace(" ", "_").replace("—", "").replace("/", "_")
    return re.sub(r"[^a-zA-Z0-9_\-àâéèêëîïôùûüç]", "", text)[:60]


def _trouver_fichier_banque(niveau, filiere, chapitre, difficulte, suffixe=""):
    if not os.path.exists(BANQUE_DIR):
        return None
    niv_n  = _normalise(niveau)
    fil_n  = _normalise(filiere)
    diff_n = _normalise(difficulte.split(" ", 1)[-1])
    chap_n = _normalise(chapitre)[:45]
    ext    = f"{suffixe}.json"

    meilleur, meilleur_score = None, 0
    for fname in os.listdir(BANQUE_DIR):
        if not fname.endswith(ext):
            continue
        if suffixe == "" and fname.endswith("_interactif.json"):
            continue
        fname_n = _normalise(fname)
        score   = sum(1 for k in [niv_n, fil_n, chap_n, diff_n] if k and k in fname_n)
        if score > meilleur_score:
            meilleur_score, meilleur = score, fname

    if meilleur_score >= 3 and meilleur:
        return os.path.join(BANQUE_DIR, meilleur)
    return None


def charger_banque(niveau, filiere, chapitre, difficulte) -> list:
    path = _trouver_fichier_banque(niveau, filiere, chapitre, difficulte, "")
    if not path:
        return []
    try:
        with open(path, encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return []


def charger_banque_interactif(niveau, filiere, chapitre, difficulte) -> list:
    path = _trouver_fichier_banque(niveau, filiere, chapitre, difficulte, "_interactif")
    if not path:
        return []
    try:
        with open(path, encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return []
        
# ============================================================
# utils.py — PARTIE 3 : Export Word
# ============================================================

# ── À coller à la suite de la Partie 2 ──────────────────────


# ============================================================
# 10. EXPORT WORD — VERSION STANDARD
# ============================================================

def clean_math(text: str) -> str:
    """Convertit la notation LaTeX/Markdown en texte lisible dans Word."""
    import re as _re
    text = _re.sub(r'```[a-zA-Z]*', '', text)
    text = _re.sub(r'\\\[(.+?)\\\]', lambda m: m.group(1).strip(), text, flags=_re.DOTALL)
    text = _re.sub(r'\\\((.+?)\\\)', lambda m: m.group(1).strip(), text, flags=_re.DOTALL)
    text = _re.sub(r'\$\$(.+?)\$\$',  lambda m: m.group(1).strip(), text, flags=_re.DOTALL)
    text = _re.sub(r'\$([^$\n]+?)\$', lambda m: m.group(1).strip(), text)
    text = _re.sub(r'\\frac\{([^}]+)\}\{([^}]+)\}', r'\1/\2', text)
    text = _re.sub(r'\^\{([^}]+)\}',  r'^\1', text)
    text = _re.sub(r'_\{([^}]+)\}',   r'_\1', text)
    for latex, uni in {
        r'\times': '×', r'\cdot': '·', r'\leq': '≤', r'\geq': '≥',
        r'\neq': '≠', r'\approx': '≈', r'\infty': '∞', r'\pi': 'π',
        r'\alpha': 'α', r'\beta': 'β', r'\sigma': 'σ', r'\mu': 'μ',
        r'\rightarrow': '→', r'\leftarrow': '←', r'\in': '∈',
        r'\subset': '⊂', r'\cup': '∪', r'\cap': '∩',
    }.items():
        text = text.replace(latex, uni)
    text = _re.sub(r'\\([a-zA-Z]+)', r'\1', text)
    return text


def parse_md_tables(md_text: str) -> list:
    """Transforme un texte Markdown en liste de blocs (text | table)."""
    import re as _re
    lines, blocks, i = md_text.split('\n'), [], 0
    while i < len(lines):
        line = lines[i]
        if line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                raw = lines[i].strip()
                if _re.match(r'^\|[\s\-:|]+\|', raw):
                    i += 1
                    continue
                table_lines.append([c.strip() for c in raw.strip('|').split('|')])
                i += 1
            if table_lines:
                blocks.append(('table', table_lines))
        else:
            blocks.append(('text', line))
            i += 1
    return blocks


def add_md_table_to_doc(doc, rows: list, font_size: int = 10):
    """Crée un vrai tableau Word depuis une liste de lignes [[cell, ...]]."""
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    table    = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(max_cols):
            cell = table.cell(i, j)
            text = clean_math(row[j]) if j < len(row) else ""
            p    = cell.paragraphs[0]
            p.clear()
            run = p.add_run(text)
            run.font.size = Pt(font_size)
            run.font.name = "Arial"
            if i == 0:
                run.bold = True
    doc.add_paragraph()


def render_inline(paragraph, text: str):
    """Écrit une ligne avec gras/italique en nettoyant le LaTeX."""
    text  = clean_math(text)
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            paragraph.add_run(part[1:-1]).italic = True
        else:
            paragraph.add_run(part)


def markdown_to_docx(md_text: str, titre: str = "Document") -> bytes:
    doc = DocxDocument()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)
    h = doc.add_heading(titre, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for kind, content in parse_md_tables(md_text):
        if kind == "table":
            add_md_table_to_doc(doc, content)
            continue
        line = content.rstrip()
        if   line.startswith("### "): doc.add_heading(clean_math(line[4:]), level=3)
        elif line.startswith("## "):  doc.add_heading(clean_math(line[3:]), level=2)
        elif line.startswith("# "):   doc.add_heading(clean_math(line[2:]), level=1)
        elif line.startswith("---"):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
        elif re.match(r"^\d+\.", line):
            p = doc.add_paragraph(style="List Number")
            render_inline(p, re.sub(r"^\d+\.\s*", "", line))
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            render_inline(p, line[2:])
        elif line == "":
            doc.add_paragraph("")
        else:
            p = doc.add_paragraph()
            render_inline(p, line)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ============================================================
# 11. EXPORT WORD — VERSION OFFICIELLE CCF
# ============================================================

def set_cell_border(cell):
    tc, tcPr   = cell._tc, cell._tc.get_or_add_tcPr()
    tcBorders  = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "AAAAAA")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def set_run(paragraph, text: str, bold: bool = False, size: int = 10):
    run = paragraph.add_run(text)
    run.bold       = bold
    run.font.size  = Pt(size)
    return run


def fill_cell(cell, text: str, bold: bool = False, size: int = 10):
    p = cell.paragraphs[0]
    p.clear()
    set_run(p, text, bold=bold, size=size)
    set_cell_border(cell)


def parse_questions_competences(content_md: str) -> dict:
    COMP_KEYWORDS = {
        "S'APPROPRIER": "S'approprier", "APPROPRIER": "S'approprier",
        "ANALYSER / RAISONNER": "Analyser / Raisonner",
        "ANALYSER": "Analyser / Raisonner", "RAISONNER": "Analyser / Raisonner",
        "RÉALISER": "Réaliser", "REALISER": "Réaliser",
        "VALIDER": "Valider", "COMMUNIQUER": "Communiquer",
    }
    result = {"S'approprier": [], "Analyser / Raisonner": [], "Réaliser": [], "Valider": [], "Communiquer": []}
    current_part = ""
    for line in content_md.split("\n"):
        ls = line.strip()
        m  = re.match(r"#+\s+PARTIE\s+([A-Z])", ls, re.IGNORECASE)
        if m:
            current_part = m.group(1).upper()
            continue
        if not current_part:
            continue
        q = re.match(r"^(\d+)\.\s+\*\*([^*]+)\*\*", ls)
        if q:
            q_ref    = f"{current_part}.{q.group(1)}"
            comp_raw = q.group(2).strip().upper()
            for kw, name in COMP_KEYWORDS.items():
                if kw in comp_raw:
                    if q_ref not in result[name]:
                        result[name].append(q_ref)
                    break
    return result


def generate_ccf_officiel_docx(content_md: str, metadata: dict,
                                nom_etablissement: str = "Mon Établissement"):
    try:
        doc = DocxDocument()
        for section in doc.sections:
            section.top_margin    = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin   = Inches(1.0)
            section.right_margin  = Inches(1.0)
        doc.styles["Normal"].font.name = "Arial"
        doc.styles["Normal"].font.size = Pt(11)

        # ── EN-TÊTE ──
        table_h    = doc.add_table(rows=3, cols=3)
        table_h.style = "Table Grid"
        col_widths = [2620, 4120, 2620]
        for row in table_h.rows:
            for j, cell in enumerate(row.cells):
                tc   = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW  = OxmlElement("w:tcW")
                tcW.set(qn("w:w"), str(col_widths[j]))
                tcW.set(qn("w:type"), "dxa")
                tcPr.append(tcW)

        cell_l    = table_h.cell(0, 0)
        logo_rep  = os.path.join(APP_DIR, "logo_republique.png")
        if os.path.exists(logo_rep):
            cell_l.paragraphs[0].add_run().add_picture(logo_rep, width=Inches(1.2))
        else:
            r = cell_l.paragraphs[0].add_run("ACADÉMIE DE CRÉTEIL")
            r.bold = True
            r.font.size = Pt(10)

        def _rb(para, text, size=10):
            r = para.add_run(text)
            r.bold = True
            r.font.size = Pt(size)
            return r

        cell_c = table_h.cell(0, 1)
        p_c    = cell_c.paragraphs[0]
        p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _rb(p_c, "CONTRÔLE EN COURS DE FORMATION\n", 12)
        _rb(p_c, "Baccalauréat professionnel\n", 11)
        _rb(p_c, f"{metadata.get('matiere', '')}\n", 11)
        _rb(p_c, f"Situation d'évaluation n°{metadata.get('num_situation', '...')}\n", 10)
        _rb(p_c, f"Intitulé du diplôme : {'.' * 31}\n", 10)
        _rb(p_c, f"Durée : {metadata.get('duree', '45 min')}", 10)

        cell_r   = table_h.cell(0, 2)
        p_r      = cell_r.paragraphs[0]
        p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        logo_mat = os.path.join(APP_DIR, "logo_matiere.png")
        if os.path.exists(logo_mat):
            p_r.add_run().add_picture(logo_mat, width=Inches(0.8))

        cell_nom = table_h.cell(1, 1)
        p_nom    = cell_nom.paragraphs[0]
        p_nom.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _rb(p_nom, "Nom, Prénom du candidat : _______________________________________________", 10)

        table_h.cell(2, 0).paragraphs[0].add_run("Date : .......................").font.size = Pt(10)
        _rb(table_h.cell(2, 1).paragraphs[0], f"Nom de l'établissement : {nom_etablissement}", 10)

        doc.add_paragraph()

        # ── NOTICE CALCULATRICE ──
        logo_appel_path = os.path.join(APP_DIR, "logo_appel.png")
        p_notice = doc.add_paragraph()
        if os.path.exists(logo_appel_path):
            p_notice.add_run().add_picture(logo_appel_path, width=Inches(0.53))
            p_notice.add_run("  ")
        p_notice.add_run("Dans la suite du document, ce symbole signifie « Appeler l'examinateur ».").font.size = Pt(10)
        for ligne in [
            "L'usage de la calculatrice avec mode examen actif est autorisé.",
            "L'usage de la calculatrice sans mémoire, « type collège » est autorisé.",
            "L'échange de calculatrices entre les candidats pendant l'épreuve est interdit.",
        ]:
            doc.add_paragraph().add_run(ligne).font.size = Pt(10)
        doc.add_paragraph()

        # ── CONTENU DU SUJET ──
        skip_kw = [
            "Établissement :", "Baccalauréat Professionnel —", "Épreuve E3",
            "Calculatrice autorisée", "Ministère de l'Éducation",
            "sujet de CCF officiel", "Voici le sujet", "stricte conformité",
            "accompagné de", "fiche d'évaluation et de son corrigé",
            "Filière :", "Niveau :", "Situation d'évaluation n°", "Durée :", "Nom & Prénom",
        ]
        filtered_lines, contenu_commence = [], False
        for line in content_md.split("\n"):
            ls = line.rstrip()
            if not contenu_commence:
                if any(ls.startswith(m) for m in [
                    "### MISE EN SITUATION", "### PARTIE", "### PROBLÉMATIQUE",
                    "## MISE EN SITUATION",  "## PARTIE",  "## PROBLÉMATIQUE",
                    "#### MISE", "#### PARTIE",
                ]):
                    contenu_commence = True
                else:
                    continue
            if any(kw in ls for kw in skip_kw):
                continue
            filtered_lines.append(ls)

        for kind, content in parse_md_tables("\n".join(filtered_lines)):
            if kind == "table":
                add_md_table_to_doc(doc, content, font_size=10)
                continue
            line = content
            if   line.startswith("#### "): doc.add_heading(clean_math(line[5:].strip()), level=3)
            elif line.startswith("### "):  doc.add_heading(clean_math(line[4:].strip()), level=2)
            elif line.startswith("## "):   doc.add_heading(clean_math(line[3:].strip()), level=2)
            elif "APPELER L'EXAMINATEUR" in line.upper() or "🛎️" in line:
                p = doc.add_paragraph()
                if os.path.exists(logo_appel_path):
                    p.add_run().add_picture(logo_appel_path, width=Inches(0.25))
                    p.add_run("  ")
                run = p.add_run("APPELER L'EXAMINATEUR")
                run.bold = True
                run.font.size = Pt(11)
                pPr  = p._p.get_or_add_pPr()
                pBdr = OxmlElement("w:pBdr")
                for side in ("top", "left", "bottom", "right"):
                    b = OxmlElement(f"w:{side}")
                    b.set(qn("w:val"), "single")
                    b.set(qn("w:sz"), "6")
                    b.set(qn("w:space"), "4")
                    b.set(qn("w:color"), "4A6CF7")
                    pBdr.append(b)
                pPr.append(pBdr)
            elif line.startswith("---"): continue
            elif line == "":             doc.add_paragraph("")
            else:
                p = doc.add_paragraph()
                render_inline(p, line)

        # ── FICHE D'ÉVALUATION (page 2) ──
        doc.add_page_break()
        titre_fiche = doc.add_heading("FICHE INDIVIDUELLE D'ÉVALUATION", level=1)
        titre_fiche.alignment = WD_ALIGN_PARAGRAPH.CENTER

        table_ef = doc.add_table(rows=2, cols=2)
        fill_cell(table_ef.cell(0, 0),
                  f"Session : {metadata.get('annee_scolaire','2025/2026')} — "
                  f"Établissement : {nom_etablissement} — Académie : Créteil",
                  bold=True, size=9)
        fill_cell(table_ef.cell(0, 1),
                  f"Spécialité : {metadata.get('matiere','')} — "
                  f"Évaluateur : _______________ — Date : _______________", size=9)
        fill_cell(table_ef.cell(1, 0),
                  f"Situation n° {metadata.get('num_situation','1')} — {metadata.get('filiere','')}",
                  bold=True, size=9)
        fill_cell(table_ef.cell(1, 1), "", size=9)

        doc.add_paragraph()
        set_run(doc.add_paragraph(),
                "Nom et prénom du candidat : _______________________________________________",
                bold=True, size=10)
        doc.add_paragraph()
        set_run(doc.add_paragraph(), "1.  Liste des capacités et connaissances évaluées",
                bold=True, size=11)

        table_cap = doc.add_table(rows=2, cols=2)
        fill_cell(table_cap.cell(0, 0), "Capacités",     bold=True, size=10)
        fill_cell(table_cap.cell(0, 1), "(voir section Fiche d'évaluation du sujet généré)", size=10)
        fill_cell(table_cap.cell(1, 0), "Connaissances", bold=True, size=10)
        fill_cell(table_cap.cell(1, 1), "(voir section Fiche d'évaluation du sujet généré)", size=10)

        doc.add_paragraph()
        set_run(doc.add_paragraph(), "2.  Évaluation", bold=True, size=11)

        total_rows = 1 + sum(len(c["indicateurs"]) for c in COMPETENCES_CCF) + 1
        grid       = doc.add_table(rows=total_rows, cols=5)
        for j, h in enumerate(["Compétences", "", "Indicateurs / Capacités",
                                "Questions", "Appréciation\n(0 / 1 / 2)"]):
            fill_cell(grid.cell(0, j), h, bold=True, size=9)

        questions_par_comp = parse_questions_competences(content_md)
        row_idx = 1
        for comp in COMPETENCES_CCF:
            refs     = questions_par_comp.get(comp["nom"], [])
            refs_str = "  ".join(refs) if refs else ""
            for k, indicateur in enumerate(comp["indicateurs"]):
                fill_cell(grid.cell(row_idx, 0), comp["nom"] if k == 0 else "",
                          bold=(k == 0), size=9)
                fill_cell(grid.cell(row_idx, 1), "", size=9)
                fill_cell(grid.cell(row_idx, 2), indicateur, size=9)
                fill_cell(grid.cell(row_idx, 3),
                          refs_str if k == 0 else "",
                          bold=(k == 0 and bool(refs_str)), size=9)
                fill_cell(grid.cell(row_idx, 4), "0    1    2", size=9)
                row_idx += 1

        fill_cell(grid.cell(row_idx, 0), "", size=9)
        fill_cell(grid.cell(row_idx, 1), "", size=9)
        fill_cell(grid.cell(row_idx, 2), "", size=9)
        fill_cell(grid.cell(row_idx, 3), "Note :", bold=True, size=10)
        fill_cell(grid.cell(row_idx, 4), "          / 10", bold=True, size=11)

        doc.add_paragraph()
        set_run(doc.add_paragraph(),
                "Observations : ___________________________________________________________________",
                size=10)

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()

    except Exception as e:
        import streamlit as st
        st.error(f"Erreur lors de la génération du Word officiel : {e}")
        return None