import google.generativeai as genai
import PIL.Image
import streamlit as st
from io import BytesIO
from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os
import time

# Chemin absolu du dossier contenant app.py — utilisé pour trouver les images
APP_DIR = os.path.dirname(os.path.abspath(__file__))

# ============================================================
# 1. DONNÉES OFFICIELLES — BO MATHS BAC PRO
# ============================================================

NIVEAUX_CATEGORIES = {
    "Collège": ["6ème", "5ème", "4ème", "3ème"],
    "Lycée Général": ["2nde", "1ère", "Terminale"],
    "Bac Pro": ["2nde Pro", "1ère Pro", "Term Pro"],
    "CAP": ["1ère année CAP", "2ème année CAP"],
    "BTS": ["BTS 1", "BTS 2"]
}

MATIERES = [
    "Mathématiques", "Français", "Histoire-Géographie", "SVT",
    "Physique-Chimie", "Anglais", "Espagnol", "SES",
    "Philosophie", "Sciences de l'ingénieur", "EMC"
]

CHAPITRES_MATHS_BAC_PRO = {
    # ── 2nde Pro ── (programme déjà conforme au BO)
    "2nde Pro": [
        "Calcul numérique et algébrique",
        "Puissances et notations scientifiques",
        "Proportionnalité et pourcentages",
        "Équations du 1er degré à une inconnue",
        "Inéquations du 1er degré",
        "Géométrie plane — figures usuelles",
        "Trigonométrie dans le triangle rectangle",
        "Vecteurs — notions de base",
        "Notion de fonction — représentation graphique",
        "Fonctions linéaires et affines",
        "Statistiques descriptives — série à une variable",
        "Notion de probabilité — expériences aléatoires",
    ],

    # ── 1ère Pro ── conforme au BO Annexe 1 (Mathématiques — Classe de première professionnelle)
    "1ère Pro": [
        # Statistique et probabilités
        "Statistique à deux variables — ajustement affine et coefficient de détermination",
        "Probabilités — événements, tableaux croisés, probabilités conditionnelles",
        # Algèbre – Analyse
        "Suites numériques — suites arithmétiques",
        "Résolution graphique d'équations et d'inéquations f(x)=g(x)",
        "Fonctions polynômes de degré 2 — racines, signe, forme factorisée",
        "Fonction dérivée — variations, extremums, fonction inverse",
        "Calculs commerciaux et financiers — intérêts simples, coûts (filières sans physique-chimie)",
        # Géométrie
        "Géométrie dans l'espace — solides usuels et sections par un plan",
        "Vecteurs du plan — coordonnées, opérations, norme (groupements A et B)",
        "Trigonométrie — cercle trigonométrique, fonctions sinus et cosinus (groupements A et B)",
        # Modules transversaux
        "Algorithmique et programmation Python — listes, fonctions, boucles",
        "Automatismes — calcul, grandeurs, lecture graphique",
    ],

    # ── Term Pro ── conforme au BO Annexe 2 (Mathématiques — Classe terminale professionnelle)
    "Term Pro": [
        # Statistique et probabilités
        "Statistiques à deux variables — ajustements non affines, changements de variable",
        "Probabilités — arbres pondérés, formule des probabilités totales, indépendance",
        # Algèbre – Analyse
        "Suites géométriques — terme général, sens de variation, somme",
        "Fonctions polynômes de degré 3 — dérivée, variations, extremums",
        "Fonctions exponentielles de base q et logarithme décimal",
        "Calculs commerciaux et financiers — intérêts composés, amortissements (filières sans physique-chimie)",
        # Géométrie
        "Vecteurs dans l'espace — coordonnées, norme, colinéarité (groupement B)",
        "Trigonométrie — équations, vecteurs de Fresnel (groupement A)",
        # Modules transversaux
        "Algorithmique et programmation Python — approfondissement listes et fonctions",
        "Automatismes — probabilités, suites, dérivation, vecteurs",
        # Programme complémentaire (poursuite d'études)
        "Calcul intégral — primitives, intégrale, aire (programme complémentaire)",
        "Fonctions logarithme népérien et exponentielle de base e (programme complémentaire)",
        "Nombres complexes — forme algébrique et trigonométrique (programme complémentaire)",
        "Produit scalaire de deux vecteurs du plan (programme complémentaire)",
    ],
}

CHAPITRES_PAR_MATIERE_GENERAL = {
    "Mathématiques": ["Nombres et opérations", "Fractions", "Équations", "Fonctions", "Géométrie", "Statistiques", "Probabilités", "Algorithmique"],
    "Français": ["Grammaire", "Orthographe", "Conjugaison", "Analyse littéraire", "Argumentation", "Expression écrite", "Oral"],
    "Histoire-Géographie": ["Antiquité", "Moyen Âge", "Époque Moderne", "Époque Contemporaine", "Géographie de la France", "Géographie mondiale", "EMC"],
    "SVT": ["Cellule et génétique", "Évolution", "Corps humain et santé", "Écosystèmes", "Géologie"],
    "Physique-Chimie": ["Mécanique", "Électricité", "Optique", "Thermodynamique", "Chimie organique", "Chimie des solutions"],
    "Anglais": ["Compréhension écrite", "Expression écrite", "Compréhension orale", "Expression orale", "Grammaire", "Vocabulaire professionnel"],
    "Espagnol": ["Compréhension écrite", "Expression écrite", "Grammaire", "Civilisation", "Vocabulaire"],
    "SES": ["Économie", "Sociologie", "Science politique", "Mondialisation"],
    "Philosophie": ["La conscience", "Le langage", "La liberté", "La morale", "La politique", "La vérité", "L'art"],
    "Sciences de l'ingénieur": ["Mécanique", "Électronique", "Informatique industrielle", "Matériaux"],
    "EMC": ["Démocratie et citoyenneté", "Droits et libertés", "Laïcité", "Engagement"],
}

# ============================================================
# 2. FILIÈRES PRO
# ============================================================

LISTE_FILIERES = ["ASSP", "MCVB", "MCVA", "AGORA", "Autre (Préciser ci-dessous)"]

CONTEXTES_FILIERES = {
    "ASSP": {
        "nom_complet": "Accompagnement, Soins et Services à la Personne",
        "contextes_maths": [
            "Calcul de doses et de dilutions médicamenteuses",
            "Gestion de plannings d'intervenants à domicile",
            "Statistiques sur données de santé (IMC, fréquence cardiaque)",
            "Calcul de coûts de prise en charge",
            "Lecture et interprétation de graphiques médicaux",
        ],
        "mots_cles": "soins, patients, résidents, personnes âgées, handicap, domicile, EHPAD, pharmacie, hygiène"
    },
    "MCVB": {
        "nom_complet": "Métiers du Commerce et de la Vente — option B (Prospection Clientèle)",
        "contextes_maths": [
            "Calcul de marges, taux de remise et prix de vente",
            "Statistiques de vente et de prospection",
            "Gestion d'un budget commercial",
            "Pourcentages d'évolution du chiffre d'affaires",
            "Représentations graphiques des ventes",
        ],
        "mots_cles": "vente, client, prospection, chiffre d'affaires, commission, remise, catalogue"
    },
    "MCVA": {
        "nom_complet": "Métiers du Commerce et de la Vente — option A (Animation et Gestion de l'Espace Commercial)",
        "contextes_maths": [
            "Calcul de surfaces de vente et d'implantation linéaire",
            "Statistiques de fréquentation et taux de transformation",
            "Gestion des stocks — taux de rotation, coefficient multiplicateur",
            "Calcul de TVA, de prix TTC et HT",
            "Optimisation de l'espace commercial",
        ],
        "mots_cles": "magasin, rayon, linéaire, stock, inventaire, promotion, merchandising, caisse"
    },
    "AGORA": {
        "nom_complet": "Assistance à la Gestion des Organisations et de leurs Activités",
        "contextes_maths": [
            "Calcul de salaires, cotisations et charges sociales",
            "Gestion de budgets d'entreprise",
            "Statistiques sur données RH",
            "Facturation et suivi comptable de base",
            "Calcul d'intérêts pour emprunts professionnels",
        ],
        "mots_cles": "entreprise, comptabilité, salaire, facture, devis, ressources humaines, secrétariat"
    },
}

# Compétences officielles du BO avec leurs indicateurs détaillés
COMPETENCES_CCF = [
    {
        "nom": "S'approprier",
        "indicateurs": [
            "Rechercher, extraire et organiser l'information.",
            "Traduire des informations, des codages.",
        ]
    },
    {
        "nom": "Analyser / Raisonner",
        "indicateurs": [
            "Émettre des conjectures, formuler des hypothèses.",
            "Proposer, choisir une méthode de résolution ou un protocole expérimental.",
            "Élaborer un algorithme.",
        ]
    },
    {
        "nom": "Réaliser",
        "indicateurs": [
            "Mettre en œuvre une méthode de résolution, des algorithmes ou un protocole expérimental en respectant les règles de sécurité.",
            "Utiliser un modèle, représenter, calculer.",
            "Expérimenter, faire une simulation.",
        ]
    },
    {
        "nom": "Valider",
        "indicateurs": [
            "Exploiter et interpréter des résultats ou des observations de façon critique et argumentée.",
            "Contrôler la vraisemblance d'une conjecture, de la valeur d'une mesure.",
            "Valider un modèle ou une hypothèse.",
            "Mener un raisonnement logique et établir une conclusion.",
        ]
    },
    {
        "nom": "Communiquer",
        "indicateurs": [
            "Rendre compte d'un résultat, à l'oral ou à l'écrit en utilisant des outils et un langage approprié.",
            "Expliquer une démarche.",
        ]
    },
]

# ============================================================
# 3. PROMPTS
# ============================================================

def get_chapitres(matiere, niveau, categorie):
    if matiere == "Mathématiques" and categorie == "Bac Pro" and niveau in CHAPITRES_MATHS_BAC_PRO:
        return CHAPITRES_MATHS_BAC_PRO[niveau]
    return CHAPITRES_PAR_MATIERE_GENERAL.get(matiere, ["Chapitre général"])


def build_contexte_filiere(filiere):
    if filiere in CONTEXTES_FILIERES:
        ctx = CONTEXTES_FILIERES[filiere]
        lignes = "\n".join(f"- {c}" for c in ctx["contextes_maths"])
        return f"\n**Filière : {ctx['nom_complet']}**\nUnivers : {ctx['mots_cles']}\nContextes maths :\n{lignes}\n"
    return f"\nFilière : {filiere}\n" if filiere else ""


SYSTEM_EXERCICES = """\
Tu es un professeur expert en pédagogie différenciée pour lycée professionnel (Bac Pro).
Tes élèves ont un niveau en mathématiques fragile : tes énoncés sont toujours clairs, bienveillants et ancrés dans des contextes professionnels concrets.

En fonction du niveau de difficulté demandé, adapte PRÉCISÉMENT la structure suivante :

━━ DÉBUTANT ━━
- Questions très guidées, découpées en micro-étapes (une opération par question).
- Résultats intermédiaires fournis pour permettre de continuer même en cas d'erreur.
- Vocabulaire ultra-simplifié, aucun terme technique sans définition immédiate.
- Rappel de cours détaillé avec exemple résolu pas à pas.
- Exercice d'application : calcul direct, données déjà extraites.
- Mise en situation : contexte simple, une seule inconnue.
- Pas de problème ouvert — remplacer par une question bilan guidée.

━━ MOYEN ━━
- Questions semi-guidées avec quelques repères (formule rappelée, première étape donnée).
- Rappel de cours synthétique avec un exemple.
- Exercice d'application : 3 questions progressives.
- Mise en situation professionnelle simple avec tableau de données.
- Problème ouvert court (1 question de synthèse guidée).

━━ CONFIRMÉ ━━
- Questions autonomes, aucune aide dans l'énoncé.
- Rappel de cours en points clés uniquement (pas d'exemple résolu).
- Exercice d'application : questions progressives avec barème.
- Mise en situation professionnelle réaliste et complète.
- Problème ouvert avec raisonnement attendu.

━━ EXPERT ━━
- Questions ouvertes sans guidage, transfert de compétences vers une situation nouvelle.
- Rappel de cours : absent ou très succinct (2 lignes max).
- Exercice d'application : données brutes à extraire soi-même.
- Mise en situation complexe avec plusieurs informations à croiser.
- Problème ouvert ambitieux — mais toujours réaliste pour un élève de Bac Pro.

Dans tous les cas : JAMAIS de calcul hors programme Bac Pro, JAMAIS de piège inutile.
La correction détaillée doit être adaptée au même niveau (plus ou moins de détails selon le niveau).

Structure de sortie (Markdown) :
1. **Rappel de cours** (adapté au niveau)
2. **Exercice d'application** (adapté au niveau)
3. **Exercice de mise en situation** (contexte professionnel de la filière)
4. **Problème ouvert** (adapté au niveau)
5. **Corrections détaillées** (avec le niveau de détail approprié)\
"""

# Descriptifs courts affichés dans l'UI
NIVEAUX_DIFFICULTE = {
    "🟢 Débutant":  "Questions très guidées, micro-étapes, résultats intermédiaires donnés.",
    "🟡 Moyen":     "Semi-guidé, quelques repères fournis, mise en situation simple.",
    "🟠 Confirmé":  "Autonome, mise en situation réaliste, raisonnement attendu.",
    "🔴 Expert":    "Transfert de compétences, données brutes, problème ouvert ambitieux.",
}


def build_prompt_exercices(niveau, categorie, matiere, chapitre, consignes, filiere="", difficulte="🟡 Moyen"):
    ctx = build_contexte_filiere(filiere)
    diff_label = difficulte.split(" ", 1)[-1].upper()  # ex: "MOYEN"
    user = (
        f"Génère un contenu pédagogique de niveau **{diff_label}** pour :\n"
        f"- Niveau scolaire : {niveau} ({categorie})\n"
        f"- Matière : {matiere}\n"
        f"- Chapitre : {chapitre}\n"
        f"{ctx}"
        f"- Instructions : {consignes or 'Aucune'}\n\n"
        f"Applique scrupuleusement les consignes du niveau {diff_label} définies dans tes instructions."
    )
    return SYSTEM_EXERCICES, user


# ============================================================
# CO-INTERVENTION — Référentiels pro par filière
# ============================================================

THEMES_COINTERVENTION = {
    "MCVB": {
        "nom_complet": "Métiers du Commerce et de la Vente — option B",
        "themes": {
            "Prix d'achat, prix de vente et marge commerciale": {
                "formules": "Marge = PV HT − PA HT  ·  Taux de marge = Marge / PA HT  ·  Taux de marque = Marge / PV HT",
                "contexte": "Un commercial calcule sa marge sur un produit, compare deux fournisseurs.",
                "documents": "bon de commande, catalogue fournisseur, tableau comparatif"
            },
            "Taux de remise et prix remisé": {
                "formules": "Prix remisé = Prix initial × (1 − taux)  ·  Montant remise = Prix initial × taux",
                "contexte": "Négociation client, grille de remises selon quantité commandée.",
                "documents": "grille tarifaire, bon de commande, facture proforma"
            },
            "TVA — calcul HT / TTC": {
                "formules": "TTC = HT × (1 + taux TVA)  ·  HT = TTC / (1 + taux TVA)  ·  TVA = TTC − HT",
                "contexte": "Établissement de devis, vérification de factures clients et fournisseurs.",
                "documents": "facture, devis, ticket de caisse"
            },
            "Taux d'évolution du chiffre d'affaires": {
                "formules": "Taux d'évolution = (VF − VI) / VI  ·  VF = VI × (1 + t)",
                "contexte": "Évolution du CA, suivi des ventes mensuelles, comparaison N/N-1.",
                "documents": "tableau de bord commercial, graphique de ventes"
            },
            "Coefficient multiplicateur": {
                "formules": "CM = PV TTC / PA HT  ·  PV TTC = PA HT × CM",
                "contexte": "Calcul rapide du prix de vente à partir du prix d'achat.",
                "documents": "catalogue, liste de prix, étiquettes"
            },
            "Commissionnement et salaire variable": {
                "formules": "Commission = CA réalisé × taux  ·  Salaire total = Fixe + Commission",
                "contexte": "Calcul de la rémunération d'un commercial selon ses performances.",
                "documents": "bulletin de salaire simplifié, tableau de suivi CA"
            },
            "Budget prévisionnel et écarts": {
                "formules": "Écart = Réalisé − Prévu  ·  Taux d'écart = Écart / Prévu",
                "contexte": "Analyse des résultats commerciaux vs objectifs fixés.",
                "documents": "tableau de bord, fiche de suivi objectifs"
            },
        }
    },
    "MCVA": {
        "nom_complet": "Métiers du Commerce et de la Vente — option A",
        "themes": {
            "TVA — calcul HT / TTC": {
                "formules": "TTC = HT × (1 + taux TVA)  ·  HT = TTC / (1 + taux TVA)  ·  TVA = TTC − HT",
                "contexte": "Gestion des prix en rayon, vérification des étiquettes, encaissement.",
                "documents": "étiquettes de rayon, facture fournisseur, ticket de caisse"
            },
            "Coefficient multiplicateur": {
                "formules": "CM = PV TTC / PA HT  ·  PV TTC = PA HT × CM",
                "contexte": "Calcul du prix de vente d'un produit à partir du tarif fournisseur.",
                "documents": "catalogue fournisseur, liste de prix magasin"
            },
            "Taux de marque et marge": {
                "formules": "Taux de marque = Marge / PV HT  ·  Marge = PV HT − PA HT",
                "contexte": "Analyse de la rentabilité d'un rayon ou d'une famille de produits.",
                "documents": "tableau de gestion rayon, inventaire"
            },
            "Gestion des stocks — taux de rotation": {
                "formules": "Taux de rotation = CA / Stock moyen  ·  Stock moyen = (SI + SF) / 2",
                "contexte": "Suivi des entrées/sorties de marchandises, optimisation du réapprovisionnement.",
                "documents": "bon de livraison, fiche de stock, inventaire"
            },
            "Taux de transformation et panier moyen": {
                "formules": "Taux de transformation = Nb acheteurs / Nb visiteurs  ·  Panier moyen = CA / Nb acheteurs",
                "contexte": "Analyse de la performance d'un point de vente.",
                "documents": "comptage client, tableau de bord magasin"
            },
            "Surfaces et linéaires": {
                "formules": "Surface de vente = longueur × largeur  ·  Linéaire développé = nb niveaux × longueur",
                "contexte": "Implantation d'un rayon, calcul de la surface allouée à une famille de produits.",
                "documents": "plan de masse, schéma d'implantation"
            },
            "Taux d'évolution du chiffre d'affaires": {
                "formules": "Taux d'évolution = (CA N − CA N-1) / CA N-1",
                "contexte": "Comparaison des performances du magasin d'une période à l'autre.",
                "documents": "tableau de bord, rapport d'activité"
            },
        }
    },
    "AGORA": {
        "nom_complet": "Assistance à la Gestion des Organisations",
        "themes": {
            "Calcul de salaire brut et net": {
                "formules": "Salaire brut = Taux horaire × Nb heures  ·  Salaire net = Brut − Cotisations salariales",
                "contexte": "Aide à la paie, vérification de bulletins de salaire simplifiés.",
                "documents": "bulletin de salaire, contrat de travail, fiche de pointage"
            },
            "TVA — calcul HT / TTC": {
                "formules": "TTC = HT × (1 + taux TVA)  ·  HT = TTC / (1 + taux TVA)  ·  TVA collectée − TVA déductible",
                "contexte": "Établissement et vérification de factures, déclaration de TVA simplifiée.",
                "documents": "facture, bon de commande, déclaration TVA"
            },
            "Intérêts simples et emprunts": {
                "formules": "Intérêts = Capital × taux × durée  ·  Montant remboursé = Capital + Intérêts",
                "contexte": "Calcul du coût d'un emprunt bancaire pour financer du matériel.",
                "documents": "tableau d'amortissement simplifié, offre de prêt"
            },
            "Budget et suivi des dépenses": {
                "formules": "Solde = Recettes − Dépenses  ·  Taux de consommation = Dépenses / Budget prévu",
                "contexte": "Gestion du budget d'une organisation, suivi des dépenses réelles vs prévisions.",
                "documents": "tableau de bord budgétaire, relevé bancaire"
            },
            "Cotisations sociales": {
                "formules": "Cotisation = Base × taux  ·  Part salariale et part patronale",
                "contexte": "Compréhension de la fiche de paie, calcul des charges sociales.",
                "documents": "bulletin de salaire, tableau des taux de cotisation"
            },
            "Taux d'évolution et indices": {
                "formules": "Taux d'évolution = (VF − VI) / VI  ·  Indice = (Valeur / Valeur de base) × 100",
                "contexte": "Suivi de l'évolution des indicateurs RH (absentéisme, masse salariale).",
                "documents": "tableau de bord RH, rapport annuel"
            },
            "Facturation et remises": {
                "formules": "Montant net = Montant brut × (1 − remise)  ·  Escompte = Montant × taux",
                "contexte": "Établissement de devis et factures, calcul de remises et d'escomptes.",
                "documents": "devis, facture, bon de commande"
            },
        }
    },
    "ASSP": {
        "nom_complet": "Accompagnement, Soins et Services à la Personne",
        "themes": {
            "Calcul de doses médicamenteuses": {
                "formules": "Dose à administrer = Dose prescrite / Concentration  ·  Règle de trois",
                "contexte": "Préparation de médicaments, dosages pour les résidents.",
                "documents": "fiche de prescription, protocole de soins"
            },
            "Calcul de dilutions": {
                "formules": "C1 × V1 = C2 × V2  ·  Volume à prélever = (C2 × V2) / C1",
                "contexte": "Préparation de solutions désinfectantes, produits d'hygiène dilués.",
                "documents": "protocole d'hygiène, fiche technique produit"
            },
            "Gestion du temps et plannings": {
                "formules": "Durée = Heure de fin − Heure de début  ·  Répartition horaire en pourcentage",
                "contexte": "Organisation des tournées d'aide à domicile, plannings.",
                "documents": "planning hebdomadaire, feuille de route"
            },
            "Calcul de coûts de prise en charge": {
                "formules": "Coût total = Tarif horaire × Nb heures  ·  Reste à charge = Coût − Aides",
                "contexte": "Estimation du coût d'une prise en charge, calcul des aides financières.",
                "documents": "devis de prise en charge, dossier APA"
            },
            "IMC et indicateurs de santé": {
                "formules": "IMC = Poids (kg) / Taille² (m)  ·  Interprétation des valeurs normales",
                "contexte": "Suivi nutritionnel des résidents, surveillance du poids.",
                "documents": "fiche de suivi patient, courbe de poids"
            },
            "Statistiques sur données de santé": {
                "formules": "Moyenne, médiane, étendue  ·  Lecture et construction de graphiques",
                "contexte": "Analyse des données de suivi (tension, glycémie, fréquence cardiaque).",
                "documents": "fiche de surveillance, graphique de suivi"
            },
        }
    },
}

SYSTEM_COINTERVENTION = """\
Tu es un professeur qui anime une séance de CO-INTERVENTION entre mathématiques et matière professionnelle en Bac Pro.
La co-intervention relie explicitement les notions mathématiques aux situations professionnelles réelles du référentiel.

Tes exercices doivent TOUJOURS :
- Partir d'un document professionnel réaliste (facture, bon de commande, bulletin de salaire, planning, étiquette...)
- Utiliser le vocabulaire professionnel exact de la filière
- Montrer explicitement POURQUOI les maths sont utiles dans le métier
- Être ancrés dans des situations que l'élève vivra réellement en stage ou en emploi
- Rappeler la formule mathématique dans son contexte pro (jamais de façon abstraite)

Structure de sortie (Markdown) :

### 📄 Document professionnel
[Reproduire un document réaliste avec données chiffrées : facture, tableau, bulletin, bon de commande...]

### 🎯 Mise en situation
[2-3 lignes de contexte professionnel concret — qui fait quoi, dans quelle entreprise]

### 📐 Rappel de la formule
[Formule + signification de chaque terme dans le vocabulaire professionnel]

### ✏️ Exercices
[3 à 5 questions progressives, du plus guidé au plus autonome]

### ✅ Corrigé détaillé
[Correction complète avec justifications dans le vocabulaire professionnel]\
"""


def build_prompt_cointervention(niveau, filiere, theme, consignes, difficulte="🟡 Moyen"):
    """Construit le prompt pour un exercice de co-intervention."""
    data_filiere = THEMES_COINTERVENTION.get(filiere, {})
    nom_filiere  = data_filiere.get("nom_complet", filiere)
    data_theme   = data_filiere.get("themes", {}).get(theme, {})
    formules     = data_theme.get("formules", "")
    contexte     = data_theme.get("contexte", "")
    documents    = data_theme.get("documents", "")
    diff_label   = difficulte.split(" ", 1)[-1].upper()
    conseils_diff = {
        "DÉBUTANT":  "formules rappelées, calculs en une étape, données déjà extraites du document.",
        "MOYEN":     "quelques étapes guidées, document partiellement exploité.",
        "CONFIRMÉ":  "questions autonomes, document complet à analyser.",
        "EXPERT":    "situation nouvelle, données brutes, raisonnement et conclusion attendus.",
    }.get(diff_label, "")

    user = (
        f"Génère un exercice de CO-INTERVENTION de niveau **{diff_label}** pour :\n"
        f"- Filière : {nom_filiere}\n"
        f"- Niveau scolaire : {niveau} (Bac Pro)\n"
        f"- Thème : {theme}\n"
        f"- Formules clés : {formules}\n"
        f"- Contexte professionnel : {contexte}\n"
        f"- Documents supports : {documents}\n"
        f"- Instructions complémentaires : {consignes or 'Aucune'}\n\n"
        f"Niveau {diff_label} : {conseils_diff}"
    )
    return SYSTEM_COINTERVENTION, user
    ctx = build_contexte_filiere(filiere)
    diff_label = difficulte.split(" ", 1)[-1].upper()  # ex: "MOYEN"
    user = (
        f"Génère un contenu pédagogique de niveau **{diff_label}** pour :\n"
        f"- Niveau scolaire : {niveau} ({categorie})\n"
        f"- Matière : {matiere}\n"
        f"- Chapitre : {chapitre}\n"
        f"{ctx}"
        f"- Instructions : {consignes or 'Aucune'}\n\n"
        f"Applique scrupuleusement les consignes du niveau {diff_label} définies dans tes instructions."
    )
    return SYSTEM_EXERCICES, user


def build_prompt_ccf_entrainement(niveau, categorie, matiere, chapitre, consignes, filiere="", avec_corrige=True, chapitre_b=""):
    ctx = build_contexte_filiere(filiere)
    bloc_corrige = "\n### CORRIGÉ DÉTAILLÉ *(document professeur)*\nCorrection complète de chaque question.\n" if avec_corrige else ""
    chap_a_label = f"Partie A : {chapitre}"
    chap_b_label = f"Partie B : {chapitre_b}" if chapitre_b else "Partie B : thème mathématique complémentaire au choix"
    return f"""Tu es un professeur de mathématiques expert en Bac Pro et en évaluation CCF conforme au BO.

Génère un SUJET D'ENTRAÎNEMENT AU CCF pour :
- Niveau : {niveau} ({categorie})
- Matière : {matiere}
- {chap_a_label}
- {chap_b_label}
{ctx}
- Instructions : {consignes or 'Aucune'}

## RÈGLES STRICTES DU FORMAT CCF OFFICIEL :

1. La MISE EN SITUATION présente le contexte professionnel avec données chiffrées réalistes.
2. La PROBLÉMATIQUE est UNE SEULE question (avec point d'interrogation ?) encadrée en gras, placée juste après la mise en situation. C'est la seule question du document avec un "?".
3. Les questions sont numérotées et chacune porte le nom de la compétence évaluée en majuscules : S'APPROPRIER / RÉALISER / ANALYSER-RAISONNER / VALIDER / COMMUNIQUER.
4. Chaque question se termine par des lignes de réponse (______).
5. Les questions servent à répondre progressivement à la problématique.
6. La dernière question COMMUNIQUER demande de répondre à la problématique initiale.
7. L'évaluation est sur 10 points (pas 20), répartis en niveaux 0/1/2 par compétence.
8. Les deux parties s'appuient sur LA MÊME mise en situation professionnelle.

## STRUCTURE À RESPECTER :

### MISE EN SITUATION PROFESSIONNELLE
[Description du contexte, de l'entreprise/structure, avec données chiffrées et document support : tableau ou graphique]

### PROBLÉMATIQUE
**[UNE SEULE question centrale se terminant par ?]**

### PARTIE A — [Titre lié à : {chapitre}]

1. **S'APPROPRIER**
- [question]
______

2. **RÉALISER**
- [question]
______

3. **ANALYSER / RAISONNER**
- [question]
______

4. **RÉALISER** 🛎️ APPELER L'EXAMINATEUR
- [question nécessitant outil/calculatrice]
______

5. **VALIDER**
- [question de vérification/interprétation]
______

6. **COMMUNIQUER**
- Répondre à la [première partie de la] problématique.
______

### PARTIE B — [Titre lié à : {chapitre_b if chapitre_b else "second thème mathématique"}]
[Mêmes règles, questions numérotées avec compétences]
{bloc_corrige}
Réponds entièrement en Markdown."""


def build_prompt_ccf_officiel(niveau, categorie, matiere, chapitre, consignes, filiere="", duree="45 min", num_sit="1", avec_corrige=True, chapitre_b=""):
    ctx = build_contexte_filiere(filiere)
    nom_filiere = CONTEXTES_FILIERES[filiere]["nom_complet"] if filiere in CONTEXTES_FILIERES else filiere
    bloc_corrige = "\n### CORRIGÉ DÉTAILLÉ *(document professeur — NE PAS DISTRIBUER)*\n[Correction complète question par question avec justifications]\n" if avec_corrige else ""
    chap_a_label = f"Partie A : {chapitre}"
    chap_b_label = f"Partie B : {chapitre_b}" if chapitre_b else "Partie B : thème mathématique complémentaire en cohérence avec la situation"
    return f"""Tu es un professeur de mathématiques expert en Bac Pro et en évaluation CCF conforme au BO de l'Éducation Nationale.

Génère un SUJET DE CCF OFFICIEL complet et prêt à imprimer pour :
- Niveau : {niveau} ({categorie})
- Filière : {nom_filiere}
- Matière : {matiere}
- {chap_a_label}
- {chap_b_label}
- Situation d'évaluation n° : {num_sit}
- Durée : {duree}
{ctx}
- Instructions complémentaires : {consignes or 'Aucune'}

## RÈGLES STRICTES DU FORMAT CCF OFFICIEL ÉDUCATION NATIONALE :

1. La MISE EN SITUATION présente le contexte professionnel avec une image mentale de la structure/entreprise, des données chiffrées réalistes et un document support (tableau de données ou graphique nommé).
2. La PROBLÉMATIQUE est UNE SEULE question (avec ?) encadrée en gras juste après la mise en situation. C'est la SEULE question du document avec un point d'interrogation.
3. Les questions sont numérotées (1. 2. 3. ...) et chacune porte le nom de la compétence BO en majuscules gras : **S'APPROPRIER** / **RÉALISER** / **ANALYSER / RAISONNER** / **VALIDER** / **COMMUNIQUER**.
4. Toutes les 5 compétences doivent être évaluées au moins une fois.
5. Les questions comportent des lignes de réponse (______).
6. La dernière question **COMMUNIQUER** demande explicitement de répondre à la problématique initiale.
7. Quand une question nécessite la calculatrice ou un outil, ajouter : 🛎️ APPELER L'EXAMINATEUR
8. L'évaluation est notée sur /10 avec niveaux d'acquisition 0/1/2 par compétence (PAS de barème en points).
9. Les deux parties s'appuient sur LA MÊME mise en situation professionnelle.

## STRUCTURE OBLIGATOIRE :

### MISE EN SITUATION PROFESSIONNELLE
[Présentation de la structure professionnelle ({nom_filiere}), contexte détaillé, données chiffrées]
[Document support : tableau ou données nommées]

L'usage de la calculatrice avec mode examen actif est autorisé.

### PROBLÉMATIQUE
**[Question centrale unique se terminant par ? — clairement encadrée]**

### PARTIE A — [Titre lié à : {chapitre}]

1. **S'APPROPRIER**
- [question d'appropriation de la situation]
______

2. **RÉALISER**
- [question de calcul/réalisation]
______

3. **ANALYSER / RAISONNER**
- [question d'analyse]
______

4. **RÉALISER** 🛎️ APPELER L'EXAMINATEUR
- [question avec outil numérique]
______

5. **VALIDER**
- [question de validation/interprétation]
______

6. **COMMUNIQUER**
- Répondre à la première partie de la problématique.
______

### PARTIE B — [Titre lié à : {chapitre_b if chapitre_b else "second thème mathématique"}]
[Mêmes règles]
{bloc_corrige}
Réponds entièrement en Markdown avec mise en page soignée et professionnelle."""


def build_prompt_correction(bareme, ton, niveau, matiere, note_sur):
    return f"""Tu es un professeur correcteur expert.

Contexte :
- Niveau : {niveau or 'Non précisé'}
- Matière : {matiere or 'Non précisée'}
- Barème : {bareme or f'Non fourni — évalue sur {note_sur}'}
- Ton : {ton}
- Note sur : {note_sur}

Mission :
1. Transcris le texte manuscrit visible.
2. Identifie et explique chaque erreur avec pédagogie.
3. Attribue une note partielle par question.
4. Calcule la note globale /{note_sur}.
5. Rédige une appréciation finale ({ton}) de 3 à 5 lignes.

Réponds en Markdown avec sections claires."""


# ============================================================
# 4. APPEL API GEMINI
# ============================================================

def call_gemini(api_key, prompt, image=None):
    genai.configure(api_key=api_key.strip())
    # Gestion du tuple (system_instruction, user_prompt) retourné par build_prompt_exercices
    if isinstance(prompt, tuple):
        system_instruction, user_prompt = prompt
        model = genai.GenerativeModel(
            model_name="gemini-2.5-flash",
            system_instruction=system_instruction
        )
        content = user_prompt
    else:
        model = genai.GenerativeModel(model_name="gemini-2.5-flash")
        content = prompt
    if image:
        img = PIL.Image.open(image)
        response = model.generate_content([content, img])
    else:
        response = model.generate_content(content)
    if response and response.text:
        return response.text
    return "L'IA a répondu mais le texte est vide. Réessayez."


# ============================================================
# 5. EXPORT WORD — VERSION STANDARD
# ============================================================

def clean_math(text):
    """
    Convertit la notation LaTeX/Markdown en texte lisible dans Word.
    Gère : \\[...\\], $$...$$, $...$, \\frac, ^, _, commandes LaTeX, backticks.
    """
    import re as _re
    # Supprimer balises ```markdown et ``` parasites
    text = _re.sub(r'```[a-zA-Z]*', '', text)
    # Blocs \[ ... \] et \( ... \)
    text = _re.sub(r'\\\[(.+?)\\\]', lambda m: m.group(1).strip(), text, flags=_re.DOTALL)
    text = _re.sub(r'\\\((.+?)\\\)', lambda m: m.group(1).strip(), text, flags=_re.DOTALL)
    # Blocs $$ ... $$
    text = _re.sub(r'\$\$(.+?)\$\$', lambda m: m.group(1).strip(), text, flags=_re.DOTALL)
    # Inline $ ... $
    text = _re.sub(r'\$([^$\n]+?)\$', lambda m: m.group(1).strip(), text)
    # \frac{a}{b} → a/b
    text = _re.sub(r'\\frac\{([^}]+)\}\{([^}]+)\}', r'\1/\2', text)
    # Exposants ^{...} → ^...
    text = _re.sub(r'\^\{([^}]+)\}', r'^\1', text)
    # Indices _{...} → _...
    text = _re.sub(r'_\{([^}]+)\}', r'_\1', text)
    # Commandes LaTeX usuelles → Unicode
    replacements = {
        r'\times': '×', r'\cdot': '·', r'\leq': '≤', r'\geq': '≥',
        r'\neq': '≠', r'\approx': '≈', r'\infty': '∞', r'\pi': 'π',
        r'\alpha': 'α', r'\beta': 'β', r'\sigma': 'σ', r'\mu': 'μ',
        r'\rightarrow': '→', r'\leftarrow': '←', r'\in': '∈',
        r'\subset': '⊂', r'\cup': '∪', r'\cap': '∩',
    }
    for latex, unicode_char in replacements.items():
        text = text.replace(latex, unicode_char)
    # Supprimer les backslashes restants isolés
    text = _re.sub(r'\\([a-zA-Z]+)', r'\1', text)
    return text


def parse_md_tables(md_text):
    """
    Transforme un texte Markdown en liste de blocs :
    - ('text', 'contenu')     pour les lignes normales
    - ('table', [[row], ...]) pour les tableaux Markdown
    """
    import re as _re
    lines = md_text.split('\n')
    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i]
        # Début d'un tableau : ligne commençant par |
        if line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                raw = lines[i].strip()
                # Ignorer les lignes séparateurs (|---|---|)
                if _re.match(r'^\|[\s\-:|]+\|', raw):
                    i += 1
                    continue
                # Découper les cellules
                cells = [c.strip() for c in raw.strip('|').split('|')]
                table_lines.append(cells)
                i += 1
            if table_lines:
                blocks.append(('table', table_lines))
        else:
            blocks.append(('text', line))
            i += 1
    return blocks


def add_md_table_to_doc(doc, rows, font_size=10):
    """Crée un vrai tableau Word à partir d'une liste de lignes [[cell, cell, ...]]."""
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(max_cols):
            cell = table.cell(i, j)
            text = clean_math(row[j]) if j < len(row) else ""
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(text)
            run.font.size = Pt(font_size)
            run.font.name = "Arial"
            # 1ère ligne en gras (en-tête)
            if i == 0:
                run.bold = True
    doc.add_paragraph()


def render_inline(paragraph, text):
    """Écrit une ligne avec gras/italique et nettoie les $ mathématiques."""
    text = clean_math(text)
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            paragraph.add_run(part[1:-1]).italic = True
        else:
            paragraph.add_run(part)


def markdown_to_docx(md_text, titre="Document"):
    doc = DocxDocument()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)

    h = doc.add_heading(titre, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    blocks = parse_md_tables(md_text)

    for kind, content in blocks:
        if kind == "table":
            add_md_table_to_doc(doc, content, font_size=10)
            continue

        line = content.rstrip()
        if line.startswith("### "):
            doc.add_heading(clean_math(line[4:]), level=3)
        elif line.startswith("## "):
            doc.add_heading(clean_math(line[3:]), level=2)
        elif line.startswith("# "):
            doc.add_heading(clean_math(line[2:]), level=1)
        elif line.startswith("---"):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
        elif re.match(r"^\d+\.", line):
            p = doc.add_paragraph(style="List Number")
            render_inline(p, re.sub(r"^\d+\.\s*", "", line))
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            render_inline(p, line[2:])
        elif line == "":
            doc.add_paragraph("")
        else:
            p = doc.add_paragraph()
            render_inline(p, line)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ============================================================
# 6. EXPORT WORD — VERSION OFFICIELLE CCF
# ============================================================

def set_cell_border(cell):
    """Ajoute une bordure fine sur toutes les faces d'une cellule."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "AAAAAA")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def set_run(paragraph, text, bold=False, size=10):
    """Ajoute un run stylé — évite le crash sur paragraphe vide."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return run


def fill_cell(cell, text, bold=False, size=10):
    """Remplit une cellule proprement sans crasher sur runs vides."""
    p = cell.paragraphs[0]
    p.clear()
    set_run(p, text, bold=bold, size=size)
    set_cell_border(cell)


def parse_questions_competences(content_md):
    """
    Parcourt le Markdown du sujet CCF et construit un dictionnaire
    { nom_compétence_normalisé : [liste de références questions ex: 'A.1', 'B.2'] }
    
    Détecte les patterns :
      ### PARTIE A  →  lettre de partie courante = 'A'
      1. **S'APPROPRIER**  →  question A.1 → compétence S'approprier
    """
    # Mapping des mots-clés détectés vers les noms officiels des compétences
    COMP_KEYWORDS = {
        "S'APPROPRIER":        "S'approprier",
        "APPROPRIER":          "S'approprier",
        "ANALYSER / RAISONNER":"Analyser / Raisonner",
        "ANALYSER":            "Analyser / Raisonner",
        "RAISONNER":           "Analyser / Raisonner",
        "RÉALISER":            "Réaliser",
        "REALISER":            "Réaliser",
        "VALIDER":             "Valider",
        "COMMUNIQUER":         "Communiquer",
    }

    # Initialisation : une liste vide par compétence officielle
    result = {
        "S'approprier":       [],
        "Analyser / Raisonner": [],
        "Réaliser":           [],
        "Valider":            [],
        "Communiquer":        [],
    }

    current_part = ""

    for line in content_md.split("\n"):
        line_s = line.strip()

        # Détecter changement de partie : ### PARTIE A, ## PARTIE B, etc.
        part_match = re.match(r"#+\s+PARTIE\s+([A-Z])", line_s, re.IGNORECASE)
        if part_match:
            current_part = part_match.group(1).upper()
            continue

        # Si pas encore de partie détectée, on ne peut pas construire de référence
        if not current_part:
            continue

        # Détecter une question numérotée avec sa compétence en gras
        # Exemples : "1. **S'APPROPRIER**"  "4. **RÉALISER** 🛎️ APPELER..."
        q_match = re.match(r"^(\d+)\.\s+\*\*([^*]+)\*\*", line_s)
        if q_match:
            q_num   = q_match.group(1)
            comp_raw = q_match.group(2).strip().upper()
            q_ref   = f"{current_part}.{q_num}"

            # Chercher la compétence correspondante
            for keyword, comp_name in COMP_KEYWORDS.items():
                if keyword in comp_raw:
                    if q_ref not in result[comp_name]:
                        result[comp_name].append(q_ref)
                    break

    return result


def generate_ccf_officiel_docx(content_md, metadata, nom_etablissement="Mon Établissement"):
    try:
        doc = DocxDocument()

        # Marges réduites pour plus d'espace
        for section in doc.sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        doc.styles["Normal"].font.name = "Arial"
        doc.styles["Normal"].font.size = Pt(11)

        # ── EN-TÊTE : tableau 3 colonnes avec bordures visibles
        # Conforme au document de référence de l'établissement
        table_h = doc.add_table(rows=3, cols=3)
        table_h.style = "Table Grid"  # bordures visibles sur toutes les cellules

        # Largeurs colonnes en DXA (total 9360 = contenu avec marges 1")
        col_widths = [2620, 4120, 2620]
        for row in table_h.rows:
            for j, cell in enumerate(row.cells):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = OxmlElement("w:tcW")
                tcW.set(qn("w:w"), str(col_widths[j]))
                tcW.set(qn("w:type"), "dxa")
                tcPr.append(tcW)

        # ── Ligne 1 : logo gauche | titre CCF | logo matière ──
        cell_l = table_h.cell(0, 0)
        logo_rep = os.path.join(APP_DIR, "logo_republique.png")
        if os.path.exists(logo_rep):
            cell_l.paragraphs[0].add_run().add_picture(logo_rep, width=Inches(1.2))
        else:
            r_fallback = cell_l.paragraphs[0].add_run("ACADÉMIE DE CRÉTEIL")
            r_fallback.bold = True
            r_fallback.font.size = Pt(10)

        cell_c = table_h.cell(0, 1)
        p_c = cell_c.paragraphs[0]
        p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        def _rb(para, text, size=10):
            r = para.add_run(text)
            r.bold = True
            r.font.size = Pt(size)
            return r
        _rb(p_c, "CONTRÔLE EN COURS DE FORMATION\n", 12)
        _rb(p_c, "Baccalauréat professionnel\n", 11)
        _rb(p_c, f"{metadata.get('matiere', '')}\n", 11)
        _rb(p_c, f"Situation d'évaluation n°{metadata.get('num_situation', '...')}\n", 10)
        _rb(p_c, f"Intitulé du diplôme : {'.' * 31}\n", 10)
        _rb(p_c, f"Durée : {metadata.get('duree', '45 min')}", 10)

        cell_r = table_h.cell(0, 2)
        p_r = cell_r.paragraphs[0]
        p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        logo_mat = os.path.join(APP_DIR, "logo_matiere.png")
        if os.path.exists(logo_mat):
            p_r.add_run().add_picture(logo_mat, width=Inches(0.8))

        # ── Ligne 2 : vide | Nom/Prénom candidat (centré) | vide ──
        # Conforme au document de référence : colonne centre uniquement
        cell_nom = table_h.cell(1, 1)
        p_nom = cell_nom.paragraphs[0]
        p_nom.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _rb(p_nom, "Nom, Prénom du candidat : _______________________________________________", 10)

        # ── Ligne 3 : Date | Établissement | vide ──
        cell_date = table_h.cell(2, 0)
        cell_date.paragraphs[0].add_run("Date : .......................").font.size = Pt(10)

        cell_etab = table_h.cell(2, 1)
        p_etab = cell_etab.paragraphs[0]
        _rb(p_etab, f"Nom de l'établissement : {nom_etablissement}", 10)

        doc.add_paragraph()

        # ── NOTICE CALCULATRICE ────────────────────────────────
        logo_appel_path = os.path.join(APP_DIR, "logo_appel.png")
        p_notice = doc.add_paragraph()
        if os.path.exists(logo_appel_path):
            p_notice.add_run().add_picture(logo_appel_path, width=Inches(0.53))
            p_notice.add_run("  ")
        p_notice.add_run("Dans la suite du document, ce symbole signifie « Appeler l'examinateur ».").font.size = Pt(10)

        for ligne_calc in [
            "L'usage de la calculatrice avec mode examen actif est autorisé.",
            "L'usage de la calculatrice sans mémoire, « type collège » est autorisé.",
            "L'échange de calculatrices entre les candidats pendant l'épreuve est interdit.",
        ]:
            doc.add_paragraph().add_run(ligne_calc).font.size = Pt(10)

        doc.add_paragraph()

        # ── CONTENU DU SUJET ──────────────────────────────────
        # Filtre les lignes parasites que l'IA peut générer avant le vrai sujet
        skip_keywords = [
            "Établissement :", "Baccalauréat Professionnel —", "Épreuve E3",
            "Calculatrice autorisée", "Ministère de l'Éducation",
            "sujet de CCF officiel", "Voici le sujet", "stricte conformité",
            "accompagné de", "fiche d'évaluation et de son corrigé",
            "Filière :", "Niveau :", "Situation d'évaluation n°",
            "Durée :", "Nom & Prénom",
        ]

        # Filtrer le contenu avant de le parser
        filtered_lines = []
        contenu_commence = False
        for line in content_md.split("\n"):
            ls = line.rstrip()
            if not contenu_commence:
                if any(ls.startswith(m) for m in [
                    "### MISE EN SITUATION", "### PARTIE", "### PROBLÉMATIQUE",
                    "## MISE EN SITUATION", "## PARTIE", "## PROBLÉMATIQUE",
                    "#### MISE", "#### PARTIE",
                ]):
                    contenu_commence = True
                else:
                    continue
            if any(kw in ls for kw in skip_keywords):
                continue
            filtered_lines.append(ls)

        filtered_md = "\n".join(filtered_lines)

        # Parser les blocs (texte et tableaux Markdown)
        blocks = parse_md_tables(filtered_md)

        for kind, content in blocks:
            # ── Vrai tableau Word ──────────────────────────────
            if kind == "table":
                add_md_table_to_doc(doc, content, font_size=10)
                continue

            line = content

            if line.startswith("#### "):
                doc.add_heading(clean_math(line[5:].strip()), level=3)
            elif line.startswith("### "):
                doc.add_heading(clean_math(line[4:].strip()), level=2)
            elif line.startswith("## "):
                doc.add_heading(clean_math(line[3:].strip()), level=2)
            elif "APPELER L'EXAMINATEUR" in line.upper() or "🛎️" in line:
                p = doc.add_paragraph()
                logo_appel = os.path.join(APP_DIR, "logo_appel.png")
                if os.path.exists(logo_appel):
                    p.add_run().add_picture(logo_appel, width=Inches(0.25))
                    p.add_run("  ")
                run = p.add_run("APPELER L'EXAMINATEUR")
                run.bold = True
                run.font.size = Pt(11)
                pPr = p._p.get_or_add_pPr()
                pBdr = OxmlElement("w:pBdr")
                for side in ("top", "left", "bottom", "right"):
                    border = OxmlElement(f"w:{side}")
                    border.set(qn("w:val"), "single")
                    border.set(qn("w:sz"), "6")
                    border.set(qn("w:space"), "4")
                    border.set(qn("w:color"), "4A6CF7")
                    pBdr.append(border)
                pPr.append(pBdr)
            elif line.startswith("---"):
                continue
            elif line == "":
                doc.add_paragraph("")
            else:
                p = doc.add_paragraph()
                render_inline(p, line)

        # ── FICHE D'ÉVALUATION OFFICIELLE (page 2) ───────────
        doc.add_page_break()

        titre_fiche = doc.add_heading("FICHE INDIVIDUELLE D'ÉVALUATION", level=1)
        titre_fiche.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # En-tête fiche
        table_ef = doc.add_table(rows=2, cols=2)
        fill_cell(table_ef.cell(0, 0),
                  f"Session : {metadata.get('annee_scolaire','2025/2026')} — Établissement : {nom_etablissement} — Académie : Créteil",
                  bold=True, size=9)
        fill_cell(table_ef.cell(0, 1),
                  f"Spécialité : {metadata.get('matiere','')} — Évaluateur : _______________ — Date : _______________",
                  size=9)
        fill_cell(table_ef.cell(1, 0),
                  f"Situation n° {metadata.get('num_situation','1')} — {metadata.get('filiere','')}",
                  bold=True, size=9)
        fill_cell(table_ef.cell(1, 1), "", size=9)

        doc.add_paragraph()
        cand_p = doc.add_paragraph()
        set_run(cand_p, "Nom et prénom du candidat : _______________________________________________",
                bold=True, size=10)

        doc.add_paragraph()

        # 1. Capacités et connaissances évaluées
        cap_titre = doc.add_paragraph()
        set_run(cap_titre, "1.  Liste des capacités et connaissances évaluées", bold=True, size=11)

        table_cap = doc.add_table(rows=2, cols=2)
        fill_cell(table_cap.cell(0, 0), "Capacités", bold=True, size=10)
        fill_cell(table_cap.cell(0, 1), "(voir section Fiche d'évaluation du sujet généré)", size=10)
        fill_cell(table_cap.cell(1, 0), "Connaissances", bold=True, size=10)
        fill_cell(table_cap.cell(1, 1), "(voir section Fiche d'évaluation du sujet généré)", size=10)

        doc.add_paragraph()

        # 2. Tableau d'évaluation officiel
        eval_titre = doc.add_paragraph()
        set_run(eval_titre, "2.  Évaluation", bold=True, size=11)

        # Calcul nb lignes : 1 en-tête + 1 ligne par indicateur + 1 ligne note
        total_rows = 1
        for comp in COMPETENCES_CCF:
            total_rows += len(comp["indicateurs"])
        total_rows += 1  # ligne note finale

        grid = doc.add_table(rows=total_rows, cols=5)

        # En-têtes colonnes
        h_texts = ["Compétences", "", "Indicateurs / Capacités", "Questions", "Appréciation\n(0 / 1 / 2)"]
        for j, h in enumerate(h_texts):
            fill_cell(grid.cell(0, j), h, bold=True, size=9)

        # Parsing des numéros de questions depuis le contenu généré
        questions_par_comp = parse_questions_competences(content_md)

        row_idx = 1
        for comp in COMPETENCES_CCF:
            # Récupérer les références de questions pour cette compétence
            refs = questions_par_comp.get(comp["nom"], [])
            refs_str = "  ".join(refs) if refs else ""

            for k, indicateur in enumerate(comp["indicateurs"]):
                fill_cell(grid.cell(row_idx, 0),
                          comp["nom"] if k == 0 else "",
                          bold=(k == 0), size=9)
                fill_cell(grid.cell(row_idx, 1), "", size=9)
                fill_cell(grid.cell(row_idx, 2), indicateur, size=9)
                # Numéros de questions sur la 1ère ligne de chaque compétence uniquement
                fill_cell(grid.cell(row_idx, 3),
                          refs_str if k == 0 else "",
                          bold=(k == 0 and bool(refs_str)), size=9)
                fill_cell(grid.cell(row_idx, 4), "0    1    2", size=9)
                row_idx += 1

        # Ligne note finale
        fill_cell(grid.cell(row_idx, 0), "", size=9)
        fill_cell(grid.cell(row_idx, 1), "", size=9)
        fill_cell(grid.cell(row_idx, 2), "", size=9)
        fill_cell(grid.cell(row_idx, 3), "Note :", bold=True, size=10)
        fill_cell(grid.cell(row_idx, 4), "          / 10", bold=True, size=11)

        doc.add_paragraph()
        obs_p = doc.add_paragraph()
        set_run(obs_p, "Observations : ___________________________________________________________________", size=10)

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()

    except Exception as e:
        st.error(f"Erreur lors de la génération du Word officiel : {e}")
        return None


# ============================================================
# 7. INTERFACE STREAMLIT
# ============================================================

st.set_page_config(
    page_title="Assistant Professeur IA",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
  .stButton>button { border-radius: 8px; font-weight: 600; }
  .info-box  { background:#f0f4ff; border-left:4px solid #4a6cf7; padding:12px 16px; border-radius:4px; margin:8px 0; font-size:.9rem; }
  .warn-box  { background:#fff8e1; border-left:4px solid #f59e0b; padding:12px 16px; border-radius:4px; margin:8px 0; font-size:.9rem; }
  .ok-box    { background:#f0fdf4; border-left:4px solid #22c55e; padding:12px 16px; border-radius:4px; margin:8px 0; font-size:.9rem; }
  .badge     { display:inline-block; background:#e0e7ff; color:#3730a3; border-radius:6px;
               padding:3px 10px; font-size:.8rem; margin-right:6px; margin-bottom:4px; }
</style>
""", unsafe_allow_html=True)

# ── SESSION STATE ─────────────────────────────────────────────
for key in ["generated_md", "generated_ccf_md", "correction_md", "meta_gen", "meta_ccf", "generated_coint_md"]:
    if key not in st.session_state:
        st.session_state[key] = None

# ── SIDEBAR ──────────────────────────────────────────────────
with st.sidebar:
    st.header("⚙️ Configuration")

    st.subheader("🔑 Clé API Gemini")
    cle_api = st.text_input("Clé Google Gemini", type="password", key="gemini_key")
    if cle_api:
        st.markdown('<div class="ok-box">✅ Clé API renseignée</div>', unsafe_allow_html=True)
    else:
        st.markdown('<div class="warn-box">⚠️ Clé API manquante</div>', unsafe_allow_html=True)

    with st.expander("📋 Obtenir une clé gratuite"):
        st.markdown("""
**3 étapes :**

**1.** → [aistudio.google.com](https://aistudio.google.com/app/apikey)

**2.** Connectez-vous → **"Create API Key"**

**3.** Copiez et collez ici

---
💡 **Quota gratuit :** 1 500 req/jour, 15/min.

🔧 **Erreur 429 ?** Attendez 1 min ou activez la facturation sur votre projet Google Cloud (gratuit).
        """)

    st.divider()
    st.subheader("🏫 Établissement")
    nom_etablissement = st.text_input(
        "Nom de l'établissement",
        value="Lino Ventura (Ozoir-la-Ferrière)",
        help="Apparaît dans l'en-tête des CCF officiels"
    )
    annee_scolaire = st.text_input("Année scolaire", value="2025/2026")

    st.divider()
    st.caption("Version 6.0 · Gemini 2.5 Flash")

# ── TITRE ────────────────────────────────────────────────────
st.title("🎓 Assistant Professeur IA")
st.caption("Génération de sujets · CCF Bac Pro · Correction de copies · Export Word")

# ── ONGLETS ──────────────────────────────────────────────────
tab_gen, tab_ccf, tab_coint, tab_correction, tab_export = st.tabs([
    "📝 Générateur de Sujets",
    "🎯 Sujets CCF",
    "🔗 Co-intervention",
    "📸 Correction de Copies",
    "📊 Export Pronote"
])


# ─────────────────────────────────────────────────────────────
# ONGLET 1 — GÉNÉRATEUR
# ─────────────────────────────────────────────────────────────
with tab_gen:
    st.subheader("📝 Générateur de sujets et exercices")
    col1, col2 = st.columns(2)

    with col1:
        cat = st.selectbox("Type d'établissement", list(NIVEAUX_CATEGORIES.keys()), key="gen_cat")
        niv = st.selectbox("Classe", NIVEAUX_CATEGORIES[cat], key="gen_niv")
        filiere = ""
        if cat in ["Bac Pro", "CAP"]:
            fil = st.selectbox("Filière", LISTE_FILIERES, key="gen_fil")
            if fil == "Autre (Préciser ci-dessous)":
                filiere = st.text_input("Filière libre", key="gen_spec")
            else:
                filiere = fil
                if fil in CONTEXTES_FILIERES:
                    st.info(f"📌 {CONTEXTES_FILIERES[fil]['nom_complet']}")

    with col2:
        mat = st.selectbox("Matière", MATIERES, key="gen_mat")
        chap = st.selectbox("Chapitre (BO)", get_chapitres(mat, niv, cat), key="gen_chap")

    consignes = st.text_area("Consignes particulières (optionnel)", height=80, key="gen_consignes",
                              placeholder="Ex : 3 exercices, niveau accessible, inclure un graphique…")

    # ── Niveau de difficulté ──────────────────────────────────
    st.markdown("**🎯 Niveau de difficulté (pédagogie différenciée)**")
    difficulte = st.select_slider(
        "Niveau",
        options=list(NIVEAUX_DIFFICULTE.keys()),
        value="🟡 Moyen",
        key="gen_diff",
        label_visibility="collapsed"
    )
    st.markdown(
        f'<div class="info-box">{difficulte} — {NIVEAUX_DIFFICULTE[difficulte]}</div>',
        unsafe_allow_html=True
    )

    if st.button("✨ Générer le sujet", type="primary", use_container_width=True):
        if not cle_api:
            st.error("🔑 Renseigne ta clé API dans le panneau gauche !")
        else:
            with st.spinner("⏳ Génération en cours…"):
                try:
                    res = call_gemini(cle_api, build_prompt_exercices(niv, cat, mat, chap, consignes, filiere, difficulte))
                    st.session_state.generated_md = res
                    st.session_state.meta_gen = {"niveau": niv, "matiere": mat, "chapitre": chap, "filiere": filiere, "difficulte": difficulte}
                    st.success("✅ Sujet généré !")
                except Exception as e:
                    if "429" in str(e):
                        st.error("⏱️ Quota dépassé (429). Attendez 1 minute et réessayez.")
                    else:
                        st.error(f"Erreur API : {e}")

    if st.session_state.generated_md:
        m = st.session_state.meta_gen or {}
        st.divider()
        badges = (
            f'<span class="badge">📚 {m.get("matiere","")}</span>'
            f'<span class="badge">🎓 {m.get("niveau","")}</span>'
            f'<span class="badge">📖 {m.get("chapitre","")}</span>'
        )
        if m.get("filiere"):
            badges += f'<span class="badge">🏢 {m["filiere"]}</span>'
        if m.get("difficulte"):
            badges += f'<span class="badge">{m["difficulte"]}</span>'
        st.markdown(badges, unsafe_allow_html=True)
        st.markdown(st.session_state.generated_md)

        titre_doc = f"Sujet — {m.get('matiere','')} {m.get('niveau','')} {m.get('difficulte','')}"
        c1, c2, c3 = st.columns(3)
        with c1:
            st.download_button("📥 .md", st.session_state.generated_md,
                               file_name="sujet.md", mime="text/markdown", key="dl1_md")
        with c2:
            st.download_button("📄 .txt", st.session_state.generated_md,
                               file_name="sujet.txt", mime="text/plain", key="dl1_txt")
        with c3:
            st.download_button("📝 .docx",
                markdown_to_docx(st.session_state.generated_md, titre_doc),
                file_name="sujet.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl1_docx")


# ─────────────────────────────────────────────────────────────
# ONGLET 2 — CCF
# ─────────────────────────────────────────────────────────────
with tab_ccf:
    st.subheader("🎯 Générateur de Sujets CCF — Bac Pro")

    mode = st.radio(
        "Mode de génération",
        ["📋 Entraînement (sans en-tête officiel)",
         "📄 Examen officiel (avec en-tête et barème imprimable)"],
        key="ccf_mode"
    )
    is_officiel = "officiel" in mode

    if is_officiel:
        st.markdown('<div class="info-box">📄 <strong>Mode Examen officiel</strong> — En-tête académique complet, cadre candidat, grille de compétences et fiche d\'évaluation. Prêt à imprimer.</div>', unsafe_allow_html=True)
    else:
        st.markdown('<div class="info-box">📋 <strong>Mode Entraînement</strong> — Structure CCF sans en-tête officiel. Idéal pour entraîner les élèves avant l\'examen.</div>', unsafe_allow_html=True)

    col1, col2 = st.columns(2)
    with col1:
        ccf_cat = st.selectbox("Catégorie", ["Bac Pro", "CAP"], key="ccf_cat")
        # Les CCF n'ont lieu qu'en 1ère et Terminale — 2nde Pro exclue
        niveaux_ccf = [n for n in NIVEAUX_CATEGORIES[ccf_cat] if n != "2nde Pro"]
        ccf_niv = st.selectbox("Classe", niveaux_ccf, key="ccf_niv")
        ccf_fil = st.selectbox("Filière", LISTE_FILIERES, key="ccf_fil")
        ccf_filiere = ""
        if ccf_fil == "Autre (Préciser ci-dessous)":
            ccf_filiere = st.text_input("Filière libre", key="ccf_spec")
        else:
            ccf_filiere = ccf_fil
            if ccf_fil in CONTEXTES_FILIERES:
                st.info(f"📌 {CONTEXTES_FILIERES[ccf_fil]['nom_complet']}")

    with col2:
        ccf_mat = st.selectbox("Matière", MATIERES, key="ccf_mat")
        ccf_duree = st.selectbox("Durée de l'épreuve", ["30 min", "45 min", "1h00", "1h30", "2h00"], index=1, key="ccf_duree")
        num_sit = st.number_input("N° de situation CCF", min_value=1, max_value=3, value=1, step=1, key="ccf_num_sit")

    # ── Chapitres ────────────────────────────────────────────
    st.markdown("**📖 Chapitres évalués**")
    chapitres_dispo = get_chapitres(ccf_mat, ccf_niv, ccf_cat)
    col_a, col_b = st.columns(2)
    with col_a:
        ccf_chap = st.selectbox("Chapitre — Partie A", chapitres_dispo, key="ccf_chap")
    with col_b:
        double_chap = st.checkbox("Ajouter un 2ème chapitre (Partie B)", key="ccf_double_chap")
        if double_chap:
            chapitres_b = [c for c in chapitres_dispo if c != ccf_chap]
            ccf_chap_b = st.selectbox("Chapitre — Partie B", chapitres_b, key="ccf_chap_b")
        else:
            ccf_chap_b = ""

    if double_chap and ccf_chap_b:
        st.markdown(
            f'<div class="info-box">📐 <strong>Partie A :</strong> {ccf_chap} &nbsp;|&nbsp; '
            f'<strong>Partie B :</strong> {ccf_chap_b}<br>'
            f'Les deux parties s\'appuieront sur la même situation professionnelle.</div>',
            unsafe_allow_html=True
        )

    avec_corrige = st.checkbox(
        "📝 Inclure le corrigé détaillé *(document professeur)*",
        value=True,
        key="ccf_avec_corrige",
        help="Décochez pour générer uniquement le sujet élève — réponse plus courte, moins de tokens consommés."
    )
    if not avec_corrige:
        st.markdown('<div class="info-box">💡 Sujet seul — génération plus rapide et moins gourmande en quota.</div>', unsafe_allow_html=True)

    ccf_consignes = st.text_area("Consignes spécifiques (optionnel)", height=80, key="ccf_consignes",
                                  placeholder="Ex : Situation en EHPAD, tableau de données, niveau accessible…")

    btn_lbl = "📄 Générer le Sujet CCF Officiel" if is_officiel else "📋 Générer le Sujet d'Entraînement CCF"
    if st.button(btn_lbl, type="primary", use_container_width=True):
        if not cle_api:
            st.error("🔑 Renseigne ta clé API dans le panneau gauche !")
        else:
            with st.spinner("⏳ Génération du sujet CCF…"):
                try:
                    if is_officiel:
                        prompt = build_prompt_ccf_officiel(
                            ccf_niv, ccf_cat, ccf_mat, ccf_chap,
                            ccf_consignes, ccf_filiere, ccf_duree, str(num_sit),
                            avec_corrige=avec_corrige,
                            chapitre_b=ccf_chap_b
                        )
                    else:
                        prompt = build_prompt_ccf_entrainement(
                            ccf_niv, ccf_cat, ccf_mat, ccf_chap,
                            ccf_consignes, ccf_filiere,
                            avec_corrige=avec_corrige,
                            chapitre_b=ccf_chap_b
                        )
                    res = call_gemini(cle_api, prompt)
                    if res:
                        st.session_state.generated_ccf_md = res
                        chap_label = ccf_chap
                        if ccf_chap_b:
                            chap_label += f" + {ccf_chap_b}"
                        st.session_state.meta_ccf = {
                            "niveau": ccf_niv,
                            "matiere": ccf_mat,
                            "chapitre": chap_label,
                            "filiere": ccf_filiere,
                            "mode": "Officiel" if is_officiel else "Entraînement",
                            "num_situation": str(num_sit),
                            "duree": ccf_duree,
                            "annee_scolaire": annee_scolaire,
                        }
                        st.success("✅ Sujet CCF généré !")
                    else:
                        st.error("❌ L'IA n'a pas renvoyé de texte. Réessayez.")
                except Exception as e:
                    if "429" in str(e):
                        st.error("⏱️ Quota dépassé (429). Attendez 1 minute et réessayez.")
                    else:
                        st.error(f"Erreur API : {e}")

    if st.session_state.generated_ccf_md:
        m = st.session_state.meta_ccf or {}
        st.divider()
        badges = (
            f'<span class="badge">🎯 CCF {m.get("mode","")}</span>'
            f'<span class="badge">📚 {m.get("matiere","")}</span>'
            f'<span class="badge">🎓 {m.get("niveau","")}</span>'
            f'<span class="badge">📖 {m.get("chapitre","")}</span>'
        )
        if m.get("filiere"):
            badges += f'<span class="badge">🏢 {m["filiere"]}</span>'
        if m.get("duree"):
            badges += f'<span class="badge">⏱ {m["duree"]}</span>'
        st.markdown(badges, unsafe_allow_html=True)
        st.markdown(st.session_state.generated_ccf_md)

        titre_doc = f"CCF_{m.get('mode','')}_{m.get('matiere','')}_{m.get('niveau','')}"
        st.subheader("📥 Télécharger")

        if is_officiel:
            c1, c2 = st.columns(2)
            with c1:
                docx_off = generate_ccf_officiel_docx(
                    st.session_state.generated_ccf_md, m, nom_etablissement
                )
                if docx_off:
                    st.download_button(
                        "🏛️ Word Officiel (avec en-tête et fiche)",
                        docx_off,
                        file_name=f"{titre_doc}_officiel.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary",
                        key="dl_ccf_officiel"
                    )
            with c2:
                st.download_button("📥 .md (brut)", st.session_state.generated_ccf_md,
                                   file_name=f"{titre_doc}.md", mime="text/markdown", key="dl2_md_off")
        else:
            c1, c2, c3 = st.columns(3)
            with c1:
                st.download_button("📝 .docx",
                    markdown_to_docx(st.session_state.generated_ccf_md, titre_doc),
                    file_name=f"{titre_doc}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="dl2_docx")
            with c2:
                st.download_button("📥 .md", st.session_state.generated_ccf_md,
                                   file_name=f"{titre_doc}.md", mime="text/markdown", key="dl2_md")
            with c3:
                st.download_button("📄 .txt", st.session_state.generated_ccf_md,
                                   file_name=f"{titre_doc}.txt", mime="text/plain", key="dl2_txt")


# ─────────────────────────────────────────────────────────────
# ONGLET 3 — CO-INTERVENTION
# ─────────────────────────────────────────────────────────────
with tab_coint:
    st.subheader("🔗 Co-intervention — Maths & Matières Professionnelles")
    st.markdown(
        '<div class="info-box">🤝 <strong>Co-intervention</strong> — Exercices qui relient '
        'les mathématiques aux situations professionnelles réelles du référentiel. '
        'Chaque exercice part d\'un document pro concret (facture, bulletin, bon de commande…).</div>',
        unsafe_allow_html=True
    )

    col1, col2 = st.columns(2)
    with col1:
        coint_fil = st.selectbox(
            "Filière",
            list(THEMES_COINTERVENTION.keys()),
            key="coint_fil",
            format_func=lambda k: f"{k} — {THEMES_COINTERVENTION[k]['nom_complet']}"
        )
        coint_niv = st.selectbox(
            "Classe",
            ["2nde Pro", "1ère Pro", "Term Pro"],
            key="coint_niv"
        )
    with col2:
        themes_dispo = list(THEMES_COINTERVENTION[coint_fil]["themes"].keys())
        coint_theme = st.selectbox("Thème de co-intervention", themes_dispo, key="coint_theme")

        # Afficher les infos du thème sélectionné
        data_th = THEMES_COINTERVENTION[coint_fil]["themes"][coint_theme]
        st.markdown(
            f'<div class="info-box" style="font-size:.82rem">'
            f'📐 <strong>Formules :</strong> {data_th["formules"]}<br>'
            f'📄 <strong>Documents :</strong> {data_th["documents"]}</div>',
            unsafe_allow_html=True
        )

    st.markdown("**🎯 Niveau de difficulté**")
    coint_diff = st.select_slider(
        "Niveau co-intervention",
        options=list(NIVEAUX_DIFFICULTE.keys()),
        value="🟡 Moyen",
        key="coint_diff",
        label_visibility="collapsed"
    )
    st.markdown(
        f'<div class="info-box">{coint_diff} — {NIVEAUX_DIFFICULTE[coint_diff]}</div>',
        unsafe_allow_html=True
    )

    coint_consignes = st.text_area(
        "Consignes particulières (optionnel)", height=80, key="coint_consignes",
        placeholder="Ex : Utiliser une facture de fournitures de bureau, insister sur le calcul de TVA à 20%…"
    )

    if st.button("🔗 Générer l'exercice de co-intervention", type="primary", use_container_width=True):
        if not cle_api:
            st.error("🔑 Renseigne ta clé API dans le panneau gauche !")
        else:
            with st.spinner("⏳ Génération de l'exercice co-intervention…"):
                try:
                    sys_instr, user_prompt = build_prompt_cointervention(
                        coint_niv, coint_fil, coint_theme, coint_consignes, coint_diff
                    )
                    res = call_gemini(cle_api, (sys_instr, user_prompt))
                    st.session_state.generated_coint_md = res
                    st.success("✅ Exercice co-intervention généré !")
                except Exception as e:
                    if "429" in str(e):
                        st.error("⏱️ Quota dépassé (429). Attendez 1 minute et réessayez.")
                    else:
                        st.error(f"Erreur API : {e}")

    if st.session_state.generated_coint_md:
        nom_filiere = THEMES_COINTERVENTION[coint_fil]["nom_complet"]
        st.divider()
        badges = (
            f'<span class="badge">🔗 Co-intervention</span>'
            f'<span class="badge">🏢 {coint_fil}</span>'
            f'<span class="badge">🎓 {coint_niv}</span>'
            f'<span class="badge">📐 {coint_theme[:40]}</span>'
            f'<span class="badge">{coint_diff}</span>'
        )
        st.markdown(badges, unsafe_allow_html=True)
        st.markdown(st.session_state.generated_coint_md)

        titre_doc = f"CoIntervention_{coint_fil}_{coint_niv}_{coint_theme[:30]}"
        c1, c2, c3 = st.columns(3)
        with c1:
            st.download_button("📥 .md", st.session_state.generated_coint_md,
                               file_name=f"{titre_doc}.md", mime="text/markdown", key="dl_coint_md")
        with c2:
            st.download_button("📄 .txt", st.session_state.generated_coint_md,
                               file_name=f"{titre_doc}.txt", mime="text/plain", key="dl_coint_txt")
        with c3:
            st.download_button("📝 .docx",
                markdown_to_docx(st.session_state.generated_coint_md,
                                 f"Co-intervention — {coint_fil} · {coint_theme}"),
                file_name=f"{titre_doc}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_coint_docx")


# ─────────────────────────────────────────────────────────────
# ONGLET 4 — CORRECTION
# ─────────────────────────────────────────────────────────────
with tab_correction:
    st.subheader("📸 Correction IA de Copies par Photo")

    col1, col2 = st.columns(2)
    with col1:
        corr_cat = st.selectbox("Catégorie", list(NIVEAUX_CATEGORIES.keys()), key="corr_cat")
        corr_niv = st.selectbox("Classe", NIVEAUX_CATEGORIES[corr_cat], key="corr_niv")
    with col2:
        corr_mat = st.selectbox("Matière", MATIERES, key="corr_mat")
        note_sur = st.number_input("Note sur :", min_value=10, max_value=100, value=20, step=5)

    img_file = st.file_uploader("📤 Téléverser la photo de la copie", type=["jpg", "jpeg", "png"])

    col1, col2 = st.columns(2)
    with col1:
        bareme = st.text_area("Barème (optionnel)", height=100,
                               placeholder="Ex : Q1 : 3 pts, Q2 : 5 pts, Q3 : 7 pts, Q4 : 5 pts")
    with col2:
        ton = st.select_slider("Ton", options=["Très bienveillant", "Bienveillant", "Encourageant", "Factuel", "Exigeant"])

    if img_file:
        col_img, col_btn = st.columns([1, 2])
        with col_img:
            st.image(img_file, caption="Copie à corriger", width=250)
        with col_btn:
            if st.button("🔍 Lancer la correction IA", type="primary", use_container_width=True):
                if not cle_api:
                    st.error("🔑 Clé API manquante !")
                else:
                    with st.spinner("🔬 Analyse de la copie…"):
                        try:
                            prompt = build_prompt_correction(bareme, ton, corr_niv, corr_mat, note_sur)
                            res = call_gemini(cle_api, prompt, image=img_file)
                            st.session_state.correction_md = res
                            st.success("✅ Correction terminée !")
                        except Exception as e:
                            if "429" in str(e):
                                st.error("⏱️ Quota dépassé (429). Attendez 1 minute.")
                            else:
                                st.error(f"Erreur : {e}")

    if st.session_state.correction_md:
        st.divider()
        st.subheader("📝 Rapport de Correction")
        st.markdown(st.session_state.correction_md)
        c1, c2 = st.columns(2)
        with c1:
            st.download_button("📥 .md", st.session_state.correction_md,
                               file_name="correction.md", mime="text/markdown", key="dl3_md")
        with c2:
            st.download_button("📝 .docx",
                markdown_to_docx(st.session_state.correction_md, "Rapport de correction"),
                file_name="correction.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl3_docx")


# ─────────────────────────────────────────────────────────────
# ONGLET 4 — EXPORT PRONOTE
# ─────────────────────────────────────────────────────────────
with tab_export:
    st.subheader("📊 Export Pronote — Module en développement")
    st.markdown('<div class="info-box">🚧 <strong>Ce module est en cours de construction.</strong> La prochaine version permettra d\'exporter les notes vers Pronote via CSV compatible.</div>', unsafe_allow_html=True)
    st.markdown("""
### 📋 Fonctionnalités prévues
- Saisie manuelle des notes après correction IA
- Export CSV au format compatible Pronote (import direct)
- Calcul automatique des moyennes de classe
- Appréciation groupée générée par IA pour chaque élève

### 💡 En attendant
Copiez les notes depuis l'onglet **Correction de Copies** et saisissez-les manuellement dans Pronote.
    """)
