import google.generativeai as genai
import PIL.Image
import streamlit as st
import re
import os
import time
import requests
from datetime import datetime
import json
from io import BytesIO
from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from zoneinfo import ZoneInfo

# Configuration des thèmes
THEMES = {
    "🌌 Nébuleuse": {
        "primary": "#6366f1", "secondary": "#8b5cf6", "bg_app": "#0d0f1a", 
        "bg_side": "#13162a", "text": "#e2e8f0", "accent": "#a5b4fc", "border": "#2d3561"
    },
    "✨ Pastel": {
        "primary": "#ec4899", "secondary": "#f472b6", "bg_app": "#fff1f2", 
        "bg_side": "#ffe4e6", "text": "#881337", "accent": "#be185d", "border": "#fecdd3"
    },
    "🦾 Cyberpunk": {
        "primary": "#00FF41", "secondary": "#008F11", "bg_app": "#000000", 
        "bg_side": "#050505", "text": "#00FF41", "accent": "#00FF41", "border": "#003B00"
    },
    "📄 Examen": {
        "primary": "#2563eb", "secondary": "#1d4ed8", "bg_app": "#ffffff", 
        "bg_side": "#f8fafc", "text": "#1e293b", "accent": "#334155", "border": "#cbd5e1"
    }
}

# 2. Le menu de choix dans la barre latérale
if 'theme_pref' not in st.session_state:
    st.session_state.theme_pref = "🌌 Nébuleuse"

# Le sélecteur utilise maintenant la valeur stockée en mémoire
theme_nom = st.sidebar.selectbox(
    "🎨 Style de l'interface", 
    list(THEMES.keys()), 
    index=list(THEMES.keys()).index(st.session_state.theme_pref)
)

st.session_state.theme_pref = theme_nom
t = THEMES[theme_nom]
st.markdown(f"""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Outfit:wght@400;600;700;800&family=JetBrains+Mono:wght@400;600&display=swap');

    /* ── Base & Fonts ── */
    .stApp {{
        background: {t['bg_app']};
        color: {t['text']};
        font-family: 'Outfit', sans-serif;
    }}
    .stApp > header {{ background: transparent !important; }}

    /* ── Sidebar ── */
    [data-testid="stSidebar"] {{
        background: {t['bg_side']} !important;
        border-right: 1px solid {t['border']};
    }}
    [data-testid="stSidebar"] * {{ color: {t['text']} !important; }}

    /* ── Widgets (Radio & Selectbox) - AJOUTÉ POUR FIXER TON PROBLÈME ── */
    .stWidgetLabel p, div[data-testid="stMarkdownContainer"] p {{ 
        color: {t['text']} !important; 
        font-weight: 600; 
    }}
    
    /* Couleur du point à l'intérieur du bouton radio sélectionné */
    div[data-testid="stRadio"] div[role="radiogroup"] div[aria-checked="true"] > div {{
        background-color: {t['primary']} !important;
        border-color: {t['primary']} !important;
    }}

    /* ── Onglets ── */
    .stTabs [data-baseweb="tab-list"] {{
        background: {t['bg_side']};
        border-radius: 12px;
        padding: 4px;
        gap: 4px;
        border: 1px solid {t['border']};
    }}
    .stTabs [data-baseweb="tab"] {{
        background: transparent;
        color: {t['text']};
        opacity: 0.7;
        border-radius: 8px;
        font-weight: 600;
        font-family: 'Outfit', sans-serif;
    }}
    .stTabs [aria-selected="true"] {{
        background: linear-gradient(135deg, {t['primary']}, {t['secondary']}) !important;
        color: white !important;
        box-shadow: 0 0 16px {t['primary']}80;
    }}

    /* ── Boutons primaires ── */
    .stButton > button[kind="primary"] {{
        background: linear-gradient(135deg, {t['primary']}, {t['secondary']});
        border: none;
        border-radius: 12px;
        font-family: 'Outfit', sans-serif;
        font-weight: 700;
        color: white;
        box-shadow: 0 0 20px {t['primary']}66;
        transition: all .25s;
    }}
    .stButton > button[kind="primary"]:hover {{
        transform: translateY(-2px);
        box-shadow: 0 0 30px {t['primary']}B3;
    }}

    /* ── Selectbox / Inputs ── */
    .stSelectbox > div > div,
    .stTextInput > div > div > input {{
        background: {t['bg_side']} !important;
        border: 1px solid {t['border']} !important;
        color: {t['text']} !important;
        border-radius: 10px !important;
    }}

    /* ── Metric cards ── */
    [data-testid="stMetric"] {{
        background: {t['bg_side']};
        border: 1px solid {t['border']};
        border-radius: 12px;
        padding: 16px;
    }}
    [data-testid="stMetricValue"] {{
        color: {t['accent']} !important;
        font-weight: 800;
    }}

    /* ── XP Banner ── */
    .xp-banner {{
        background: linear-gradient(135deg, {t['bg_side']}, {t['bg_app']});
        border: 1px solid {t['border']};
        border-radius: 14px;
        padding: 14px 20px;
        margin-bottom: 16px;
        display: flex;
        align-items: center;
        gap: 16px;
    }}
    .xp-value {{
        font-family: 'Outfit', sans-serif;
        font-weight: 800;
        font-size: 1.4rem;
        color: {t['primary']};
    }}

    /* ── Boss banner ── */
    .boss-banner {{
        background: linear-gradient(135deg, {t['bg_side']}, {t['secondary']}44);
        border: 2px solid {t['primary']};
        border-radius: 16px;
        padding: 24px;
        text-align: center;
        margin: 16px 0;
        box-shadow: 0 0 40px {t['primary']}33;
    }}

    /* ── Scrollbar ── */
    ::-webkit-scrollbar {{ width: 6px; }}
    ::-webkit-scrollbar-track {{ background: {t['bg_app']}; }}
    ::-webkit-scrollbar-thumb {{ background: {t['border']}; border-radius: 3px; }}
    ::-webkit-scrollbar-thumb:hover {{ background: {t['primary']}; }}

    /* ── Markdown content ── */
    .stMarkdown h1, .stMarkdown h2, .stMarkdown h3 {{ color: {t['text']}; }}
    .stMarkdown p, .stMarkdown li {{ color: {t['text']}; opacity: 0.9; }}
    </style>
""", unsafe_allow_html=True)

# 3. Interface de choix (Genre & Avatar)
col1, col2 = st.columns(2)
with col1:
    genre = st.radio("Comment souhaites-tu que l'on s'adresse à toi ?", 
                    ["Neutre (Aventurier)", "Féminin (Aventurière)", "Masculin (Aventurier)"],
                    key="radio_genre")
with col2:
    avatar_style = st.selectbox("Choisis ton style d'avatar", 
                               ["Robotique 🤖", "Mage 🧙‍♂️", "Guerrier/ère 🛡️", "Animalier 🐾"],
                               key="select_avatar")

# 4. Sauvegarde des préférences d'identité
if 'genre_pref' not in st.session_state:
    st.session_state.genre_pref = "Neutre (Aventurier)"
if 'avatar_pref' not in st.session_state:
    st.session_state.avatar_pref = "Robotique 🤖"

# 5. Mise à jour si l'utilisateur change ses choix
st.session_state.genre_pref = genre
st.session_state.avatar_pref = avatar_style
st.session_state.theme_pref = theme_nom


# ============================================================
# 0. INTÉGRATION GRIST — Suivi des élèves
# ============================================================

def envoyer_grist(code_eleve, type_activite, meta, auto_evaluation=""):
    """Envoie une ligne dans Grist. Silencieux en cas d'erreur."""
    try:
        api_key  = st.secrets.get("GRIST_API_KEY", "")
        doc_id   = st.secrets.get("GRIST_DOC_ID", "")
        base_url = st.secrets.get("GRIST_URL", "https://grist.numerique.gouv.fr")
        if not api_key or not doc_id:
            return
        # ✅ Heure Paris (était UTC avant)
        now = datetime.now(ZoneInfo("Europe/Paris"))
        url = f"{base_url}/api/docs/{doc_id}/tables/Suivi_eleves/records"  # ✅ casse corrigée
        headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
        payload = {"records": [{"fields": {
            "code_eleve":        str(code_eleve),
            "date":              now.strftime("%Y-%m-%d"),
            "heure":             now.strftime("%H:%M"),
            "type_activite":     type_activite,
            "classe":            str(meta.get("niveau", "")),
            "filiere":           str(meta.get("filiere", "")),
            "matiere":           str(meta.get("matiere", "")),
            "chapitre":          str(meta.get("chapitre", "")),
            "niveau_difficulte": str(meta.get("difficulte", "")),
            "auto_evaluation":   str(auto_evaluation),
            "score_auto":  str(meta.get("score_auto", "")),  # ✅ ID corrigé score_auto → score_evaluation
            "source":            str(meta.get("source", "Gemini")),
            "genre":             str(st.session_state.get("genre_pref", "Neutre (Aventurier)")),  # ✅ session_state garanti
            "avatar":            str(st.session_state.get("avatar_pref", "Robotique 🤖")),        # ✅ ID corrigé
        }}]}
        requests.post(url, headers=headers, json=payload, timeout=5)
    except Exception:
        pass


def envoyer_proposition_grist(code_eleve, meta, contenu, auto_evaluation=""):
    """Envoie un sujet généré dans Banque_propositions pour validation prof."""
    try:
        api_key  = st.secrets.get("GRIST_API_KEY", "")
        doc_id   = st.secrets.get("GRIST_DOC_ID", "")
        base_url = st.secrets.get("GRIST_URL", "https://grist.numerique.gouv.fr")
        
        now = datetime.now(ZoneInfo("Europe/Paris"))
        url = f"{base_url}/api/docs/{doc_id}/tables/Banque_propositions/records"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        payload = {
            "records": [{
                "fields": {
                    "code_eleve": str(code_eleve),
                    "date": now.strftime("%Y-%m-%d"),
                    "heure": now.strftime("%H:%M"),
                    "niveau": str(meta.get("niveau", "")),
                    "filiere": str(meta.get("filiere", "")),
                    "matiere": str(meta.get("matiere", "")),
                    "chapitre": str(meta.get("chapitre", "")),
                    "niveau_difficulte": str(meta.get("difficulte", "")),
                    "contenu": str(contenu)[:5000],  # Limite pour éviter les erreurs
                    "auto_evaluation": str(auto_evaluation),
                    "statut": "en attente"
                }
            }]
        }
        
        resp = requests.post(url, headers=headers, json=payload, timeout=8)
        if resp.status_code in (200, 201):
            st.success("✅ Sujet envoyé dans Banque_propositions !")
            return True
        else:
            st.error(f"❌ Erreur envoi {resp.status_code}: {resp.text[:200]}")
            return False
            
    except Exception as e:
        st.error(f"❌ Erreur proposition: {e}")
        return False



def lire_progression_grist(code_eleve: str) -> list:
    """
    Lit toutes les lignes Grist pour un élève donné.
    Retourne une liste de dicts, ou [] en cas d'erreur.
    """
    try:
        api_key  = st.secrets.get("GRIST_API_KEY", "")
        doc_id   = st.secrets.get("GRIST_DOC_ID", "")
        base_url = st.secrets.get("GRIST_URL", "https://grist.numerique.gouv.fr")
        if not api_key or not doc_id or not code_eleve:
            return []
        import urllib.parse
        filtre = urllib.parse.quote(json.dumps({"code_eleve": [code_eleve]}))
        url = f"{base_url}/api/docs/{doc_id}/tables/Suivi_eleves/records?filter={filtre}"
        headers = {"Authorization": f"Bearer {api_key}"}
        resp = requests.get(url, headers=headers, timeout=8)
        if resp.status_code == 200:
            records = resp.json().get("records", [])
            return [r["fields"] for r in records]
    except Exception:
        pass
    return []


def calculer_streak(records: list) -> dict:
    """
    Calcule le streak quotidien d'un élève.
    Retourne :
      - streak       : nb de jours consécutifs jusqu'à aujourd'hui (ou hier)
      - record       : meilleur streak historique
      - nouveau_record : True si le streak actuel bat le record
      - derniere_date  : date ISO de la dernière activité
    """
    from datetime import date, timedelta

    # Extraire les dates uniques avec au moins une bonne éval
    dates_ok = set()
    for r in records:
        if r.get("auto_evaluation", "") in ("😊 Bien", "🌟 Très bien"):
            d = r.get("date", "")
            if d:
                try:
                    dates_ok.add(date.fromisoformat(str(d)[:10]))
                except ValueError:
                    pass

    if not dates_ok:
        return {"streak": 0, "record": 0, "nouveau_record": False, "derniere_date": None}

    today     = date.today()
    yesterday = today - timedelta(days=1)
    dates_ok_sorted = sorted(dates_ok, reverse=True)
    derniere  = dates_ok_sorted[0]

    # Le streak ne compte que si l'élève a travaillé aujourd'hui ou hier
    if derniere < yesterday:
        streak_actuel = 0
    else:
        streak_actuel = 1
        curseur = derniere
        for d in dates_ok_sorted[1:]:
            if curseur - d == timedelta(days=1):
                streak_actuel += 1
                curseur = d
            else:
                break

    # Calcul du record historique (toutes les séquences)
    record = 0
    if dates_ok_sorted:
        seq_cur = 1
        for i in range(1, len(dates_ok_sorted)):
            if dates_ok_sorted[i-1] - dates_ok_sorted[i] == timedelta(days=1):
                seq_cur += 1
            else:
                record = max(record, seq_cur)
                seq_cur = 1
        record = max(record, seq_cur)

    return {
        "streak":         streak_actuel,
        "record":         record,
        "nouveau_record": streak_actuel > 0 and streak_actuel >= record,
        "derniere_date":  derniere,
    }


# ── Banque d'exercices ────────────────────────────────────────


def _slug(text: str) -> str:
    text = text.replace(" ", "_").replace("—", "").replace("/", "_")
    text = re.sub(r"[^a-zA-Z0-9_\-àâéèêëîïôùûüç]", "", text)
    return text[:60]


def _normalise(text: str) -> str:
    """Normalise un texte pour comparaison souple : minuscules, sans accents, sans ponctuation."""
    import unicodedata
    text = unicodedata.normalize("NFD", text)
    text = "".join(c for c in text if unicodedata.category(c) != "Mn")
    text = re.sub(r"[^a-z0-9]", "", text.lower())
    return text


def _trouver_fichier_banque(niveau, filiere, chapitre, difficulte, suffixe="") -> str | None:
    """
    Cherche le fichier le plus proche dans BANQUE_DIR par correspondance souple.
    Ignore les différences d'encodage, d'accents, de ponctuation et de casse.
    suffixe = "" pour classique, "_interactif" pour interactif.
    """
    if not os.path.exists(BANQUE_DIR):
        return None

    niv_n  = _normalise(niveau)
    fil_n  = _normalise(filiere)
    diff_n = _normalise(difficulte.split(" ", 1)[-1])
    # Tronquer le chapitre normalisé à 45 chars — les noms de fichiers sont tronqués par _slug
    chap_n = _normalise(chapitre)[:45]
    ext    = f"{suffixe}.json"

    meilleur = None
    meilleur_score = 0

    for fname in os.listdir(BANQUE_DIR):
        if not fname.endswith(ext):
            continue
        # Pour les fichiers classiques, exclure les fichiers interactifs
        if suffixe == "" and fname.endswith("_interactif.json"):
            continue
        fname_n = _normalise(fname)
        score = sum(1 for k in [niv_n, fil_n, chap_n, diff_n] if k and k in fname_n)
        if score > meilleur_score:
            meilleur_score = score
            meilleur = fname

    if meilleur_score >= 4:
        return os.path.join(BANQUE_DIR, meilleur)
    # Si on ne trouve pas avec 4/4, essayer avec 3/4 (chapitre peut différer légèrement)
    if meilleur_score == 3 and meilleur:
        return os.path.join(BANQUE_DIR, meilleur)
    return None


def charger_banque(niveau, filiere, chapitre, difficulte) -> list:
    """Charge les sujets disponibles pour une combinaison donnée."""
    path = _trouver_fichier_banque(niveau, filiere, chapitre, difficulte, "")
    if not path:
        return []
    try:
        with open(path, encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return []


def charger_banque_interactif(niveau, filiere, chapitre, difficulte) -> list:
    """Charge les sujets interactifs (JSON structuré avec questions) pour une combinaison."""
    path = _trouver_fichier_banque(niveau, filiere, chapitre, difficulte, "_interactif")
    if not path:
        return []
    try:
        with open(path, encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return []


# ============================================================
# GAMIFICATION — Système de boss et de niveaux
# ============================================================

SEUIL_VALIDATION  = 2   # bonnes évals minimum pour valider un niveau
ORDRE_NIVEAUX_DIFF = ["🟢 Débutant", "🟡 Moyen", "🟠 Confirmé", "🔴 Expert"]

# Mascotte de chaque niveau — affiché dans l'onglet progression
MASCOTTES = {
    "🟢 Débutant":  {"animal": "🐣", "nom": "Poussin",  "couleur": "#22c55e"},
    "🟡 Moyen":     {"animal": "🦊", "nom": "Renard",   "couleur": "#f59e0b"},
    "🟠 Confirmé":  {"animal": "🦁", "nom": "Lion",     "couleur": "#f97316"},
    "🔴 Expert":    {"animal": "🐉", "nom": "Dragon",   "couleur": "#ef4444"},
}

# Message d'encouragement affiché quand un boss est vaincu
MESSAGES_VICTOIRE = {
    "🟢 Débutant":  "🐣 Poussin vaincu ! Tu maîtrises les bases — le Renard t'attend !",
    "🟡 Moyen":     "🦊 Renard vaincu ! Tu commences à être redoutable — au Lion !",
    "🟠 Confirmé":  "🦁 Lion vaincu ! Tu es vraiment solide — ose affronter le Dragon !",
    "🔴 Expert":    "🐉 Dragon vaincu ! Tu es un expert — bravo, c'est le sommet !",
}


def calculer_progression(records: list, chapitre: str) -> dict:
    """
    Pour un chapitre donné, retourne pour chaque niveau :
    - nb_bonnes    : nombre de bonnes évaluations (😊 ou 🌟)
    - valide       : True si nb_bonnes >= SEUIL_VALIDATION
    - boss_vaincu  : True si une éval de type "Boss" a été réussie sur ce niveau
    """
    from collections import defaultdict
    bonnes   = defaultdict(int)
    boss_ok  = set()

    for r in records:
        if r.get("chapitre", "") != chapitre:
            continue
        diff   = r.get("niveau_difficulte", "")
        eval_v = r.get("auto_evaluation", "")
        type_a = r.get("type_activite", "")

        if eval_v in ("😊 Bien", "🌟 Très bien"):
            if "Boss" in type_a:
                boss_ok.add(diff)
            else:
                bonnes[diff] += 1

    result = {}
    for diff in ORDRE_NIVEAUX_DIFF:
        nb    = bonnes.get(diff, 0)
        valid = nb >= SEUIL_VALIDATION
        result[diff] = {
            "nb_bonnes":   nb,
            "valide":      valid,
            "boss_vaincu": diff in boss_ok,
        }
    return result


def niveau_suivant(diff: str) -> str | None:
    """Retourne le niveau suivant, ou None si on est au maximum."""
    idx = ORDRE_NIVEAUX_DIFF.index(diff)
    if idx < len(ORDRE_NIVEAUX_DIFF) - 1:
        return ORDRE_NIVEAUX_DIFF[idx + 1]
    return None


def build_prompt_boss(niveau, filiere, chapitre, niveau_valide):
    """
    Génère un CCF-Boss : sujet complet, situation complexe,
    SANS corrigé (l'élève doit se corriger seul ou avec le prof).
    """
    ctx = build_contexte_filiere(filiere)
    diff_label = niveau_valide.split(" ", 1)[-1].upper()
    mascotte   = MASCOTTES[niveau_valide]["animal"]
    return f"""Tu es un professeur expert en Bac Pro qui crée un DÉFI BOSS pour un élève.

L'élève vient de valider le niveau {diff_label} sur le chapitre "{chapitre}".
Il doit maintenant affronter le Boss {mascotte} pour passer au niveau suivant.

Crée un sujet d'exercices BOSS pour :
- Niveau scolaire : {niveau} (Bac Pro)
- Matière : Mathématiques
- Chapitre : {chapitre}
- Difficulté : niveau {diff_label} AVANCÉ — plus ambitieux que d'habitude
{ctx}

RÈGLES DU BOSS :
- Mise en situation professionnelle réaliste et complète, sans données guidées.
- 4 à 5 questions progressives sans aide, sans formules rappelées.
- La dernière question demande un raisonnement complet et une conclusion rédigée.
- PAS DE CORRIGÉ — l'élève doit s'auto-corriger ou demander au professeur.
- Ton encourageant mais exigeant : c'est un défi, pas un exercice ordinaire.

Commence le sujet par :
### ⚔️ DÉFI BOSS — {mascotte} {MASCOTTES[niveau_valide]["nom"]}
*Prouve que tu maîtrises vraiment ce chapitre !*

Réponds entièrement en Markdown."""

# Chemin absolu du dossier contenant app.py — utilisé pour trouver les images
APP_DIR = os.path.dirname(os.path.abspath(__file__))
BANQUE_DIR = os.path.join(APP_DIR, "banque")

# ============================================================
# 1. DONNÉES OFFICIELLES — BO MATHS BAC PRO
# ============================================================

NIVEAUX_CATEGORIES = {
    "Collège": ["6ème", "5ème", "4ème", "3ème"],
    "Lycée Général": ["2nde", "1ère", "Terminale"],
    "Bac Pro": ["2nde Pro", "1ère Pro", "Term Pro"],
    "CAP": ["1ère année CAP", "2ème année CAP"],
    "BTS": ["BTS 1", "BTS 2"]
}

MATIERES = [
    "Mathématiques", "Français", "Histoire-Géographie", "SVT",
    "Physique-Chimie", "Anglais", "Espagnol", "SES",
    "Philosophie", "Sciences de l'ingénieur", "EMC"
]

CHAPITRES_MATHS_BAC_PRO = {
    # ── 2nde Pro ── (programme déjà conforme au BO)
    "2nde Pro": [
        "Calcul numérique et algébrique",
        "Puissances et notations scientifiques",
        "Proportionnalité et pourcentages",
        "Équations du 1er degré à une inconnue",
        "Inéquations du 1er degré",
        "Géométrie plane — figures usuelles",
        "Trigonométrie dans le triangle rectangle",
        "Vecteurs — notions de base",
        "Notion de fonction — représentation graphique",
        "Fonctions linéaires et affines",
        "Statistiques descriptives — série à une variable",
        "Notion de probabilité — expériences aléatoires",
    ],

    # ── 1ère Pro ── conforme au BO Annexe 1 (Mathématiques — Classe de première professionnelle)
    "1ère Pro": [
        # Statistique et probabilités
        "Statistique à deux variables — ajustement affine et coefficient de détermination",
        "Probabilités — événements, tableaux croisés, probabilités conditionnelles",
        # Algèbre – Analyse
        "Suites numériques — suites arithmétiques",
        "Résolution graphique d'équations et d'inéquations f(x)=g(x)",
        "Fonctions polynômes de degré 2 — racines, signe, forme factorisée",
        "Fonction dérivée — variations, extremums, fonction inverse",
        "Calculs commerciaux et financiers — intérêts simples, coûts (filières sans physique-chimie)",
        # Géométrie
        "Géométrie dans l'espace — solides usuels et sections par un plan",
        "Vecteurs du plan — coordonnées, opérations, norme (groupements A et B)",
        "Trigonométrie — cercle trigonométrique, fonctions sinus et cosinus (groupements A et B)",
        # Modules transversaux
        "Algorithmique et programmation Python — listes, fonctions, boucles",
        "Automatismes — calcul, grandeurs, lecture graphique",
    ],

    # ── Term Pro ── conforme au BO Annexe 2 (Mathématiques — Classe terminale professionnelle)
    "Term Pro": [
        # Statistique et probabilités
        "Statistiques à deux variables — ajustements non affines, changements de variable",
        "Probabilités — arbres pondérés, formule des probabilités totales, indépendance",
        # Algèbre – Analyse
        "Suites géométriques — terme général, sens de variation, somme",
        "Fonctions polynômes de degré 3 — dérivée, variations, extremums",
        "Fonctions exponentielles de base q et logarithme décimal",
        "Calculs commerciaux et financiers — intérêts composés, amortissements (filières sans physique-chimie)",
        # Géométrie
        "Vecteurs dans l'espace — coordonnées, norme, colinéarité (groupement B)",
        "Trigonométrie — équations, vecteurs de Fresnel (groupement A)",
        # Modules transversaux
        "Algorithmique et programmation Python — approfondissement listes et fonctions",
        "Automatismes — probabilités, suites, dérivation, vecteurs",
        # Programme complémentaire (poursuite d'études)
        "Calcul intégral — primitives, intégrale, aire (programme complémentaire)",
        "Fonctions logarithme népérien et exponentielle de base e (programme complémentaire)",
        "Nombres complexes — forme algébrique et trigonométrique (programme complémentaire)",
        "Produit scalaire de deux vecteurs du plan (programme complémentaire)",
    ],
}

CHAPITRES_PAR_MATIERE_GENERAL = {
    "Mathématiques": ["Nombres et opérations", "Fractions", "Équations", "Fonctions", "Géométrie", "Statistiques", "Probabilités", "Algorithmique"],
    "Français": ["Grammaire", "Orthographe", "Conjugaison", "Analyse littéraire", "Argumentation", "Expression écrite", "Oral"],
    "Histoire-Géographie": ["Antiquité", "Moyen Âge", "Époque Moderne", "Époque Contemporaine", "Géographie de la France", "Géographie mondiale", "EMC"],
    "SVT": ["Cellule et génétique", "Évolution", "Corps humain et santé", "Écosystèmes", "Géologie"],
    "Physique-Chimie": ["Mécanique", "Électricité", "Optique", "Thermodynamique", "Chimie organique", "Chimie des solutions"],
    "Anglais": ["Compréhension écrite", "Expression écrite", "Compréhension orale", "Expression orale", "Grammaire", "Vocabulaire professionnel"],
    "Espagnol": ["Compréhension écrite", "Expression écrite", "Grammaire", "Civilisation", "Vocabulaire"],
    "SES": ["Économie", "Sociologie", "Science politique", "Mondialisation"],
    "Philosophie": ["La conscience", "Le langage", "La liberté", "La morale", "La politique", "La vérité", "L'art"],
    "Sciences de l'ingénieur": ["Mécanique", "Électronique", "Informatique industrielle", "Matériaux"],
    "EMC": ["Démocratie et citoyenneté", "Droits et libertés", "Laïcité", "Engagement"],
}

# ============================================================
# 2. FILIÈRES PRO
# ============================================================

LISTE_FILIERES = ["ASSP", "MCVB", "MCVA", "AGORA", "Autre (Préciser ci-dessous)"]

CONTEXTES_FILIERES = {
    "ASSP": {
        "nom_complet": "Accompagnement, Soins et Services à la Personne",
        "contextes_maths": [
            "Calcul de doses et de dilutions médicamenteuses",
            "Gestion de plannings d'intervenants à domicile",
            "Statistiques sur données de santé (IMC, fréquence cardiaque)",
            "Calcul de coûts de prise en charge",
            "Lecture et interprétation de graphiques médicaux",
        ],
        "mots_cles": "soins, patients, résidents, personnes âgées, handicap, domicile, EHPAD, pharmacie, hygiène"
    },
    "MCVB": {
        "nom_complet": "Métiers du Commerce et de la Vente — option B (Prospection Clientèle)",
        "contextes_maths": [
            "Calcul de marges, taux de remise et prix de vente",
            "Statistiques de vente et de prospection",
            "Gestion d'un budget commercial",
            "Pourcentages d'évolution du chiffre d'affaires",
            "Représentations graphiques des ventes",
        ],
        "mots_cles": "vente, client, prospection, chiffre d'affaires, commission, remise, catalogue"
    },
    "MCVA": {
        "nom_complet": "Métiers du Commerce et de la Vente — option A (Animation et Gestion de l'Espace Commercial)",
        "contextes_maths": [
            "Calcul de surfaces de vente et d'implantation linéaire",
            "Statistiques de fréquentation et taux de transformation",
            "Gestion des stocks — taux de rotation, coefficient multiplicateur",
            "Calcul de TVA, de prix TTC et HT",
            "Optimisation de l'espace commercial",
        ],
        "mots_cles": "magasin, rayon, linéaire, stock, inventaire, promotion, merchandising, caisse"
    },
    "AGORA": {
        "nom_complet": "Assistance à la Gestion des Organisations et de leurs Activités",
        "contextes_maths": [
            "Calcul de salaires, cotisations et charges sociales",
            "Gestion de budgets d'entreprise",
            "Statistiques sur données RH",
            "Facturation et suivi comptable de base",
            "Calcul d'intérêts pour emprunts professionnels",
        ],
        "mots_cles": "entreprise, comptabilité, salaire, facture, devis, ressources humaines, secrétariat"
    },
}

# Compétences officielles du BO avec leurs indicateurs détaillés
COMPETENCES_CCF = [
    {
        "nom": "S'approprier",
        "indicateurs": [
            "Rechercher, extraire et organiser l'information.",
            "Traduire des informations, des codages.",
        ]
    },
    {
        "nom": "Analyser / Raisonner",
        "indicateurs": [
            "Émettre des conjectures, formuler des hypothèses.",
            "Proposer, choisir une méthode de résolution ou un protocole expérimental.",
            "Élaborer un algorithme.",
        ]
    },
    {
        "nom": "Réaliser",
        "indicateurs": [
            "Mettre en œuvre une méthode de résolution, des algorithmes ou un protocole expérimental en respectant les règles de sécurité.",
            "Utiliser un modèle, représenter, calculer.",
            "Expérimenter, faire une simulation.",
        ]
    },
    {
        "nom": "Valider",
        "indicateurs": [
            "Exploiter et interpréter des résultats ou des observations de façon critique et argumentée.",
            "Contrôler la vraisemblance d'une conjecture, de la valeur d'une mesure.",
            "Valider un modèle ou une hypothèse.",
            "Mener un raisonnement logique et établir une conclusion.",
        ]
    },
    {
        "nom": "Communiquer",
        "indicateurs": [
            "Rendre compte d'un résultat, à l'oral ou à l'écrit en utilisant des outils et un langage approprié.",
            "Expliquer une démarche.",
        ]
    },
]

# ============================================================
# 3. PROMPTS
# ============================================================

def get_chapitres(matiere, niveau, categorie):
    if matiere == "Mathématiques" and categorie == "Bac Pro" and niveau in CHAPITRES_MATHS_BAC_PRO:
        return CHAPITRES_MATHS_BAC_PRO[niveau]
    return CHAPITRES_PAR_MATIERE_GENERAL.get(matiere, ["Chapitre général"])


def build_contexte_filiere(filiere):
    if filiere in CONTEXTES_FILIERES:
        ctx = CONTEXTES_FILIERES[filiere]
        lignes = "\n".join(f"- {c}" for c in ctx["contextes_maths"])
        return f"\n**Filière : {ctx['nom_complet']}**\nUnivers : {ctx['mots_cles']}\nContextes maths :\n{lignes}\n"
    return f"\nFilière : {filiere}\n" if filiere else ""


SYSTEM_EXERCICES = """\
Tu es un professeur expert en pédagogie différenciée pour lycée professionnel (Bac Pro).
Tes élèves ont un niveau en mathématiques fragile : tes énoncés sont toujours clairs, bienveillants et ancrés dans des contextes professionnels concrets.

En fonction du niveau de difficulté demandé, adapte PRÉCISÉMENT la structure suivante :

━━ DÉBUTANT ━━
- Questions très guidées, découpées en micro-étapes (une opération par question).
- Résultats intermédiaires fournis pour permettre de continuer même en cas d'erreur.
- Vocabulaire ultra-simplifié, aucun terme technique sans définition immédiate.
- Rappel de cours détaillé avec exemple résolu pas à pas.
- Exercice d'application : calcul direct, données déjà extraites.
- Mise en situation : contexte simple, une seule inconnue.
- Pas de problème ouvert — remplacer par une question bilan guidée.

━━ MOYEN ━━
- Questions semi-guidées avec quelques repères (formule rappelée, première étape donnée).
- Rappel de cours synthétique avec un exemple.
- Exercice d'application : 3 questions progressives.
- Mise en situation professionnelle simple avec tableau de données.
- Problème ouvert court (1 question de synthèse guidée).

━━ CONFIRMÉ ━━
- Questions autonomes, aucune aide dans l'énoncé.
- Rappel de cours en points clés uniquement (pas d'exemple résolu).
- Exercice d'application : questions progressives avec barème.
- Mise en situation professionnelle réaliste et complète.
- Problème ouvert avec raisonnement attendu.

━━ EXPERT ━━
- Questions ouvertes sans guidage, transfert de compétences vers une situation nouvelle.
- Rappel de cours : absent ou très succinct (2 lignes max).
- Exercice d'application : données brutes à extraire soi-même.
- Mise en situation complexe avec plusieurs informations à croiser.
- Problème ouvert ambitieux — mais toujours réaliste pour un élève de Bac Pro.

Dans tous les cas : JAMAIS de calcul hors programme Bac Pro, JAMAIS de piège inutile.
La correction détaillée doit être adaptée au même niveau (plus ou moins de détails selon le niveau).

Structure de sortie (Markdown) :
1. **Rappel de cours** (adapté au niveau)
2. **Exercice d'application** (adapté au niveau)
3. **Exercice de mise en situation** (contexte professionnel de la filière)
4. **Problème ouvert** (adapté au niveau)
5. **Corrections détaillées** (avec le niveau de détail approprié)\
"""

# Descriptifs courts affichés dans l'UI
NIVEAUX_DIFFICULTE = {
    "🟢 Débutant":  "Questions très guidées, micro-étapes, résultats intermédiaires donnés.",
    "🟡 Moyen":     "Semi-guidé, quelques repères fournis, mise en situation simple.",
    "🟠 Confirmé":  "Autonome, mise en situation réaliste, raisonnement attendu.",
    "🔴 Expert":    "Transfert de compétences, données brutes, problème ouvert ambitieux.",
}


def build_prompt_exercices(niveau, categorie, matiere, chapitre, consignes, filiere="", difficulte="🟡 Moyen"):
    ctx = build_contexte_filiere(filiere)
    diff_label = difficulte.split(" ", 1)[-1].upper()  # ex: "MOYEN"
    user = (
        f"Génère un contenu pédagogique de niveau **{diff_label}** pour :\n"
        f"- Niveau scolaire : {niveau} ({categorie})\n"
        f"- Matière : {matiere}\n"
        f"- Chapitre : {chapitre}\n"
        f"{ctx}"
        f"- Instructions : {consignes or 'Aucune'}\n\n"
        f"Applique scrupuleusement les consignes du niveau {diff_label} définies dans tes instructions."
    )
    return SYSTEM_EXERCICES, user


def build_prompt_ccf_entrainement(niveau, categorie, matiere, chapitre, consignes, filiere="", avec_corrige=True, chapitre_b=""):
    ctx = build_contexte_filiere(filiere)
    bloc_corrige = "\n### CORRIGÉ DÉTAILLÉ *(document professeur)*\nCorrection complète de chaque question.\n" if avec_corrige else ""
    chap_a_label = f"Partie A : {chapitre}"
    chap_b_label = f"Partie B : {chapitre_b}" if chapitre_b else "Partie B : thème mathématique complémentaire au choix"
    return f"""Tu es un professeur de mathématiques expert en Bac Pro et en évaluation CCF conforme au BO.

Génère un SUJET D'ENTRAÎNEMENT AU CCF pour :
- Niveau : {niveau} ({categorie})
- Matière : {matiere}
- {chap_a_label}
- {chap_b_label}
{ctx}
- Instructions : {consignes or 'Aucune'}

## RÈGLES STRICTES DU FORMAT CCF OFFICIEL :

1. La MISE EN SITUATION présente le contexte professionnel avec données chiffrées réalistes.
2. La PROBLÉMATIQUE est UNE SEULE question (avec point d'interrogation ?) encadrée en gras, placée juste après la mise en situation. C'est la seule question du document avec un "?".
3. Les questions sont numérotées et chacune porte le nom de la compétence évaluée en majuscules : S'APPROPRIER / RÉALISER / ANALYSER-RAISONNER / VALIDER / COMMUNIQUER.
4. Chaque question se termine par des lignes de réponse (______).
5. Les questions servent à répondre progressivement à la problématique.
6. La dernière question COMMUNIQUER demande de répondre à la problématique initiale.
7. L'évaluation est sur 10 points (pas 20), répartis en niveaux 0/1/2 par compétence.
8. Les deux parties s'appuient sur LA MÊME mise en situation professionnelle.

## STRUCTURE À RESPECTER :

### MISE EN SITUATION PROFESSIONNELLE
[Description du contexte, de l'entreprise/structure, avec données chiffrées et document support : tableau ou graphique]

### PROBLÉMATIQUE
**[UNE SEULE question centrale se terminant par ?]**

### PARTIE A — [Titre lié à : {chapitre}]

1. **S'APPROPRIER**
- [question]
______

2. **RÉALISER**
- [question]
______

3. **ANALYSER / RAISONNER**
- [question]
______

4. **RÉALISER** 🛎️ APPELER L'EXAMINATEUR
- [question nécessitant outil/calculatrice]
______

5. **VALIDER**
- [question de vérification/interprétation]
______

6. **COMMUNIQUER**
- Répondre à la [première partie de la] problématique.
______

### PARTIE B — [Titre lié à : {chapitre_b if chapitre_b else "second thème mathématique"}]
[Mêmes règles, questions numérotées avec compétences]
{bloc_corrige}
Réponds entièrement en Markdown."""


def build_prompt_ccf_officiel(niveau, categorie, matiere, chapitre, consignes, filiere="", duree="45 min", num_sit="1", avec_corrige=True, chapitre_b=""):
    ctx = build_contexte_filiere(filiere)
    nom_filiere = CONTEXTES_FILIERES[filiere]["nom_complet"] if filiere in CONTEXTES_FILIERES else filiere
    bloc_corrige = "\n### CORRIGÉ DÉTAILLÉ *(document professeur — NE PAS DISTRIBUER)*\n[Correction complète question par question avec justifications]\n" if avec_corrige else ""
    chap_a_label = f"Partie A : {chapitre}"
    chap_b_label = f"Partie B : {chapitre_b}" if chapitre_b else "Partie B : thème mathématique complémentaire en cohérence avec la situation"
    return f"""Tu es un professeur de mathématiques expert en Bac Pro et en évaluation CCF conforme au BO de l'Éducation Nationale.

Génère un SUJET DE CCF OFFICIEL complet et prêt à imprimer pour :
- Niveau : {niveau} ({categorie})
- Filière : {nom_filiere}
- Matière : {matiere}
- {chap_a_label}
- {chap_b_label}
- Situation d'évaluation n° : {num_sit}
- Durée : {duree}
{ctx}
- Instructions complémentaires : {consignes or 'Aucune'}

## RÈGLES STRICTES DU FORMAT CCF OFFICIEL ÉDUCATION NATIONALE :

1. La MISE EN SITUATION présente le contexte professionnel avec une image mentale de la structure/entreprise, des données chiffrées réalistes et un document support (tableau de données ou graphique nommé).
2. La PROBLÉMATIQUE est UNE SEULE question (avec ?) encadrée en gras juste après la mise en situation. C'est la SEULE question du document avec un point d'interrogation.
3. Les questions sont numérotées (1. 2. 3. ...) et chacune porte le nom de la compétence BO en majuscules gras : **S'APPROPRIER** / **RÉALISER** / **ANALYSER / RAISONNER** / **VALIDER** / **COMMUNIQUER**.
4. Toutes les 5 compétences doivent être évaluées au moins une fois.
5. Les questions comportent des lignes de réponse (______).
6. La dernière question **COMMUNIQUER** demande explicitement de répondre à la problématique initiale.
7. Quand une question nécessite la calculatrice ou un outil, ajouter : 🛎️ APPELER L'EXAMINATEUR
8. L'évaluation est notée sur /10 avec niveaux d'acquisition 0/1/2 par compétence (PAS de barème en points).
9. Les deux parties s'appuient sur LA MÊME mise en situation professionnelle.

## STRUCTURE OBLIGATOIRE :

### MISE EN SITUATION PROFESSIONNELLE
[Présentation de la structure professionnelle ({nom_filiere}), contexte détaillé, données chiffrées]
[Document support : tableau ou données nommées]

L'usage de la calculatrice avec mode examen actif est autorisé.

### PROBLÉMATIQUE
**[Question centrale unique se terminant par ? — clairement encadrée]**

### PARTIE A — [Titre lié à : {chapitre}]

1. **S'APPROPRIER**
- [question d'appropriation de la situation]
______

2. **RÉALISER**
- [question de calcul/réalisation]
______

3. **ANALYSER / RAISONNER**
- [question d'analyse]
______

4. **RÉALISER** 🛎️ APPELER L'EXAMINATEUR
- [question avec outil numérique]
______

5. **VALIDER**
- [question de validation/interprétation]
______

6. **COMMUNIQUER**
- Répondre à la première partie de la problématique.
______

### PARTIE B — [Titre lié à : {chapitre_b if chapitre_b else "second thème mathématique"}]
[Mêmes règles]
{bloc_corrige}
Réponds entièrement en Markdown avec mise en page soignée et professionnelle."""


#def build_prompt_correction(bareme, ton, niveau, matiere, note_sur):
#    return f"""Tu es un professeur correcteur expert.

#Contexte :
#- Niveau : {niveau or 'Non précisé'}
#- Matière : {matiere or 'Non précisée'}
# - Barème : {bareme or f'Non fourni — évalue sur {note_sur}'}
#- Ton : {ton}
#- Note sur : {note_sur}

# Mission :
#1. Transcris le texte manuscrit visible.
#2. Identifie et explique chaque erreur avec pédagogie.
#3. Attribue une note partielle par question.
#4. Calcule la note globale /{note_sur}.
#5. Rédige une appréciation finale ({ton}) de 3 à 5 lignes.

#Réponds en Markdown avec sections claires."""/*


# ============================================================
# 4. APPEL API GEMINI
# ============================================================

def call_gemini(api_key, prompt, image=None):
    genai.configure(api_key=api_key.strip())
    # Gestion du tuple (system_instruction, user_prompt) retourné par build_prompt_exercices
    if isinstance(prompt, tuple):
        system_instruction, user_prompt = prompt
        model = genai.GenerativeModel(
            model_name="gemini-2.5-flash",
            system_instruction=system_instruction
        )
        content = user_prompt
    else:
        model = genai.GenerativeModel(model_name="gemini-2.5-flash")
        content = prompt
    if image:
        img = PIL.Image.open(image)
        response = model.generate_content([content, img])
    else:
        response = model.generate_content(content)
    if response and response.text:
        return response.text
    return "L'IA a répondu mais le texte est vide. Réessayez."


# ============================================================
# 5. EXPORT WORD — VERSION STANDARD
# ============================================================

def clean_math(text):
    """
    Convertit la notation LaTeX/Markdown en texte lisible dans Word.
    Gère : \\[...\\], $$...$$, $...$, \\frac, ^, _, commandes LaTeX, backticks.
    """
    import re as _re
    # Supprimer balises ```markdown et ``` parasites
    text = _re.sub(r'```[a-zA-Z]*', '', text)
    # Blocs \[ ... \] et \( ... \)
    text = _re.sub(r'\\\[(.+?)\\\]', lambda m: m.group(1).strip(), text, flags=_re.DOTALL)
    text = _re.sub(r'\\\((.+?)\\\)', lambda m: m.group(1).strip(), text, flags=_re.DOTALL)
    # Blocs $$ ... $$
    text = _re.sub(r'\$\$(.+?)\$\$', lambda m: m.group(1).strip(), text, flags=_re.DOTALL)
    # Inline $ ... $
    text = _re.sub(r'\$([^$\n]+?)\$', lambda m: m.group(1).strip(), text)
    # \frac{a}{b} → a/b
    text = _re.sub(r'\\frac\{([^}]+)\}\{([^}]+)\}', r'\1/\2', text)
    # Exposants ^{...} → ^...
    text = _re.sub(r'\^\{([^}]+)\}', r'^\1', text)
    # Indices _{...} → _...
    text = _re.sub(r'_\{([^}]+)\}', r'_\1', text)
    # Commandes LaTeX usuelles → Unicode
    replacements = {
        r'\times': '×', r'\cdot': '·', r'\leq': '≤', r'\geq': '≥',
        r'\neq': '≠', r'\approx': '≈', r'\infty': '∞', r'\pi': 'π',
        r'\alpha': 'α', r'\beta': 'β', r'\sigma': 'σ', r'\mu': 'μ',
        r'\rightarrow': '→', r'\leftarrow': '←', r'\in': '∈',
        r'\subset': '⊂', r'\cup': '∪', r'\cap': '∩',
    }
    for latex, unicode_char in replacements.items():
        text = text.replace(latex, unicode_char)
    # Supprimer les backslashes restants isolés
    text = _re.sub(r'\\([a-zA-Z]+)', r'\1', text)
    return text


def parse_md_tables(md_text):
    """
    Transforme un texte Markdown en liste de blocs :
    - ('text', 'contenu')     pour les lignes normales
    - ('table', [[row], ...]) pour les tableaux Markdown
    """
    import re as _re
    lines = md_text.split('\n')
    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i]
        # Début d'un tableau : ligne commençant par |
        if line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                raw = lines[i].strip()
                # Ignorer les lignes séparateurs (|---|---|)
                if _re.match(r'^\|[\s\-:|]+\|', raw):
                    i += 1
                    continue
                # Découper les cellules
                cells = [c.strip() for c in raw.strip('|').split('|')]
                table_lines.append(cells)
                i += 1
            if table_lines:
                blocks.append(('table', table_lines))
        else:
            blocks.append(('text', line))
            i += 1
    return blocks


def add_md_table_to_doc(doc, rows, font_size=10):
    """Crée un vrai tableau Word à partir d'une liste de lignes [[cell, cell, ...]]."""
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(max_cols):
            cell = table.cell(i, j)
            text = clean_math(row[j]) if j < len(row) else ""
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(text)
            run.font.size = Pt(font_size)
            run.font.name = "Arial"
            # 1ère ligne en gras (en-tête)
            if i == 0:
                run.bold = True
    doc.add_paragraph()


def render_inline(paragraph, text):
    """Écrit une ligne avec gras/italique et nettoie les $ mathématiques."""
    text = clean_math(text)
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            paragraph.add_run(part[1:-1]).italic = True
        else:
            paragraph.add_run(part)


def markdown_to_docx(md_text, titre="Document"):
    doc = DocxDocument()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)

    h = doc.add_heading(titre, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    blocks = parse_md_tables(md_text)

    for kind, content in blocks:
        if kind == "table":
            add_md_table_to_doc(doc, content, font_size=10)
            continue

        line = content.rstrip()
        if line.startswith("### "):
            doc.add_heading(clean_math(line[4:]), level=3)
        elif line.startswith("## "):
            doc.add_heading(clean_math(line[3:]), level=2)
        elif line.startswith("# "):
            doc.add_heading(clean_math(line[2:]), level=1)
        elif line.startswith("---"):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
        elif re.match(r"^\d+\.", line):
            p = doc.add_paragraph(style="List Number")
            render_inline(p, re.sub(r"^\d+\.\s*", "", line))
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            render_inline(p, line[2:])
        elif line == "":
            doc.add_paragraph("")
        else:
            p = doc.add_paragraph()
            render_inline(p, line)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ============================================================
# 6. EXPORT WORD — VERSION OFFICIELLE CCF
# ============================================================

def set_cell_border(cell):
    """Ajoute une bordure fine sur toutes les faces d'une cellule."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "AAAAAA")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def set_run(paragraph, text, bold=False, size=10):
    """Ajoute un run stylé — évite le crash sur paragraphe vide."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return run


def fill_cell(cell, text, bold=False, size=10):
    """Remplit une cellule proprement sans crasher sur runs vides."""
    p = cell.paragraphs[0]
    p.clear()
    set_run(p, text, bold=bold, size=size)
    set_cell_border(cell)


def parse_questions_competences(content_md):
    """
    Parcourt le Markdown du sujet CCF et construit un dictionnaire
    { nom_compétence_normalisé : [liste de références questions ex: 'A.1', 'B.2'] }
    
    Détecte les patterns :
      ### PARTIE A  →  lettre de partie courante = 'A'
      1. **S'APPROPRIER**  →  question A.1 → compétence S'approprier
    """
    # Mapping des mots-clés détectés vers les noms officiels des compétences
    COMP_KEYWORDS = {
        "S'APPROPRIER":        "S'approprier",
        "APPROPRIER":          "S'approprier",
        "ANALYSER / RAISONNER":"Analyser / Raisonner",
        "ANALYSER":            "Analyser / Raisonner",
        "RAISONNER":           "Analyser / Raisonner",
        "RÉALISER":            "Réaliser",
        "REALISER":            "Réaliser",
        "VALIDER":             "Valider",
        "COMMUNIQUER":         "Communiquer",
    }

    # Initialisation : une liste vide par compétence officielle
    result = {
        "S'approprier":       [],
        "Analyser / Raisonner": [],
        "Réaliser":           [],
        "Valider":            [],
        "Communiquer":        [],
    }

    current_part = ""

    for line in content_md.split("\n"):
        line_s = line.strip()

        # Détecter changement de partie : ### PARTIE A, ## PARTIE B, etc.
        part_match = re.match(r"#+\s+PARTIE\s+([A-Z])", line_s, re.IGNORECASE)
        if part_match:
            current_part = part_match.group(1).upper()
            continue

        # Si pas encore de partie détectée, on ne peut pas construire de référence
        if not current_part:
            continue

        # Détecter une question numérotée avec sa compétence en gras
        # Exemples : "1. **S'APPROPRIER**"  "4. **RÉALISER** 🛎️ APPELER..."
        q_match = re.match(r"^(\d+)\.\s+\*\*([^*]+)\*\*", line_s)
        if q_match:
            q_num   = q_match.group(1)
            comp_raw = q_match.group(2).strip().upper()
            q_ref   = f"{current_part}.{q_num}"

            # Chercher la compétence correspondante
            for keyword, comp_name in COMP_KEYWORDS.items():
                if keyword in comp_raw:
                    if q_ref not in result[comp_name]:
                        result[comp_name].append(q_ref)
                    break

    return result


def generate_ccf_officiel_docx(content_md, metadata, nom_etablissement="Mon Établissement"):
    try:
        doc = DocxDocument()

        # Marges réduites pour plus d'espace
        for section in doc.sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        doc.styles["Normal"].font.name = "Arial"
        doc.styles["Normal"].font.size = Pt(11)

        # ── EN-TÊTE : tableau 3 colonnes avec bordures visibles
        # Conforme au document de référence de l'établissement
        table_h = doc.add_table(rows=3, cols=3)
        table_h.style = "Table Grid"  # bordures visibles sur toutes les cellules

        # Largeurs colonnes en DXA (total 9360 = contenu avec marges 1")
        col_widths = [2620, 4120, 2620]
        for row in table_h.rows:
            for j, cell in enumerate(row.cells):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = OxmlElement("w:tcW")
                tcW.set(qn("w:w"), str(col_widths[j]))
                tcW.set(qn("w:type"), "dxa")
                tcPr.append(tcW)

        # ── Ligne 1 : logo gauche | titre CCF | logo matière ──
        cell_l = table_h.cell(0, 0)
        logo_rep = os.path.join(APP_DIR, "logo_republique.png")
        if os.path.exists(logo_rep):
            cell_l.paragraphs[0].add_run().add_picture(logo_rep, width=Inches(1.2))
        else:
            r_fallback = cell_l.paragraphs[0].add_run("ACADÉMIE DE CRÉTEIL")
            r_fallback.bold = True
            r_fallback.font.size = Pt(10)

        cell_c = table_h.cell(0, 1)
        p_c = cell_c.paragraphs[0]
        p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        def _rb(para, text, size=10):
            r = para.add_run(text)
            r.bold = True
            r.font.size = Pt(size)
            return r
        _rb(p_c, "CONTRÔLE EN COURS DE FORMATION\n", 12)
        _rb(p_c, "Baccalauréat professionnel\n", 11)
        _rb(p_c, f"{metadata.get('matiere', '')}\n", 11)
        _rb(p_c, f"Situation d'évaluation n°{metadata.get('num_situation', '...')}\n", 10)
        _rb(p_c, f"Intitulé du diplôme : {'.' * 31}\n", 10)
        _rb(p_c, f"Durée : {metadata.get('duree', '45 min')}", 10)

        cell_r = table_h.cell(0, 2)
        p_r = cell_r.paragraphs[0]
        p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        logo_mat = os.path.join(APP_DIR, "logo_matiere.png")
        if os.path.exists(logo_mat):
            p_r.add_run().add_picture(logo_mat, width=Inches(0.8))

        # ── Ligne 2 : vide | Nom/Prénom candidat (centré) | vide ──
        # Conforme au document de référence : colonne centre uniquement
        cell_nom = table_h.cell(1, 1)
        p_nom = cell_nom.paragraphs[0]
        p_nom.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _rb(p_nom, "Nom, Prénom du candidat : _______________________________________________", 10)

        # ── Ligne 3 : Date | Établissement | vide ──
        cell_date = table_h.cell(2, 0)
        cell_date.paragraphs[0].add_run("Date : .......................").font.size = Pt(10)

        cell_etab = table_h.cell(2, 1)
        p_etab = cell_etab.paragraphs[0]
        _rb(p_etab, f"Nom de l'établissement : {nom_etablissement}", 10)

        doc.add_paragraph()

        # ── NOTICE CALCULATRICE ────────────────────────────────
        logo_appel_path = os.path.join(APP_DIR, "logo_appel.png")
        p_notice = doc.add_paragraph()
        if os.path.exists(logo_appel_path):
            p_notice.add_run().add_picture(logo_appel_path, width=Inches(0.53))
            p_notice.add_run("  ")
        p_notice.add_run("Dans la suite du document, ce symbole signifie « Appeler l'examinateur ».").font.size = Pt(10)

        for ligne_calc in [
            "L'usage de la calculatrice avec mode examen actif est autorisé.",
            "L'usage de la calculatrice sans mémoire, « type collège » est autorisé.",
            "L'échange de calculatrices entre les candidats pendant l'épreuve est interdit.",
        ]:
            doc.add_paragraph().add_run(ligne_calc).font.size = Pt(10)

        doc.add_paragraph()

        # ── CONTENU DU SUJET ──────────────────────────────────
        # Filtre les lignes parasites que l'IA peut générer avant le vrai sujet
        skip_keywords = [
            "Établissement :", "Baccalauréat Professionnel —", "Épreuve E3",
            "Calculatrice autorisée", "Ministère de l'Éducation",
            "sujet de CCF officiel", "Voici le sujet", "stricte conformité",
            "accompagné de", "fiche d'évaluation et de son corrigé",
            "Filière :", "Niveau :", "Situation d'évaluation n°",
            "Durée :", "Nom & Prénom",
        ]

        # Filtrer le contenu avant de le parser
        filtered_lines = []
        contenu_commence = False
        for line in content_md.split("\n"):
            ls = line.rstrip()
            if not contenu_commence:
                if any(ls.startswith(m) for m in [
                    "### MISE EN SITUATION", "### PARTIE", "### PROBLÉMATIQUE",
                    "## MISE EN SITUATION", "## PARTIE", "## PROBLÉMATIQUE",
                    "#### MISE", "#### PARTIE",
                ]):
                    contenu_commence = True
                else:
                    continue
            if any(kw in ls for kw in skip_keywords):
                continue
            filtered_lines.append(ls)

        filtered_md = "\n".join(filtered_lines)

        # Parser les blocs (texte et tableaux Markdown)
        blocks = parse_md_tables(filtered_md)

        for kind, content in blocks:
            # ── Vrai tableau Word ──────────────────────────────
            if kind == "table":
                add_md_table_to_doc(doc, content, font_size=10)
                continue

            line = content

            if line.startswith("#### "):
                doc.add_heading(clean_math(line[5:].strip()), level=3)
            elif line.startswith("### "):
                doc.add_heading(clean_math(line[4:].strip()), level=2)
            elif line.startswith("## "):
                doc.add_heading(clean_math(line[3:].strip()), level=2)
            elif "APPELER L'EXAMINATEUR" in line.upper() or "🛎️" in line:
                p = doc.add_paragraph()
                logo_appel = os.path.join(APP_DIR, "logo_appel.png")
                if os.path.exists(logo_appel):
                    p.add_run().add_picture(logo_appel, width=Inches(0.25))
                    p.add_run("  ")
                run = p.add_run("APPELER L'EXAMINATEUR")
                run.bold = True
                run.font.size = Pt(11)
                pPr = p._p.get_or_add_pPr()
                pBdr = OxmlElement("w:pBdr")
                for side in ("top", "left", "bottom", "right"):
                    border = OxmlElement(f"w:{side}")
                    border.set(qn("w:val"), "single")
                    border.set(qn("w:sz"), "6")
                    border.set(qn("w:space"), "4")
                    border.set(qn("w:color"), "4A6CF7")
                    pBdr.append(border)
                pPr.append(pBdr)
            elif line.startswith("---"):
                continue
            elif line == "":
                doc.add_paragraph("")
            else:
                p = doc.add_paragraph()
                render_inline(p, line)

        # ── FICHE D'ÉVALUATION OFFICIELLE (page 2) ───────────
        doc.add_page_break()

        titre_fiche = doc.add_heading("FICHE INDIVIDUELLE D'ÉVALUATION", level=1)
        titre_fiche.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # En-tête fiche
        table_ef = doc.add_table(rows=2, cols=2)
        fill_cell(table_ef.cell(0, 0),
                  f"Session : {metadata.get('annee_scolaire','2025/2026')} — Établissement : {nom_etablissement} — Académie : Créteil",
                  bold=True, size=9)
        fill_cell(table_ef.cell(0, 1),
                  f"Spécialité : {metadata.get('matiere','')} — Évaluateur : _______________ — Date : _______________",
                  size=9)
        fill_cell(table_ef.cell(1, 0),
                  f"Situation n° {metadata.get('num_situation','1')} — {metadata.get('filiere','')}",
                  bold=True, size=9)
        fill_cell(table_ef.cell(1, 1), "", size=9)

        doc.add_paragraph()
        cand_p = doc.add_paragraph()
        set_run(cand_p, "Nom et prénom du candidat : _______________________________________________",
                bold=True, size=10)

        doc.add_paragraph()

        # 1. Capacités et connaissances évaluées
        cap_titre = doc.add_paragraph()
        set_run(cap_titre, "1.  Liste des capacités et connaissances évaluées", bold=True, size=11)

        table_cap = doc.add_table(rows=2, cols=2)
        fill_cell(table_cap.cell(0, 0), "Capacités", bold=True, size=10)
        fill_cell(table_cap.cell(0, 1), "(voir section Fiche d'évaluation du sujet généré)", size=10)
        fill_cell(table_cap.cell(1, 0), "Connaissances", bold=True, size=10)
        fill_cell(table_cap.cell(1, 1), "(voir section Fiche d'évaluation du sujet généré)", size=10)

        doc.add_paragraph()

        # 2. Tableau d'évaluation officiel
        eval_titre = doc.add_paragraph()
        set_run(eval_titre, "2.  Évaluation", bold=True, size=11)

        # Calcul nb lignes : 1 en-tête + 1 ligne par indicateur + 1 ligne note
        total_rows = 1
        for comp in COMPETENCES_CCF:
            total_rows += len(comp["indicateurs"])
        total_rows += 1  # ligne note finale

        grid = doc.add_table(rows=total_rows, cols=5)

        # En-têtes colonnes
        h_texts = ["Compétences", "", "Indicateurs / Capacités", "Questions", "Appréciation\n(0 / 1 / 2)"]
        for j, h in enumerate(h_texts):
            fill_cell(grid.cell(0, j), h, bold=True, size=9)

        # Parsing des numéros de questions depuis le contenu généré
        questions_par_comp = parse_questions_competences(content_md)

        row_idx = 1
        for comp in COMPETENCES_CCF:
            # Récupérer les références de questions pour cette compétence
            refs = questions_par_comp.get(comp["nom"], [])
            refs_str = "  ".join(refs) if refs else ""

            for k, indicateur in enumerate(comp["indicateurs"]):
                fill_cell(grid.cell(row_idx, 0),
                          comp["nom"] if k == 0 else "",
                          bold=(k == 0), size=9)
                fill_cell(grid.cell(row_idx, 1), "", size=9)
                fill_cell(grid.cell(row_idx, 2), indicateur, size=9)
                # Numéros de questions sur la 1ère ligne de chaque compétence uniquement
                fill_cell(grid.cell(row_idx, 3),
                          refs_str if k == 0 else "",
                          bold=(k == 0 and bool(refs_str)), size=9)
                fill_cell(grid.cell(row_idx, 4), "0    1    2", size=9)
                row_idx += 1

        # Ligne note finale
        fill_cell(grid.cell(row_idx, 0), "", size=9)
        fill_cell(grid.cell(row_idx, 1), "", size=9)
        fill_cell(grid.cell(row_idx, 2), "", size=9)
        fill_cell(grid.cell(row_idx, 3), "Note :", bold=True, size=10)
        fill_cell(grid.cell(row_idx, 4), "          / 10", bold=True, size=11)

        doc.add_paragraph()
        obs_p = doc.add_paragraph()
        set_run(obs_p, "Observations : ___________________________________________________________________", size=10)

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()

    except Exception as e:
        st.error(f"Erreur lors de la génération du Word officiel : {e}")
        return None


# ============================================================
# 7. INTERFACE STREAMLIT
# ============================================================

st.set_page_config(
    page_title="Entraînement Bac Pro — Maths",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
  /* Override minimal — le thème principal est dans le premier bloc */
  .stCheckbox > label { color: #94a3b8 !important; }
  .stRadio > label { color: #94a3b8 !important; }
  .stExpander { background: #13162a !important; border: 1px solid #2d3561 !important; border-radius: 10px !important; }
</style>
""", unsafe_allow_html=True)

# ── Système XP ───────────────────────────────────────────────
XP_PAR_ACTION = {
    "Exercice-🟢 Débutant":  10,
    "Exercice-🟡 Moyen":     20,
    "Exercice-🟠 Confirmé":  35,
    "Exercice-🔴 Expert":    50,
    "CCF":                   40,
    "Boss":                  80,
}

def calculer_xp(records: list) -> int:
    """Calcule le total XP d'un élève depuis ses records Grist."""
    total = 0
    for r in records:
        if r.get("auto_evaluation", "") not in ("😊 Bien", "🌟 Très bien"):
            continue
        type_a = r.get("type_activite", "")
        diff   = r.get("niveau_difficulte", "")
        if "Boss" in type_a:
            total += XP_PAR_ACTION["Boss"]
        elif "CCF" in type_a:
            total += XP_PAR_ACTION["CCF"]
        else:
            total += XP_PAR_ACTION.get(f"Exercice-{diff}", 10)
    return total


def lire_classement_grist() -> list:
    """
    Lit les 10 meilleurs élèves par XP depuis Grist.
    Retourne une liste triée [{code_eleve, xp, streak}].
    """
    try:
        api_key  = st.secrets.get("GRIST_API_KEY", "")
        doc_id   = st.secrets.get("GRIST_DOC_ID", "")
        base_url = st.secrets.get("GRIST_URL", "https://grist.numerique.gouv.fr")
        if not api_key or not doc_id:
            return []
        url = f"{base_url}/api/docs/{doc_id}/tables/Suivi_eleves/records"
        headers = {"Authorization": f"Bearer {api_key}"}
        resp = requests.get(url, headers=headers, timeout=8)
        if resp.status_code != 200:
            return []
        records = [r["fields"] for r in resp.json().get("records", [])]

        # Grouper par élève
        from collections import defaultdict
        par_eleve = defaultdict(list)
        for r in records:
            code = r.get("code_eleve", "")
            if code and code != "anonyme":
                par_eleve[code].append(r)

        # Calculer XP et streak pour chaque élève
        classement = []
        for code, recs in par_eleve.items():
            xp     = calculer_xp(recs)
            streak = calculer_streak(recs).get("streak", 0)
            classement.append({"code": code, "xp": xp, "streak": streak})

        return sorted(classement, key=lambda x: x["xp"], reverse=True)[:10]
    except Exception:
        return []


def calculer_objectif(records: list, chapitre_actif: str, difficulte_active: str) -> dict:
    """
    Calcule un objectif personnalisé pour l'élève selon sa progression.
    Retourne : {message, emoji, priorite (0-3)}
    """
    from datetime import date
    s = calculer_streak(records)
    streak = s.get("streak", 0)
    derniere = s.get("derniere_date")
    today = date.today()

    # Priorité 1 — streak en danger
    if derniere and derniere < today and streak > 0:
        return {
            "emoji": "🔥",
            "message": f"Ton streak de {streak} jour{'s' if streak > 1 else ''} est en danger ! Fais au moins 1 exercice aujourd'hui.",
            "couleur": "#ef4444",
            "priorite": 3
        }

    # Priorité 2 — boss disponible sur le chapitre actif
    if chapitre_actif and records:
        prog = calculer_progression(records, chapitre_actif)
        for diff in ORDRE_NIVEAUX_DIFF:
            p = prog.get(diff, {})
            if p.get("valide") and not p.get("boss_vaincu"):
                mascotte = MASCOTTES[diff]["animal"]
                return {
                    "emoji": "⚔️",
                    "message": f"Le boss {mascotte} t'attend sur '{chapitre_actif[:35]}' ! Affronte-le.",
                    "couleur": "#7c3aed",
                    "priorite": 2
                }

    # Priorité 3 — proche de valider un niveau
    if chapitre_actif and records:
        prog = calculer_progression(records, chapitre_actif)
        for diff in ORDRE_NIVEAUX_DIFF:
            p = prog.get(diff, {})
            nb = p.get("nb_bonnes", 0)
            if 0 < nb < SEUIL_VALIDATION and not p.get("valide"):
                reste = SEUIL_VALIDATION - nb
                return {
                    "emoji": "🎯",
                    "message": f"Plus que {reste} bonne{'s' if reste > 1 else ''} éval en {diff.split()[-1]} sur ce chapitre pour valider le niveau !",
                    "couleur": "#f59e0b",
                    "priorite": 1
                }

    # Défaut — encouragement simple
    if streak == 0:
        return {
            "emoji": "🚀",
            "message": "Commence ta session du jour — même 1 exercice suffit pour lancer ton streak !",
            "couleur": "#6366f1",
            "priorite": 0
        }
    return {
        "emoji": "💪",
        "message": f"Super ! {streak} jour{'s' if streak > 1 else ''} d'affilée. Continue comme ça !",
        "couleur": "#22c55e",
        "priorite": 0
    }
    
    # Défaut — encouragement simple
    if streak == 0:
        return {
            "emoji": "🚀",
            "message": "Commence ta session du jour — même 1 exercice suffit pour lancer ton streak !",
            "couleur": "#6366f1",
            "priorite": 0
        }
    return {
        "emoji": "💪",
        "message": f"Super ! {streak} jour{'s' if streak > 1 else ''} d'affilée. Continue comme ça !",
        "couleur": "#22c55e",
        "priorite": 0
    }
    """Calcule le total XP d'un élève depuis ses records Grist."""
    total = 0
    for r in records:
        if r.get("auto_evaluation", "") not in ("😊 Bien", "🌟 Très bien"):
            continue
        type_a = r.get("type_activite", "")
        diff   = r.get("niveau_difficulte", "")
        if "Boss" in type_a:
            total += XP_PAR_ACTION["Boss"]
        elif "CCF" in type_a:
            total += XP_PAR_ACTION["CCF"]
        else:
            total += XP_PAR_ACTION.get(f"Exercice-{diff}", 10)
    return total

# ── SESSION STATE ─────────────────────────────────────────────
for key in ["generated_md", "generated_ccf_md", "meta_gen", "meta_ccf",
            "eval_gen_done", "eval_ccf_done", "grist_debug", "progression_cache",
            "boss_actif", "boss_niveau", "boss_chapitre", "boss_md", "eval_boss_done",
            "streak_cache", "interactif_sujet", "interactif_idx",
            "interactif_reponses", "interactif_termine", "classement_cache"]:
    if key not in st.session_state:
        st.session_state[key] = None

# ── SIDEBAR ──────────────────────────────────────────────────
with st.sidebar:
    st.markdown("""
    <div style="text-align:center;padding:16px 0 8px">
        <div style="font-size:2.5rem">🎓</div>
        <div style="font-family:'Outfit',sans-serif;font-weight:800;font-size:1.1rem;
                    color:#a5b4fc;letter-spacing:.5px">MATHS BAC PRO</div>
        <div style="font-size:.75rem;color:#475569;margin-top:2px">Entraînement IA</div>
    </div>
    """, unsafe_allow_html=True)

    st.divider()

    # ── Code élève ───────────────────────────────────────────
    st.markdown('<div style="font-family:Outfit,sans-serif;font-weight:700;color:#94a3b8;font-size:.8rem;letter-spacing:1px;margin-bottom:6px">🎮 TON IDENTIFIANT</div>', unsafe_allow_html=True)
    code_eleve = st.text_input(
        "Code élève",
        placeholder="Ex : ASSP-03",
        key="code_eleve",
        label_visibility="collapsed",
        help="Code distribué par ton professeur en début d'année."
    ).strip().upper()
    if code_eleve:
        # ── Chargement streak ────────────────────────────────
        if not st.session_state.streak_cache:
            records_streak = lire_progression_grist(code_eleve)
            st.session_state.streak_cache = calculer_streak(records_streak)

        s      = st.session_state.streak_cache or {}
        streak = s.get("streak", 0)
        record = s.get("record", 0)

        # ── Message de bienvenue personnalisé ────────────────
        from datetime import date
        derniere = s.get("derniere_date")
        if derniere == date.today():
            salut = "Bonne continuation"
        elif streak > 0:
            salut = "Bon retour"
        else:
            salut = "Bienvenue"

        if streak >= 7:   feu, couleur_bienv = "🔥🔥🔥", "#ef4444"
        elif streak >= 3: feu, couleur_bienv = "🔥🔥",   "#f97316"
        elif streak >= 1: feu, couleur_bienv = "🔥",     "#f59e0b"
        else:             feu, couleur_bienv = "👋",     "#6366f1"

        streak_txt = (
            f"{feu} {streak} jour{'s' if streak > 1 else ''} d'affilée !"
            if streak > 0 else "Lance ton premier streak aujourd'hui !"
        )

        st.markdown(
            f'<div style="background:linear-gradient(135deg,#13162a,#1e2235);'
            f'border:1px solid {couleur_bienv}44;border-radius:10px;'
            f'padding:12px 14px;margin:6px 0">'
            f'<div style="font-size:.82rem;color:#94a3b8">{salut} 👋</div>'
            f'<div style="font-family:Outfit,sans-serif;font-weight:800;'
            f'color:#e2e8f0;font-size:1rem;margin:2px 0">{code_eleve}</div>'
            f'<div style="font-size:.82rem;color:{couleur_bienv};margin-top:4px">'
            f'{streak_txt}</div>'
            + (f'<div style="font-size:.72rem;color:#475569;margin-top:2px">'
               f'Record : {record} jour{"s" if record > 1 else ""}</div>'
               if record > 1 else "")
            + f'</div>',
            unsafe_allow_html=True
        )
    else:
        st.markdown('<div class="warn-box">⚠️ Entre ton code pour sauvegarder ta progression.</div>', unsafe_allow_html=True)

    st.divider()

    st.markdown('<div style="font-family:Outfit,sans-serif;font-weight:700;color:#94a3b8;font-size:.8rem;letter-spacing:1px;margin-bottom:6px">🔑 CLÉ GEMINI</div>', unsafe_allow_html=True)
    cle_api = st.text_input("Clé API", type="password", key="gemini_key", label_visibility="collapsed")
    if cle_api:
        st.markdown('<div class="ok-box">✅ Prêt à générer !</div>', unsafe_allow_html=True)
    else:
        st.markdown('<div class="warn-box">⚠️ Clé API manquante.</div>', unsafe_allow_html=True)

    with st.expander("📋 Obtenir une clé gratuite (2 min)"):
        st.markdown("""
**1.** → [aistudio.google.com](https://aistudio.google.com/app/apikey)

**2.** Connecte-toi avec ton compte Google

**3.** Clique **"Create API Key"**

**4.** Copie et colle la clé ici

---
💡 **C'est gratuit** — pas de carte bancaire requise.
        """)

    st.divider()
    st.caption("Entraînement Bac Pro · Gemini 2.5 Flash")

# ── TITRE ────────────────────────────────────────────────────
st.markdown("""
<div style="padding:24px 0 8px">
    <div style="font-family:'Outfit',sans-serif;font-weight:800;font-size:2rem;
                background:linear-gradient(90deg,#a5b4fc,#c084fc,#67e8f9);
                -webkit-background-clip:text;-webkit-text-fill-color:transparent;
                background-clip:text;line-height:1.2">
        📚 Entraînement Bac Pro
    </div>
    <div style="color:#475569;font-size:.9rem;margin-top:4px">
        Exercices · CCF · Progression — propulsé par l'IA
    </div>
</div>
""", unsafe_allow_html=True)

# ── ONGLETS ──────────────────────────────────────────────────
tab_gen, tab_ccf, tab_graphique, tab_progression = st.tabs([
    "📝 Exercices d'entraînement",
    "🎯 Sujets CCF",
    "📈 Graphique GeoGebra",
    "📊 Ma progression",
])


# ─────────────────────────────────────────────────────────────
# ONGLET 1 — GÉNÉRATEUR
# ─────────────────────────────────────────────────────────────
with tab_gen:
    # ── Objectif du jour ─────────────────────────────────────
    if code_eleve and st.session_state.progression_cache:
        chap_actif = (st.session_state.meta_gen or {}).get("chapitre", "")
        diff_active = (st.session_state.meta_gen or {}).get("difficulte", "")
        obj = calculer_objectif(st.session_state.progression_cache, chap_actif, diff_active)
        st.markdown(
            f'<div style="background:#0d0f1a;border:1px solid {obj["couleur"]}55;'
            f'border-left:4px solid {obj["couleur"]};border-radius:0 10px 10px 0;'
            f'padding:10px 16px;margin-bottom:12px">'
            f'<span style="font-size:1.1rem">{obj["emoji"]}</span> '
            f'<span style="color:#e2e8f0;font-size:.88rem">{obj["message"]}</span>'
            f'</div>',
            unsafe_allow_html=True
        )

    st.subheader("📝 Générateur de sujets et exercices")
    col1, col2 = st.columns(2)

    with col1:
        cat = st.selectbox("Type d'établissement", list(NIVEAUX_CATEGORIES.keys()), key="gen_cat")
        niv = st.selectbox("Classe", NIVEAUX_CATEGORIES[cat], key="gen_niv")
        filiere = ""
        if cat in ["Bac Pro", "CAP"]:
            fil = st.selectbox("Filière", LISTE_FILIERES, key="gen_fil")
            if fil == "Autre (Préciser ci-dessous)":
                filiere = st.text_input("Filière libre", key="gen_spec")
            else:
                filiere = fil
                if fil in CONTEXTES_FILIERES:
                    st.info(f"📌 {CONTEXTES_FILIERES[fil]['nom_complet']}")

    with col2:
        mat = st.selectbox("Matière", MATIERES, key="gen_mat")
        chap = st.selectbox("Chapitre (BO)", get_chapitres(mat, niv, cat), key="gen_chap")

    consignes = st.text_area("Consignes particulières (optionnel)", height=80, key="gen_consignes",
                              placeholder="Ex : 3 exercices, niveau accessible, inclure un graphique…")

    # ── Niveau de difficulté ──────────────────────────────────
    st.markdown("**🎯 Niveau de difficulté (pédagogie différenciée)**")
    difficulte = st.select_slider(
        "Niveau",
        options=list(NIVEAUX_DIFFICULTE.keys()),
        value="🟡 Moyen",
        key="gen_diff",
        label_visibility="collapsed"
    )
    st.markdown(
        f'<div class="info-box">{difficulte} — {NIVEAUX_DIFFICULTE[difficulte]}</div>',
        unsafe_allow_html=True
    )

    # ── Mode banque ou génération ─────────────────────────────
    sujets_banque      = charger_banque(niv, filiere, chap, difficulte)
    sujets_interactifs = charger_banque_interactif(niv, filiere, chap, difficulte)
    nb_banque      = len(sujets_banque)
    nb_interactifs = len(sujets_interactifs)

    # ── Debug temporaire ──────────────────────────────────────
    with st.expander("🔍 Debug banque (temporaire)", expanded=False):
        st.write(f"**BANQUE_DIR** : `{BANQUE_DIR}`")
        st.write(f"**Dossier existe** : {os.path.exists(BANQUE_DIR)}")
        if os.path.exists(BANQUE_DIR):
            fichiers = sorted(os.listdir(BANQUE_DIR))
            st.write(f"**Fichiers présents ({len(fichiers)})** :")
            for f in fichiers:
                st.write(f"  - `{f}`")
        st.write(f"**Recherche** : niveau=`{niv}` filière=`{filiere}` chapitre=`{chap[:40]}` diff=`{difficulte}`")
        st.write(f"**Normalisé** : niv=`{_normalise(niv)}` fil=`{_normalise(filiere)}` diff=`{_normalise(difficulte.split(' ',1)[-1])}`")
        st.write(f"**Chap normalisé (45c)** : `{_normalise(chap)[:45]}`")
        st.write(f"**Résultat classique** : {nb_banque} sujet(s)")
        st.write(f"**Résultat interactif** : {nb_interactifs} sujet(s)")

    col_btn1, col_btn2, col_btn3 = st.columns(3)
    with col_btn1:
        btn_interactif = st.button(
            f"⚡ Mode interactif ({nb_interactifs} dispo)",
            type="primary",
            use_container_width=True,
            disabled=(nb_interactifs == 0),
            key="btn_interactif"
        )
        if nb_interactifs == 0:
            st.caption("Aucun sujet interactif en banque.")
    with col_btn2:
        btn_banque = st.button(
            f"📚 Sujet classique ({nb_banque} dispo)",
            use_container_width=True,
            disabled=(nb_banque == 0),
            key="btn_banque"
        )
        if nb_banque == 0:
            st.caption("Aucun sujet classique en banque.")
    with col_btn3:
        btn_gemini = st.button(
            "✨ Générer (Gemini)",
            use_container_width=True,
            key="btn_gemini"
        )

    if btn_interactif and nb_interactifs > 0:
        import random
        sujet_i = random.choice(sujets_interactifs)
        st.session_state.interactif_sujet   = sujet_i
        st.session_state.interactif_idx     = 0
        st.session_state.interactif_reponses = {}
        st.session_state.interactif_termine  = False
        st.session_state.generated_md        = None
        st.session_state.eval_gen_done       = None
        st.session_state.meta_gen = {
            "niveau": niv, "matiere": mat, "chapitre": chap,
            "filiere": filiere, "difficulte": difficulte, "source": "Banque interactif"
        }
        st.rerun()

    if btn_banque and nb_banque > 0:
        import random
        sujet = random.choice(sujets_banque)
        st.session_state.generated_md        = sujet["contenu"]
        st.session_state.interactif_sujet    = None
        st.session_state.eval_gen_done       = None
        st.session_state.meta_gen = {
            "niveau": niv, "matiere": mat, "chapitre": chap,
            "filiere": filiere, "difficulte": difficulte, "source": "Banque"
        }
        st.rerun()

    if btn_gemini:
        if not cle_api:
            st.error("🔑 Renseigne ta clé API dans le panneau gauche !")
        else:
            with st.spinner("⏳ Génération en cours…"):
                try:
                    res = call_gemini(cle_api, build_prompt_exercices(niv, cat, mat, chap, consignes, filiere, difficulte))
                    st.session_state.generated_md     = res
                    st.session_state.interactif_sujet = None
                    st.session_state.eval_gen_done    = None
                    st.session_state.meta_gen = {
                        "niveau": niv, "matiere": mat, "chapitre": chap,
                        "filiere": filiere, "difficulte": difficulte, "source": "Gemini"
                    }
                    # Envoyer dans Banque_propositions pour validation prof
                    envoyer_proposition_grist(
                        code_eleve or "anonyme",
                        st.session_state.meta_gen,
                        res
                    )
                    st.success("✅ Sujet généré !")
                except Exception as e:
                    if "429" in str(e):
                        st.error("⏱️ Quota dépassé (429). Attends 1 minute ou utilise un sujet de la banque.")
                    else:
                        st.error(f"Erreur API : {e}")

# ── BOSS DÉBLOQUÉ ─────────────────────────────────────

        m = st.session_state.metagen or {}
        if st.session_state.boss_actif and st.session_state.boss_chapitre == m.get("chapitre", ""):
            diff_boss = st.session_state.boss_niveau          # ← 12 espaces (3 niveaux)
            mascotte  = MASCOTTES.get(diff_boss, {})
            st.markdown("---")
            st.markdown(
                f'<div class="boss-banner">'
                f'<div style="font-size:4rem;margin-bottom:8px">{mascotte.get("animal","⚔️")}</div>'
                f'<div style="font-family:Outfit,sans-serif;font-size:1.5rem;font-weight:800;'
                f'color:#e9d5ff;letter-spacing:1px">BOSS DÉBLOQUÉ !</div>'
                f'<div style="font-size:1.1rem;font-weight:700;color:#c084fc;margin:6px 0">'
                f'{mascotte.get("nom","Boss").upper()} t\'attend…</div>'
                f'<div style="font-size:.9rem;color:#94a3b8;max-width:400px;margin:0 auto">'
                f'Tu as validé le niveau <strong style="color:#a5b4fc">'
                f'{diff_boss.split(" ",1)[-1]}</strong> sur '
                f'<strong style="color:#a5b4fc">{m.get("chapitre","")[:45]}</strong>.<br>'
                f'Bats ce boss pour débloquer le niveau suivant !</div>'
                f'</div>',
                unsafe_allow_html=True
            )
            col_boss1, col_boss2 = st.columns(2)
            with col_boss1:
                if st.button(f"⚔️ Affronter le {mascotte.get('nom','Boss')} !",
                             type="primary", use_container_width=True, key="btn_boss"):
                    if not cle_api:
                        st.error("🔑 Clé API manquante !")
                    else:
                        with st.spinner(f"⚔️ Le {mascotte.get('nom','Boss')} se prépare…"):
                            try:
                                prompt_boss = build_prompt_boss(
                                    niv, filiere,
                                    st.session_state.boss_chapitre,
                                    st.session_state.boss_niveau
                                )
                                res_boss = call_gemini(cle_api, prompt_boss)
                                st.session_state.boss_md        = res_boss
                                st.session_state.eval_boss_done = None
                                st.session_state.boss_actif     = False
                                st.rerun()
                            except Exception as e:
                                st.error(f"Erreur : {e}")
            with col_boss2:
                if st.button("⏭️ Plus tard", use_container_width=True, key="btn_boss_skip"):
                    st.session_state.boss_actif = False
                    st.rerun()

# ── MODE INTERACTIF ──────────────────────────────────────
    # À 4 espaces — HORS de generated_md pour fonctionner quand generated_md=None
    if st.session_state.get("interactif_sujet"):
        sujet_i   = st.session_state.interactif_sujet
        questions = sujet_i.get("questions", [])
        idx       = st.session_state.get("interactif_idx", 0)
        reponses  = st.session_state.get("interactif_reponses", {})
        termine   = st.session_state.get("interactif_termine", False)

        if sujet_i.get("contexte"):
            st.markdown(
                f'<div class="info-box">🏢 <strong>Mise en situation</strong><br>'
                f'{sujet_i["contexte"]}</div>',
                unsafe_allow_html=True
            )
        if sujet_i.get("rappel_cours"):
            st.markdown(
                f'<div style="background:#1a2040;border:1px solid #3d4480;border-radius:10px;'
                f'padding:12px 16px;margin:8px 0;font-size:.88rem;color:#a5b4fc">'
                f'📐 <strong>Rappel</strong> : {sujet_i["rappel_cours"]}</div>',
                unsafe_allow_html=True
            )

        if not termine:
            st.markdown(
                f'<div style="margin:12px 0 4px;font-size:.82rem;color:#64748b">'
                f'Question {idx + 1} / {len(questions)}</div>',
                unsafe_allow_html=True
            )
            st.progress((idx) / len(questions))

            q = questions[idx]
            st.markdown(
                f'<div style="background:#13162a;border:1px solid #3d4480;border-radius:12px;'
                f'padding:18px 20px;margin:12px 0">'
                f'<div style="font-size:.8rem;color:#6366f1;font-weight:700;margin-bottom:6px">'
                f'QUESTION {q["numero"]}</div>'
                f'<div style="color:#e2e8f0;font-size:.95rem">{q["enonce"]}</div>'
                f'</div>',
                unsafe_allow_html=True
            )

            col_rep, col_ind = st.columns([3, 1])
            with col_rep:
                unite = q.get("unite", "")
                rep_saisie = st.number_input(
                    f"Ta réponse {f'({unite})' if unite else ''}",
                    value=0.0, step=0.01, format="%.2f",
                    key=f"rep_q_{idx}"
                )
            with col_ind:
                if q.get("indice") and st.button("💡 Indice", key=f"ind_{idx}", use_container_width=True):
                    st.info(q["indice"])

            col_v, col_p = st.columns(2)
            with col_v:
                if st.button("✅ Valider", type="primary", use_container_width=True, key=f"val_{idx}"):
                    tolerance = float(q.get("tolerance", 0.5))
                    bonne = abs(rep_saisie - float(q["reponse"])) <= tolerance
                    reponses[idx] = {"saisie": rep_saisie, "bonne": bonne,
                                     "reponse": q["reponse"], "unite": unite,
                                     "corrige": q.get("corrige", "")}
                    st.session_state.interactif_reponses = reponses
                    if idx + 1 < len(questions):
                        st.session_state.interactif_idx = idx + 1
                    else:
                        st.session_state.interactif_termine = True
                    st.rerun()
            with col_p:
                if idx > 0 and st.button("⬅️ Précédente", use_container_width=True, key=f"prev_{idx}"):
                    st.session_state.interactif_idx = idx - 1
                    st.rerun()

        else:
            # ── RÉSULTATS FINAUX ──────────────────────────────
            nb_bonnes_i = sum(1 for r in reponses.values() if r["bonne"])
            nb_total_i  = len(questions)
            score_pct   = int(nb_bonnes_i / nb_total_i * 100) if nb_total_i else 0

            if score_pct == 100:
                st.balloons()
                verdict, couleur = "🏆 Parfait !", "#065f46"
            elif score_pct >= 70:
                verdict, couleur = "👍 Bien joué !", "#1e3a5f"
            elif score_pct >= 50:
                verdict, couleur = "😐 Peut mieux faire", "#3f2a00"
            else:
                verdict, couleur = "💪 Continue à t'entraîner !", "#3b0000"

            st.markdown(
                f'<div style="background:{couleur};border-radius:14px;padding:20px;'
                f'text-align:center;margin:12px 0">'
                f'<div style="font-size:2rem">{verdict}</div>'
                f'<div style="font-size:1.5rem;font-weight:800;color:white;margin:8px 0">'
                f'{nb_bonnes_i} / {nb_total_i} — {score_pct}%</div>'
                f'</div>',
                unsafe_allow_html=True
            )

            st.markdown("### 📋 Correction détaillée")
            for i, q in enumerate(questions):
                r     = reponses.get(i, {})
                icone = "✅" if r.get("bonne") else "❌"
                unite = r.get("unite", "")
                st.markdown(
                    f'<div style="background:#13162a;border-left:3px solid '
                    f'{"#22c55e" if r.get("bonne") else "#ef4444"};'
                    f'border-radius:0 10px 10px 0;padding:12px 16px;margin:6px 0">'
                    f'<div style="font-size:.8rem;color:#64748b">Q{q["numero"]}</div>'
                    f'<div style="color:#e2e8f0;margin:4px 0">{q["enonce"]}</div>'
                    f'<div style="font-size:.88rem;margin-top:6px">'
                    f'{icone} Ta réponse : <strong>{r.get("saisie","?"):.2f} {unite}</strong> — '
                    f'Bonne réponse : <strong>{float(q["reponse"]):.2f} {unite}</strong></div>'
                    f'<div style="font-size:.82rem;color:#94a3b8;margin-top:4px">'
                    f'💡 {r.get("corrige","")}</div>'
                    f'</div>',
                    unsafe_allow_html=True
                )

            st.divider()
            st.markdown("**🎯 Et ton ressenti ?**")
            col_e1, col_e2, col_e3, col_e4 = st.columns(4)
            eval_choix_i = None
            with col_e1:
                if st.button("😕 Difficile", use_container_width=True, key="eval_i_1"):
                    eval_choix_i = "😕 Difficile"
            with col_e2:
                if st.button("😐 Moyen", use_container_width=True, key="eval_i_2"):
                    eval_choix_i = "😐 Moyen"
            with col_e3:
                if st.button("😊 Bien", use_container_width=True, key="eval_i_3"):
                    eval_choix_i = "😊 Bien"
            with col_e4:
                if st.button("🌟 Très bien", use_container_width=True, key="eval_i_4"):
                    eval_choix_i = "🌟 Très bien"

            if eval_choix_i:
                m_i = st.session_state.meta_gen or {}
                envoyer_grist(
                    code_eleve or "anonyme",
                    f"Exercice-interactif ({score_pct}%)",
                    {**m_i, "score_auto": f"{nb_bonnes_i}/{nb_total_i}"},
                    eval_choix_i
                )
                st.session_state.eval_gen_done    = eval_choix_i
                st.session_state.streak_cache     = None
                st.session_state.interactif_sujet = None
                st.rerun()

            if st.button("🔄 Nouvel exercice", use_container_width=True, key="btn_reset_i"):
                st.session_state.interactif_sujet    = None
                st.session_state.interactif_reponses = {}
                st.session_state.interactif_termine  = False
                st.rerun()

    # ── SUJET CLASSIQUE ───────────────────────────────────────
    # À 4 espaces — séparé du bloc interactif
    if st.session_state.generated_md:
        m = st.session_state.meta_gen or {}
        st.divider()
        badges = (
            f'<span class="badge">📚 {m.get("matiere","")}</span>'
            f'<span class="badge">🎓 {m.get("niveau","")}</span>'
            f'<span class="badge">📖 {m.get("chapitre","")}</span>'
        )
        if m.get("filiere"):
            badges += f'<span class="badge">🏢 {m["filiere"]}</span>'
        if m.get("difficulte"):
            badges += f'<span class="badge">{m["difficulte"]}</span>'
        st.markdown(badges, unsafe_allow_html=True)
        st.markdown(st.session_state.generated_md)

        # ── Auto-évaluation ───────────────────────────────────
        st.divider()
        st.markdown("**🎯 Comment tu t'en es sorti ?**")
        col_e1, col_e2, col_e3, col_e4 = st.columns(4)
        eval_choix = None
        with col_e1:
            if st.button("😕 Difficile", use_container_width=True, key="eval_gen_1"):
                eval_choix = "😕 Difficile"
        with col_e2:
            if st.button("😐 Moyen", use_container_width=True, key="eval_gen_2"):
                eval_choix = "😐 Moyen"
        with col_e3:
            if st.button("😊 Bien", use_container_width=True, key="eval_gen_3"):
                eval_choix = "😊 Bien"
        with col_e4:
            if st.button("🌟 Très bien", use_container_width=True, key="eval_gen_4"):
                eval_choix = "🌟 Très bien"

        if eval_choix:
            envoyer_grist(code_eleve or "anonyme", "Exercice", m, eval_choix)
            if m.get("source") == "Gemini" and st.session_state.generated_md:
                success = envoyer_proposition_grist(
                    code_eleve or "anonyme", m,
                    st.session_state.generated_md, eval_choix
                )
                if success:
                    st.success("📤 Ton sujet Gemini proposé au prof ! 👏")        
            st.session_state.eval_gen_done = eval_choix
            st.session_state.streak_cache  = None
            if eval_choix in ("😊 Bien", "🌟 Très bien") and code_eleve:
                records_check = lire_progression_grist(code_eleve)
                records_check.append({
                    "chapitre": m.get("chapitre", ""),
                    "niveau_difficulte": m.get("difficulte", ""),
                    "auto_evaluation": eval_choix,
                    "type_activite": "Exercice",
                })
                prog = calculer_progression(records_check, m.get("chapitre", ""))
                diff_actuel = m.get("difficulte", "")
                if (prog.get(diff_actuel, {}).get("valide") and
                        not prog.get(diff_actuel, {}).get("boss_vaincu") and
                        diff_actuel in ORDRE_NIVEAUX_DIFF):
                    st.session_state.boss_actif    = True
                    st.session_state.boss_niveau   = diff_actuel
                    st.session_state.boss_chapitre = m.get("chapitre", "")
            st.rerun()

        if st.session_state.eval_gen_done:
            st.markdown(
                f'<div class="ok-box">✅ Auto-évaluation enregistrée : '
                f'<strong>{st.session_state.eval_gen_done}</strong> — continue comme ça !</div>',
                unsafe_allow_html=True
            )
            # Neige si nouveau record de streak
            if st.session_state.streak_cache is None:
                records_s = lire_progression_grist(code_eleve or "")
                s_data    = calculer_streak(records_s)
                st.session_state.streak_cache = s_data
                if s_data.get("nouveau_record") and s_data.get("streak", 0) >= 3:
                    st.snow()  # ✅ corrigé — était st.snow()(diff_boss, {})

        # ── BOSS DÉBLOQUÉ ─────────────────────────────────────
        # À 8 espaces — dans generated_md, affiché après eval_gen_done
        if st.session_state.boss_actif and st.session_state.boss_chapitre == m.get("chapitre", ""):
            diff_boss = st.session_state.boss_niveau
            mascotte  = MASCOTTES.get(diff_boss, {})
            st.markdown("---")
            st.markdown(
                f'<div class="boss-banner">'
                f'<div style="font-size:4rem;margin-bottom:8px">{mascotte.get("animal","⚔️")}</div>'
                f'<div style="font-family:Outfit,sans-serif;font-size:1.5rem;font-weight:800;'
                f'color:#e9d5ff;letter-spacing:1px">BOSS DÉBLOQUÉ !</div>'
                f'<div style="font-size:1.1rem;font-weight:700;color:#c084fc;margin:6px 0">'
                f'{mascotte.get("nom","Boss").upper()} t\'attend…</div>'
                f'<div style="font-size:.9rem;color:#94a3b8;max-width:400px;margin:0 auto">'
                f'Tu as validé le niveau <strong style="color:#a5b4fc">'
                f'{diff_boss.split(" ",1)[-1]}</strong> sur '
                f'<strong style="color:#a5b4fc">{m.get("chapitre","")[:45]}</strong>.<br>'
                f'Bats ce boss pour débloquer le niveau suivant !</div>'
                f'</div>',
                unsafe_allow_html=True
            )
            col_boss1, col_boss2 = st.columns(2)
            with col_boss1:
                if st.button(f"⚔️ Affronter le {mascotte.get('nom','Boss')} !",
                             type="primary", use_container_width=True, key="btn_boss"):
                    if not cle_api:
                        st.error("🔑 Clé API manquante !")
                    else:
                        with st.spinner(f"⚔️ Le {mascotte.get('nom','Boss')} se prépare…"):
                            try:
                                prompt_boss = build_prompt_boss(
                                    niv, filiere,
                                    st.session_state.boss_chapitre,
                                    st.session_state.boss_niveau
                                )
                                res_boss = call_gemini(cle_api, prompt_boss)
                                st.session_state.boss_md        = res_boss
                                st.session_state.eval_boss_done = None
                                st.session_state.boss_actif     = False
                                st.rerun()
                            except Exception as e:
                                st.error(f"Erreur : {e}")
            with col_boss2:
                if st.button("⏭️ Plus tard", use_container_width=True, key="btn_boss_skip"):
                    st.session_state.boss_actif = False
                    st.rerun()

        # ── SUJET BOSS AFFICHÉ ────────────────────────────────
        if st.session_state.boss_md:
            diff_boss = st.session_state.boss_niveau or "🟢 Débutant"
            mascotte  = MASCOTTES.get(diff_boss, {})
            st.markdown(
                f'<div style="background:#1e1b4b;color:white;padding:12px 16px;'
                f'border-radius:8px;margin:8px 0;font-size:1.1rem;font-weight:bold">'
                f'{mascotte.get("animal","⚔️")} DÉFI BOSS — {mascotte.get("nom","Boss").upper()}'
                f'</div>',
                unsafe_allow_html=True
            )
            st.markdown(st.session_state.boss_md)

            st.divider()
            st.markdown("**⚔️ As-tu vaincu le Boss ?**")
            col_b1, col_b2 = st.columns(2)
            with col_b1:
                if st.button("😤 Pas encore…", use_container_width=True, key="boss_fail"):
                    envoyer_grist(
                        code_eleve or "anonyme", f"Boss-{diff_boss}",
                        {"chapitre": st.session_state.boss_chapitre or "",
                         "niveau": niv, "filiere": filiere, "matiere": mat},
                        "😕 Difficile"
                    )
                    st.session_state.eval_boss_done = "fail"
                    st.session_state.boss_md = None
                    st.rerun()
            with col_b2:
                if st.button("🏆 Boss vaincu !", use_container_width=True, key="boss_win"):
                    envoyer_grist(
                        code_eleve or "anonyme", f"Boss-{diff_boss}",
                        {"chapitre": st.session_state.boss_chapitre or "",
                         "niveau": niv, "filiere": filiere, "matiere": mat},
                        "🌟 Très bien"
                    )
                    st.session_state.eval_boss_done    = "win"
                    st.session_state.boss_md           = None
                    st.session_state.progression_cache = None
                    st.rerun()

        if st.session_state.eval_boss_done == "win":
            st.balloons()
            diff_boss = st.session_state.boss_niveau or ""
            msg = MESSAGES_VICTOIRE.get(diff_boss, "🏆 Félicitations !")
            st.markdown(
                f'<div style="background:linear-gradient(135deg,#052e16,#14532d);'
                f'border:2px solid #22c55e;color:white;padding:20px;border-radius:14px;'
                f'text-align:center;font-family:Outfit,sans-serif;'
                f'box-shadow:0 0 30px rgba(34,197,94,.3);margin:8px 0">'
                f'<div style="font-size:2.5rem;margin-bottom:8px">🏆</div>'
                f'<div style="font-size:1.2rem;font-weight:800;color:#86efac">{msg}</div>'
                f'</div>',
                unsafe_allow_html=True
            )
        elif st.session_state.eval_boss_done == "fail":
            st.markdown(
                '<div class="info-box">💪 Pas grave ! Continue à t\'entraîner et reviens affronter le Boss.</div>',
                unsafe_allow_html=True
            )

        # ── Téléchargements ───────────────────────────────────
        titre_doc = f"Sujet — {m.get('matiere','')} {m.get('niveau','')} {m.get('difficulte','')}"
        c1, c2, c3 = st.columns(3)
        with c1:
            st.download_button("📥 .md", st.session_state.generated_md,
                               file_name="sujet.md", mime="text/markdown", key="dl1_md")
        with c2:
            st.download_button("📄 .txt", st.session_state.generated_md,
                               file_name="sujet.txt", mime="text/plain", key="dl1_txt")
        with c3:
            st.download_button("📝 .docx",
                markdown_to_docx(st.session_state.generated_md, titre_doc),
                file_name="sujet.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl1_docx")


# ─────────────────────────────────────────────────────────────
# ONGLET 2 — CCF ENTRAÎNEMENT UNIQUEMENT
# ─────────────────────────────────────────────────────────────
with tab_ccf:
    st.subheader("🎯 Sujets CCF — Entraînement")
    st.markdown('<div class="info-box">📋 <strong>Mode Entraînement</strong> — Structure CCF officielle pour t\'entraîner avant l\'examen. Sans en-tête officiel.</div>', unsafe_allow_html=True)

    # Forcer le mode entraînement — pas d'officiel pour les élèves
    is_officiel = False

    col1, col2 = st.columns(2)
    with col1:
        ccf_cat = st.selectbox("Catégorie", ["Bac Pro", "CAP"], key="ccf_cat")
        # Les CCF n'ont lieu qu'en 1ère et Terminale — 2nde Pro exclue
        niveaux_ccf = [n for n in NIVEAUX_CATEGORIES[ccf_cat] if n != "2nde Pro"]
        ccf_niv = st.selectbox("Classe", niveaux_ccf, key="ccf_niv")
        ccf_fil = st.selectbox("Filière", LISTE_FILIERES, key="ccf_fil")
        ccf_filiere = ""
        if ccf_fil == "Autre (Préciser ci-dessous)":
            ccf_filiere = st.text_input("Filière libre", key="ccf_spec")
        else:
            ccf_filiere = ccf_fil
            if ccf_fil in CONTEXTES_FILIERES:
                st.info(f"📌 {CONTEXTES_FILIERES[ccf_fil]['nom_complet']}")

    with col2:
        ccf_mat = st.selectbox("Matière", MATIERES, key="ccf_mat")
        ccf_duree = st.selectbox("Durée de l'épreuve", ["30 min", "45 min", "1h00", "1h30", "2h00"], index=1, key="ccf_duree")
        num_sit = st.number_input("N° de situation CCF", min_value=1, max_value=3, value=1, step=1, key="ccf_num_sit")

    # ── Chapitres ────────────────────────────────────────────
    st.markdown("**📖 Chapitres évalués**")
    chapitres_dispo = get_chapitres(ccf_mat, ccf_niv, ccf_cat)
    col_a, col_b = st.columns(2)
    with col_a:
        ccf_chap = st.selectbox("Chapitre — Partie A", chapitres_dispo, key="ccf_chap")
    with col_b:
        double_chap = st.checkbox("Ajouter un 2ème chapitre (Partie B)", key="ccf_double_chap")
        if double_chap:
            chapitres_b = [c for c in chapitres_dispo if c != ccf_chap]
            ccf_chap_b = st.selectbox("Chapitre — Partie B", chapitres_b, key="ccf_chap_b")
        else:
            ccf_chap_b = ""

    if double_chap and ccf_chap_b:
        st.markdown(
            f'<div class="info-box">📐 <strong>Partie A :</strong> {ccf_chap} &nbsp;|&nbsp; '
            f'<strong>Partie B :</strong> {ccf_chap_b}<br>'
            f'Les deux parties s\'appuieront sur la même situation professionnelle.</div>',
            unsafe_allow_html=True
        )

    avec_corrige = st.checkbox(
        "📝 Inclure le corrigé détaillé *(document professeur)*",
        value=True,
        key="ccf_avec_corrige",
        help="Décochez pour générer uniquement le sujet élève — réponse plus courte, moins de tokens consommés."
    )
    if not avec_corrige:
        st.markdown('<div class="info-box">💡 Sujet seul — génération plus rapide et moins gourmande en quota.</div>', unsafe_allow_html=True)

    ccf_consignes = st.text_area("Consignes spécifiques (optionnel)", height=80, key="ccf_consignes",
                                  placeholder="Ex : Situation en EHPAD, tableau de données, niveau accessible…")

    btn_lbl = "📄 Générer le Sujet CCF Officiel" if is_officiel else "📋 Générer le Sujet d'Entraînement CCF"
    if st.button(btn_lbl, type="primary", use_container_width=True):
        if not cle_api:
            st.error("🔑 Renseigne ta clé API dans le panneau gauche !")
        else:
            with st.spinner("⏳ Génération du sujet CCF…"):
                try:
                    if is_officiel:
                        prompt = build_prompt_ccf_officiel(
                            ccf_niv, ccf_cat, ccf_mat, ccf_chap,
                            ccf_consignes, ccf_filiere, ccf_duree, str(num_sit),
                            avec_corrige=avec_corrige,
                            chapitre_b=ccf_chap_b
                        )
                    else:
                        prompt = build_prompt_ccf_entrainement(
                            ccf_niv, ccf_cat, ccf_mat, ccf_chap,
                            ccf_consignes, ccf_filiere,
                            avec_corrige=avec_corrige,
                            chapitre_b=ccf_chap_b
                        )
                    res = call_gemini(cle_api, prompt)
                    if res:
                        st.session_state.generated_ccf_md = res
                        st.session_state.eval_ccf_done = None  # reset éval
                        chap_label = ccf_chap
                        if ccf_chap_b:
                            chap_label += f" + {ccf_chap_b}"
                        st.session_state.meta_ccf = {
                            "niveau": ccf_niv,
                            "matiere": ccf_mat,
                            "chapitre": chap_label,
                            "filiere": ccf_filiere,
                            "mode": "Entraînement",
                            "num_situation": str(num_sit),
                            "duree": ccf_duree,
                            "annee_scolaire": "2025/2026",
                        }
                        st.success("✅ Sujet CCF généré !")
                    else:
                        st.error("❌ L'IA n'a pas renvoyé de texte. Réessayez.")
                except Exception as e:
                    if "429" in str(e):
                        st.error("⏱️ Quota dépassé (429). Attendez 1 minute et réessayez.")
                    else:
                        st.error(f"Erreur API : {e}")

    if st.session_state.generated_ccf_md:
        m = st.session_state.meta_ccf or {}
        st.divider()
        badges = (
            f'<span class="badge">🎯 CCF {m.get("mode","")}</span>'
            f'<span class="badge">📚 {m.get("matiere","")}</span>'
            f'<span class="badge">🎓 {m.get("niveau","")}</span>'
            f'<span class="badge">📖 {m.get("chapitre","")}</span>'
        )
        if m.get("filiere"):
            badges += f'<span class="badge">🏢 {m["filiere"]}</span>'
        if m.get("duree"):
            badges += f'<span class="badge">⏱ {m["duree"]}</span>'
        st.markdown(badges, unsafe_allow_html=True)
        st.markdown(st.session_state.generated_ccf_md)

        # ── Auto-évaluation + envoi Grist ────────────────────
        st.divider()
        st.markdown("**🎯 Comment tu t'en es sorti ?**")
        col_c1, col_c2, col_c3, col_c4 = st.columns(4)
        eval_ccf_choix = None
        with col_c1:
            if st.button("😕 Difficile", use_container_width=True, key="eval_ccf_1"):
                eval_ccf_choix = "😕 Difficile"
        with col_c2:
            if st.button("😐 Moyen", use_container_width=True, key="eval_ccf_2"):
                eval_ccf_choix = "😐 Moyen"
        with col_c3:
            if st.button("😊 Bien", use_container_width=True, key="eval_ccf_3"):
                eval_ccf_choix = "😊 Bien"
        with col_c4:
            if st.button("🌟 Très bien", use_container_width=True, key="eval_ccf_4"):
                eval_ccf_choix = "🌟 Très bien"

        if eval_ccf_choix:
            envoyer_grist(code_eleve or "anonyme", "CCF", m, eval_ccf_choix)
            st.session_state.eval_ccf_done = eval_ccf_choix
            st.session_state.streak_cache  = None  # forcer recalcul streak
            st.rerun()

        if st.session_state.eval_ccf_done:
            st.markdown(f'<div class="ok-box">✅ Auto-évaluation enregistrée : <strong>{st.session_state.eval_ccf_done}</strong> — continue comme ça !</div>', unsafe_allow_html=True)

        titre_doc = f"CCF_{m.get('mode','')}_{m.get('matiere','')}_{m.get('niveau','')}"
        st.subheader("📥 Télécharger")

        if is_officiel:
            c1, c2 = st.columns(2)
            with c1:
                docx_off = generate_ccf_officiel_docx(
                    st.session_state.generated_ccf_md, m, "Lino Ventura (Ozoir-la-Ferrière)"
                )
                if docx_off:
                    st.download_button(
                        "🏛️ Word Officiel (avec en-tête et fiche)",
                        docx_off,
                        file_name=f"{titre_doc}_officiel.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary",
                        key="dl_ccf_officiel"
                    )
            with c2:
                st.download_button("📥 .md (brut)", st.session_state.generated_ccf_md,
                                   file_name=f"{titre_doc}.md", mime="text/markdown", key="dl2_md_off")
        else:
            c1, c2, c3 = st.columns(3)
            with c1:
                st.download_button("📝 .docx",
                    markdown_to_docx(st.session_state.generated_ccf_md, titre_doc),
                    file_name=f"{titre_doc}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="dl2_docx")
            with c2:
                st.download_button("📥 .md", st.session_state.generated_ccf_md,
                                   file_name=f"{titre_doc}.md", mime="text/markdown", key="dl2_md")
            with c3:
                st.download_button("📄 .txt", st.session_state.generated_ccf_md,
                                   file_name=f"{titre_doc}.txt", mime="text/plain", key="dl2_txt")

# ─────────────────────────────────────────────────────────────
# ONGLET 3 — MA PROGRESSION
# ─────────────────────────────────────────────────────────────
with tab_graphique:
    st.markdown("""
    <div style="font-family:'Outfit',sans-serif;font-weight:800;font-size:1.4rem;
                color:#e2e8f0;padding:8px 0 4px">📈 Laboratoire Graphique</div>
    """, unsafe_allow_html=True)
    st.markdown(
        '<div class="info-box">📐 Trace tes fonctions, vérifie tes dérivées, '
        'explore les paraboles — tape directement dans la barre de saisie en bas à gauche.</div>',
        unsafe_allow_html=True
    )

    # Suggestions de fonctions selon le chapitre sélectionné
    suggestions = {
        "Fonctions polynômes de degré 2": ["f(x) = x^2 - 4x + 3", "g(x) = -2x^2 + 4x"],
        "Fonctions polynômes de degré 3": ["f(x) = x^3 - 3x", "g(x) = x^3 - x^2 - x + 1"],
        "Fonction dérivée": ["f(x) = x^2 + 2x", "f'(x) = 2x + 2"],
        "Suites": ["f(x) = 2x + 1", "g(x) = 3 * 1.5^x"],
        "Fonctions exponentielles": ["f(x) = exp(x)", "g(x) = 2^x"],
    }

    # Trouver une suggestion selon le chapitre actif
    chap_actif = (st.session_state.meta_gen or {}).get("chapitre", "")
    suggestion_actuelle = None
    for mot_cle, exemples in suggestions.items():
        if mot_cle.lower() in chap_actif.lower():
            suggestion_actuelle = exemples
            break

    if suggestion_actuelle:
        st.markdown(
            f'<div class="ok-box">💡 Exercice en cours : essaie de tracer '
            f'<code>{suggestion_actuelle[0]}</code></div>',
            unsafe_allow_html=True
        )

    # Intégration GeoGebra via l'API officielle
    st.components.v1.html("""
    <!DOCTYPE html>
    <html>
    <head>
        <script src="https://www.geogebra.org/apps/deployggb.js"></script>
    </head>
    <body style="margin:0;padding:0;background:#13162a;">
        <div id="ggb-element"></div>
        <script>
            var params = {
                "appName": "graphing",
                "width": 780,
                "height": 520,
                "showToolBar": true,
                "showAlgebraInput": true,
                "showMenuBar": false,
                "showFullscreenButton": true,
                "language": "fr",
                "showResetIcon": true,
                "algebraInputPosition": "bottom"
            };
            var applet = new GGBApplet(params, true);
            window.addEventListener("load", function() {
                applet.inject("ggb-element");
            });
        </script>
    </body>
    </html>
    """, height=540)

# ─────────────────────────────────────────────────────────────
# ONGLET 4 — MA PROGRESSION
# ─────────────────────────────────────────────────────────────
with tab_progression:
    st.markdown("""
    <div style="font-family:'Outfit',sans-serif;font-weight:800;font-size:1.4rem;
                color:#e2e8f0;padding:8px 0 4px">📊 Ma progression</div>
    """, unsafe_allow_html=True)

    if not code_eleve:
        st.markdown('<div class="warn-box">⚠️ Entre ton code élève dans le panneau gauche pour voir ta progression.</div>', unsafe_allow_html=True)
    else:
        col_refresh, col_info = st.columns([1, 3])
        with col_refresh:
            if st.button("🔄 Actualiser", use_container_width=True):
                st.session_state.progression_cache = None

        with col_info:
            st.markdown(f'<span style="color:#475569;font-size:.85rem">Joueur : <strong style="color:#a5b4fc">{code_eleve}</strong></span>', unsafe_allow_html=True)

        # Chargement depuis Grist (avec cache session)
        if not st.session_state.progression_cache:
            with st.spinner("⚡ Chargement de ta progression…"):
                records = lire_progression_grist(code_eleve)
                st.session_state.progression_cache = records
        else:
            records = st.session_state.progression_cache

        if not records:
            st.markdown('<div class="info-box">📭 Aucune activité enregistrée. Commence par faire des exercices !</div>', unsafe_allow_html=True)
        else:
            # ── XP Banner ────────────────────────────────────
            total_xp  = calculer_xp(records)
            nb_total  = len(records)
            nb_bonnes = sum(1 for r in records if r.get("auto_evaluation","") in ("😊 Bien","🌟 Très bien"))
            taux      = int(nb_bonnes / nb_total * 100) if nb_total else 0

            # Niveau XP
            if total_xp < 100:    rang, rang_label = "🥉", "Apprenti"
            elif total_xp < 300:  rang, rang_label = "🥈", "Initié"
            elif total_xp < 600:  rang, rang_label = "🥇", "Confirmé"
            elif total_xp < 1000: rang, rang_label = "💎", "Expert"
            else:                 rang, rang_label = "👑", "Maître"

            st.markdown(
                f'<div class="xp-banner">'
                f'<div style="font-size:2.5rem">{rang}</div>'
                f'<div style="flex:1">'
                f'<div style="font-family:Outfit,sans-serif;font-weight:700;color:#e2e8f0">{rang_label}</div>'
                f'<div class="xp-label">Rang actuel</div>'
                f'</div>'
                f'<div style="text-align:center">'
                f'<div class="xp-value">⚡ {total_xp} XP</div>'
                f'<div class="xp-label">points d\'expérience</div>'
                f'</div>'
                f'<div style="text-align:center;border-left:1px solid #2d3561;padding-left:16px">'
                f'<div style="font-family:Outfit,sans-serif;font-weight:800;font-size:1.3rem;color:#86efac">{taux}%</div>'
                f'<div class="xp-label">taux de réussite</div>'
                f'</div>'
                f'</div>',
                unsafe_allow_html=True
            )

            st.divider()

            # ── Progression par chapitre × niveau ────────────
            st.markdown("### 📖 Progression par chapitre")

            from collections import defaultdict
            tous_chap = sorted(set(r.get("chapitre", "") for r in records if r.get("chapitre")))

            for chap in tous_chap:
                prog_chap = calculer_progression(records, chap)
                niveaux_html = ""
                for diff in ORDRE_NIVEAUX_DIFF:
                    p         = prog_chap[diff]
                    mascotte  = MASCOTTES[diff]
                    animal    = mascotte["animal"]
                    nom       = mascotte["nom"]
                    couleur   = mascotte["couleur"]

                    if p["boss_vaincu"]:
                        cell = (f'<span title="{nom} vaincu !" style="font-size:1.5rem;'
                                f'filter:drop-shadow(0 0 4px {couleur})">{animal}👑</span>')
                    elif p["valide"]:
                        nb_ok = p["nb_bonnes"]
                        cell = (f'<span title="{nom} — Boss disponible ! ({nb_ok}/{SEUIL_VALIDATION})" '
                                f'style="font-size:1.5rem;opacity:.9">{animal}⚔️</span>')
                    elif p["nb_bonnes"] > 0:
                        nb_ok = p["nb_bonnes"]
                        cell = (f'<span title="{nom} — en cours ({nb_ok}/{SEUIL_VALIDATION})" '
                                f'style="font-size:1.5rem;filter:grayscale(60%)">{animal}</span>')
                    else:
                        cell = (f'<span title="{nom} — pas encore commencé" '
                                f'style="font-size:1.5rem;filter:grayscale(100%);opacity:.3">{animal}</span>')

                    niveaux_html += f'<span style="margin-right:8px">{cell}</span>'

                nb_act_chap = sum(1 for r in records if r.get("chapitre") == chap)
                chap_court  = chap[:50] + "…" if len(chap) > 50 else chap
                st.markdown(
                    f'<div style="display:flex;align-items:center;gap:12px;padding:10px 0;'
                    f'border-bottom:1px solid #eee">'
                    f'<div style="flex:1;font-size:.9rem"><strong>{chap_court}</strong></div>'
                    f'<div style="display:flex;align-items:center">{niveaux_html}</div>'
                    f'<div style="color:#888;font-size:.8rem;white-space:nowrap">'
                    f'{nb_act_chap} activité(s)</div>'
                    f'</div>',
                    unsafe_allow_html=True
                )

            st.markdown(
                '<div style="margin-top:14px;font-size:.82rem;color:#888;line-height:1.8">'
                '🐣 Poussin = Débutant &nbsp;|&nbsp; 🦊 Renard = Moyen &nbsp;|&nbsp; '
                '🦁 Lion = Confirmé &nbsp;|&nbsp; 🐉 Dragon = Expert<br>'
                f'⚔️ = Boss disponible &nbsp;|&nbsp; 👑 = Boss vaincu &nbsp;|&nbsp; '
                f'🔘 grisé = en cours &nbsp;|&nbsp; ⬜ = pas encore commencé<br>'
                f'Seuil de validation : {SEUIL_VALIDATION} bonnes évals minimum par niveau'
                f'</div>',
                unsafe_allow_html=True
            )

            st.divider()

            # ── Compétences BO ────────────────────────────────
            st.markdown("### 🎯 Compétences BO travaillées")
            st.caption("Basé sur les types d'exercices effectués et les chapitres travaillés.")

            COMP_PAR_CHAP = {
                "Probabilités": ["Analyser / Raisonner", "Réaliser", "Valider"],
                "Statistiques": ["S'approprier", "Réaliser", "Communiquer"],
                "Suites": ["Réaliser", "Valider", "Communiquer"],
                "Fonctions": ["S'approprier", "Analyser / Raisonner", "Réaliser"],
                "Vecteurs": ["S'approprier", "Réaliser"],
                "Trigonométrie": ["Réaliser", "Valider"],
                "Algorithmique": ["Analyser / Raisonner", "Réaliser"],
                "Calculs commerciaux": ["S'approprier", "Réaliser", "Communiquer"],
                "Géométrie": ["S'approprier", "Réaliser", "Valider"],
            }
            TOUTES_COMP = ["S'approprier", "Analyser / Raisonner", "Réaliser", "Valider", "Communiquer"]

            # Compter les activités par compétence (estimation)
            comp_count = defaultdict(int)
            for r in records:
                chap_r = r.get("chapitre", "")
                for mot_cle, comps in COMP_PAR_CHAP.items():
                    if mot_cle.lower() in chap_r.lower():
                        for c in comps:
                            comp_count[c] += 1

            # Si aucune correspondance trouvée, toutes les comp reçoivent 1
            if not any(comp_count.values()):
                for c in TOUTES_COMP:
                    comp_count[c] = nb_total // 5 or 1

            max_count = max(comp_count.values()) if comp_count else 1
            for comp in TOUTES_COMP:
                val  = comp_count.get(comp, 0)
                pct  = int(val / max_count * 100)
                bar  = "█" * (pct // 10) + "░" * (10 - pct // 10)
                st.markdown(
                    f'<div style="display:flex;align-items:center;gap:10px;margin:4px 0">'
                    f'<div style="width:180px;font-size:.85rem">{comp}</div>'
                    f'<div style="font-family:monospace;color:#4a6cf7">{bar}</div>'
                    f'<div style="font-size:.8rem;color:#888">{val} activité(s)</div>'
                    f'</div>',
                    unsafe_allow_html=True
                )

            st.divider()

            # ── Historique récent ─────────────────────────────
            with st.expander("📅 Voir l'historique complet"):
                import pandas as pd
                df = pd.DataFrame(records)
                cols_affich = [c for c in ["date", "heure", "type_activite", "chapitre",
                                            "niveau_difficulte", "auto_evaluation", "source"]
                               if c in df.columns]
                st.dataframe(df[cols_affich].sort_values("date", ascending=False),
                             use_container_width=True, hide_index=True)

            st.divider()

            # ── Classement de la classe ───────────────────────
            st.markdown("### 🏆 Classement de la classe")
            col_ref_class, _ = st.columns([1, 3])
            with col_ref_class:
                if st.button("🔄 Actualiser le classement", use_container_width=True, key="refresh_class"):
                    st.session_state.classement_cache = None

            if not st.session_state.classement_cache:
                with st.spinner("Chargement du classement…"):
                    st.session_state.classement_cache = lire_classement_grist()

            classement = st.session_state.classement_cache or []

            if not classement:
                st.markdown('<div class="info-box">📭 Pas encore assez de données pour afficher le classement.</div>', unsafe_allow_html=True)
            else:
                MEDAILLES = {1: "🥇", 2: "🥈", 3: "🥉"}
                for i, joueur in enumerate(classement, 1):
                    est_moi   = joueur["code"] == code_eleve
                    medaille  = MEDAILLES.get(i, f"**{i}.**")
                    streak_j  = joueur.get("streak", 0)
                    feu_j     = "🔥" * min(streak_j, 3) if streak_j > 0 else ""
                    bg_color  = "#1e2a1e" if est_moi else "#13162a"
                    border    = "#22c55e" if est_moi else "#2d3561"
                    toi       = " ← Toi" if est_moi else ""

                    st.markdown(
                        f'<div style="display:flex;align-items:center;gap:12px;'
                        f'background:{bg_color};border:1px solid {border};'
                        f'border-radius:10px;padding:10px 16px;margin:4px 0">'
                        f'<div style="font-size:1.2rem;min-width:32px">{medaille}</div>'
                        f'<div style="flex:1;font-family:Outfit,sans-serif;'
                        f'font-weight:{"800" if est_moi else "600"};color:#e2e8f0">'
                        f'{joueur["code"]}{toi}</div>'
                        f'<div style="color:#fbbf24;font-weight:700">⚡ {joueur["xp"]} XP</div>'
                        f'<div style="color:#f97316;font-size:.85rem;min-width:40px">{feu_j}</div>'
                        f'</div>',
                        unsafe_allow_html=True
                    )

                st.markdown(
                    '<div style="font-size:.75rem;color:#475569;margin-top:8px">'
                    '🔒 Seuls les codes élèves sont affichés — aucun nom réel.</div>',
                    unsafe_allow_html=True
                )

# ============================================================
# PIED DE PAGE (CRÉDITS & LICENCE)
# ============================================================
st.divider()

st.markdown(f"""
    <div style="text-align: center; color: #888; font-size: 0.8rem; padding: 20px;">
        Conçu et développé par <b>Fabrice GUZZINATI</b> & <b>Gemini&Claude</b> (Architecte IA)<br>
        Version 1.5 — 2026 • Ozoir-la-Ferrière<br>
        <br>
        <i>Distribué sous licence <b>Creative Commons BY-NC-SA 4.0</b></i>
    </div>
    """, unsafe_allow_html=True)
